from __future__ import annotations

import csv
import json
import re
import shutil
import uuid
from collections import Counter, defaultdict
from copy import deepcopy
from datetime import date, datetime, timedelta
from io import BytesIO
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document
from docx.shared import Pt

from app.database.connection import DatabaseManager
from app.repositories.service_mapping_repository import ServiceMappingRepository
from app.core.input_layer import (
    OfficialMessageParser,
    aplicar_regra_primeira_equipe,
    carregar_dicionario_servicos,
    extrair_primeira_equipe,
    normalizar_texto,
)
from app.core.nucleo_master import (
    get_nucleo_profile,
    load_nucleo_registry,
    reconcile_parsed_with_registry,
    reconcile_rows_with_registry,
    save_nucleo_registry,
    split_registry_text,
)
from app.services.base_builder import build_management_workbook
from app.services.master_builder import update_master_from_output
from app.services.report_system import ReportGenerator, ServiceDictionary, WhatsAppReportParser, save_parsed_outputs


SERVICE_ALIAS_BOOTSTRAP: dict[str, dict[str, object]] = {
    "PRA": {
        "categoria": "rede_agua",
        "unidade_padrao": "m",
        "aliases": [
            "PRA",
        ],
    },
    "PRE": {
        "categoria": "rede_esgoto",
        "unidade_padrao": "m",
        "aliases": [
            "PRE",
        ],
    },
    "rede_agua": {
        "categoria": "rede_agua",
        "unidade_padrao": "m",
        "aliases": [
            "rede_agua",
            "Rede de agua",
            "Rede agua",
            "Prolongamento de rede de agua",
            "Prolongamento de rede",
            "Prolongamento rede agua",
            "Prolongamento rede",
        ],
    },
    "rede_esgoto": {
        "categoria": "rede_esgoto",
        "unidade_padrao": "m",
        "aliases": [
            "rede_esgoto",
            "Execucao de rede de esgoto",
            "Rede de esgoto",
        ],
    },
    "ramal_agua": {
        "categoria": "rede_agua",
        "unidade_padrao": "un",
        "aliases": [
            "Execucao de ramais de agua",
            "Ramais de agua",
            "Ramal de agua",
        ],
    },
    "ramal_esgoto": {
        "categoria": "rede_esgoto",
        "unidade_padrao": "un",
        "aliases": [
            "Execucao de ramais de esgoto",
            "Ramais de esgoto",
            "Ramal de esgoto",
        ],
    },
    "furo_direcional": {
        "categoria": "rede_agua",
        "unidade_padrao": "m",
        "aliases": [
            "furo_direcional",
            "Furo direcional",
            "Execucao de furo direcional",
        ],
    },
    "intradomiciliar": {
        "categoria": "intradomiciliar",
        "unidade_padrao": "un",
        "aliases": [
            "Ligacoes intradomiciliares",
            "Ligacao intradomiciliar",
            "Intradomiciliar",
        ],
    },
    "hidrometro": {
        "categoria": "hidrometro",
        "unidade_padrao": "un",
        "aliases": [
            "Instalacao de hidrometros",
            "Instalacao de hidrometro",
            "Hidrometro",
        ],
    },
    "caixa_uma": {
        "categoria": "caixa_uma",
        "unidade_padrao": "un",
        "aliases": [
            "Instalacao de caixas UMA",
            "Instalacao de caixa UMA",
            "Instalacao de caixas uma",
        ],
    },
    "interligacao": {
        "categoria": "rede_agua",
        "unidade_padrao": "un",
        "aliases": [
            "Execucao de interligacao de rede",
            "Execucao de interligacao de rede de agua",
            "Interligacao de rede",
            "Interligacao",
        ],
    },
    "adesao_agua": {
        "categoria": "adesao_agua",
        "unidade_padrao": "un",
        "aliases": [
            "Adesoes realizadas",
            "Adesao de agua",
            "Adesao agua",
        ],
    },
    "caixa_dagua": {
        "categoria": "caixa_dagua",
        "unidade_padrao": "un",
        "aliases": [
            "Caixa d'água",
            "Instalacao de caixa dagua",
            "Instalacao de caixa d'agua",
            "Instalacao de caixa de agua",
            "Caixa dagua",
            "Caixa dagua instalada",
            "Caixas dagua instaladas",
            "Caixas d'agua instaladas",
            "Caixa dgua instalada",
            "Caixa d agua instalada",
        ],
    },
    "Economias": {
        "categoria": "economias",
        "unidade_padrao": "un",
        "aliases": [
            "Economias",
            "economias",
        ],
    },
    "caixa_inspecao": {
        "categoria": "esgoto",
        "unidade_padrao": "un",
        "aliases": [
            "Instalacao de caixa de inspecao",
            "Caixa de inspecao",
        ],
    },
    "recomposicao_passeio": {
        "categoria": "recomposicao",
        "unidade_padrao": "m",
        "aliases": [
            "Recomposicao de valas",
            "Recomposicao de vala",
            "Recomposicao",
        ],
    },
    "concretagem_vala": {
        "categoria": "recomposicao",
        "unidade_padrao": "m",
        "aliases": [
            "Concretagem de vala realizada",
            "Concretagem de vala",
        ],
    },
    "caps": {
        "categoria": "insumo",
        "unidade_padrao": "un",
        "aliases": [
            "Caps utilizados",
            "Capeamento de rede",
            "Capeamento",
        ],
    },
    "corte_pavimento": {
        "categoria": "escavacao",
        "unidade_padrao": "m",
        "aliases": [
            "Corte de pavimento com serra clip",
            "Corte de pavimento",
            "Serra clip",
        ],
    },
    "retirada_entulho": {
        "categoria": "apoio",
        "unidade_padrao": "un",
        "aliases": [
            "Retirada de Big Bag",
            "Retirada de entulho",
            "Big Bag",
        ],
    },
    "luvas": {
        "categoria": "insumo",
        "unidade_padrao": "un",
        "aliases": [
            "Luvas executadas",
            "Luvas",
        ],
    },
    "manutencao_poco_inspecao": {
        "categoria": "manutencao",
        "unidade_padrao": "un",
        "aliases": [
            "Manutencao de pocos de inspecao",
            "Manutencao poco inspecao",
            "Manutencao de poco de inspecao",
            "Poco de inspecao",
        ],
    },
    "servico_complementar": {
        "categoria": "apoio",
        "unidade_padrao": "un",
        "aliases": [
            "Execucao de e degraus",
            "Retrabalhos em ramais",
            "Retrabalho em ramais",
        ],
    },
}


class WebPipelineService:
    def __init__(
        self,
        base_dir: Path,
        outputs_root: Path | None = None,
        master_dir: Path | None = None,
        history_file: Path | None = None,
        draft_dir: Path | None = None,
        nucleo_reference_file: Path | None = None,
        db_manager: DatabaseManager | None = None,
    ):
        self.base_dir = Path(base_dir)
        self.outputs_root = Path(outputs_root) if outputs_root else self.base_dir / "saidas"
        self.master_dir = Path(master_dir) if master_dir else self.base_dir / "BASE_MESTRA"
        self.history_file = (
            Path(history_file)
            if history_file
            else self.base_dir / "data" / "runtime" / "processing_history.csv"
        )
        self.draft_dir = Path(draft_dir) if draft_dir else self.base_dir / "data" / "drafts" / "web"
        self.unmapped_candidates_file = (
            Path(history_file).parent if history_file else self.base_dir / "data" / "runtime"
        ) / "unmapped_alias_candidates.csv"

        self.outputs_root.mkdir(parents=True, exist_ok=True)
        self.master_dir.mkdir(parents=True, exist_ok=True)
        self.history_file.parent.mkdir(parents=True, exist_ok=True)
        self.draft_dir.mkdir(parents=True, exist_ok=True)

        self.dictionary_csv = self.base_dir / "config" / "service_dictionary.csv"
        self.service_dictionary_v2_json = self.base_dir / "config" / "service_dictionary_v2.json"
        self.nucleo_reference_file = Path(nucleo_reference_file) if nucleo_reference_file else self.base_dir / "config" / "nucleo_reference.json"
        self.db_manager = db_manager
        self.service_mapping_repository: ServiceMappingRepository | None = None

        self.legacy_dictionary = ServiceDictionary(self.dictionary_csv)
        self.legacy_parser = WhatsAppReportParser(self.legacy_dictionary)
        self.report_generator = ReportGenerator(self.legacy_dictionary)

        self.official_parser = None
        if self.service_dictionary_v2_json.exists():
            self.official_parser = OfficialMessageParser(carregar_dicionario_servicos(self.service_dictionary_v2_json))

        self.nucleo_reference = self._load_nucleo_reference()
        self.service_catalog = self._build_service_catalog()

    def set_database_manager(self, db_manager: DatabaseManager | None) -> None:
        self.db_manager = db_manager
        self.nucleo_reference = self._load_nucleo_reference()

    def set_service_mapping_repository(self, repository: ServiceMappingRepository | None) -> None:
        self.service_mapping_repository = repository
        self.service_catalog = self._build_service_catalog()

    def refresh_service_catalog(self) -> None:
        self.service_catalog = self._build_service_catalog()

    def parse_message(self, raw_message: str, source_name: str | None = None) -> Tuple[dict, str]:
        source_name = source_name or f"mensagem_web_{datetime.now():%Y%m%d_%H%M%S}.txt"
        parser_mode = "legado"
        parsed = None

        if self.official_parser:
            parsed = self.official_parser.parse_text(raw_message, source_name=source_name)
            if parsed is not None:
                parser_mode = "oficial"

        if parsed is None:
            parsed = self.legacy_parser.parse_text(raw_message, source_name=source_name)

        parsed.setdefault("frentes", [])
        parsed.setdefault("execucao", [])
        parsed.setdefault("ocorrencias", [])
        parsed.setdefault("observacoes", [])
        parsed.setdefault("servicos_nao_mapeados", [])
        parsed.setdefault("data_referencia", "")
        aplicar_regra_primeira_equipe(parsed)
        return parsed, parser_mode

    def _normalize_nucleo_key(self, value: object) -> str:
        return normalizar_texto(str(value or "")).strip()

    def _compact_normalized_key(self, value: object) -> str:
        key = normalizar_texto(str(value or "")).strip()
        if not key:
            return ""
        return re.sub(r"(.)\1+", r"\1", key)

    def _split_multi_text(self, value: object) -> List[str]:
        text = str(value or "").replace("\r\n", "\n").replace("\r", "\n").strip()
        if not text:
            return []

        parts = re.split(r"\s*(?:/|;|\n)+\s*", text)
        out: List[str] = []
        seen = set()
        for part in parts:
            clean = re.sub(r"^[\u2022\-\*]+\s*", "", str(part or "")).strip(" -")
            clean = " ".join(clean.split()).strip()
            if not clean:
                continue
            key = normalizar_texto(clean)
            if not key or key in seen:
                continue
            seen.add(key)
            out.append(clean)
        return out

    def _read_nucleo_reference_payload_from_db(self) -> dict | None:
        if self.db_manager is None:
            return None

        try:
            with self.db_manager.connection() as conn:
                with conn.cursor() as cur:
                    cur.execute(
                        "SELECT payload_json FROM nucleo_registry_state WHERE id = 1"
                    )
                    row = cur.fetchone()
            if not row:
                return None
            payload = row[0]
            if isinstance(payload, str):
                payload = json.loads(payload)
            if not isinstance(payload, dict):
                return None
            version = str(payload.get("version", "2.0") or "2.0")
            raw_entries = payload.get("nucleos", [])
            if not isinstance(raw_entries, list):
                raw_entries = []
            return {"version": version, "nucleos": raw_entries}
        except Exception:
            return None

    def _write_nucleo_reference_payload_to_file(self, payload: dict) -> None:
        if not isinstance(payload, dict):
            return

        version = str(payload.get("version", "2.0") or "2.0")
        raw_entries = payload.get("nucleos", [])
        if not isinstance(raw_entries, list):
            raw_entries = []

        file_payload = {
            "version": version,
            "nucleos": raw_entries,
        }

        self.nucleo_reference_file.parent.mkdir(parents=True, exist_ok=True)
        with self.nucleo_reference_file.open("w", encoding="utf-8") as f:
            json.dump(file_payload, f, ensure_ascii=False, indent=2)

    def _sync_nucleo_reference_to_db(self, registry: dict) -> None:
        if self.db_manager is None:
            return

        entries = list((registry or {}).get("entries", []))
        version = str((registry or {}).get("version", "2.0") or "2.0")
        payload = {
            "version": version,
            "nucleos": entries,
        }

        try:
            with self.db_manager.connection() as conn:
                with conn.cursor() as cur:
                    cur.execute(
                        """
                        INSERT INTO nucleo_registry_state (id, version, payload_json, updated_at)
                        VALUES (1, %s, %s::jsonb, NOW())
                        ON CONFLICT (id)
                        DO UPDATE SET
                            version = EXCLUDED.version,
                            payload_json = EXCLUDED.payload_json,
                            updated_at = NOW()
                        """,
                        (version, json.dumps(payload, ensure_ascii=False)),
                    )
                conn.commit()
        except Exception:
            return

    def _load_nucleo_reference(self) -> dict:
        db_payload = self._read_nucleo_reference_payload_from_db()
        if db_payload is not None:
            try:
                self._write_nucleo_reference_payload_to_file(db_payload)
            except Exception:
                pass
            return load_nucleo_registry(self.nucleo_reference_file)

        registry = load_nucleo_registry(self.nucleo_reference_file)
        self._sync_nucleo_reference_to_db(registry)
        return registry

    def _get_nucleo_profile(self, nucleo: object) -> dict | None:
        return get_nucleo_profile(self.nucleo_reference, nucleo)

    def get_nucleo_reference_ui(self) -> dict:
        entries = list(self.nucleo_reference.get("entries", []))
        active_entries = list(self.nucleo_reference.get("active_entries", []) or entries)
        options = sorted(
            {
                str(e.get("nucleo", "") or "").strip()
                for e in active_entries
                if str(e.get("nucleo", "") or "").strip()
            }
        )

        profiles_by_key: Dict[str, dict] = {}
        for entry in entries:
            key = self._normalize_nucleo_key(entry.get("nucleo", ""))
            if not key:
                continue
            profiles_by_key[key] = {
                "nucleo": str(entry.get("nucleo", "") or "").strip(),
                "municipio": str(entry.get("municipio", "") or "").strip(),
                "status": str(entry.get("status", "ativo") or "ativo").strip(),
                "observacoes": str(entry.get("observacoes", "") or "").strip(),
                "logradouro_principal": str(entry.get("logradouro_principal", "") or "").strip(),
                "logradouros_padrao": list(entry.get("logradouros_padrao", [])),
                "equipes_padrao": list(entry.get("equipes_padrao", [])),
                "aliases": list(entry.get("aliases", [])),
                "created_at": str(entry.get("created_at", "") or "").strip(),
                "updated_at": str(entry.get("updated_at", "") or "").strip(),
            }

        return {
            "options": options,
            "entries": entries,
            "profiles_by_key": profiles_by_key,
            "total": len(options),
        }

    def list_nucleo_registry(self, search: str = "", status: str = "") -> List[dict]:
        search_norm = normalizar_texto(str(search or "").strip())
        status_norm = normalizar_texto(str(status or "").strip())
        if status_norm not in {"", "ativo", "inativo"}:
            status_norm = ""

        rows: List[dict] = []
        for entry in list(self.nucleo_reference.get("entries", [])):
            item = dict(entry)
            item["status"] = str(item.get("status", "ativo") or "ativo").strip() or "ativo"
            item["aliases_text"] = " | ".join(item.get("aliases", []))
            item["logradouros_padrao_text"] = " | ".join(item.get("logradouros_padrao", []))
            item["equipes_padrao_text"] = " | ".join(item.get("equipes_padrao", []))

            if status_norm and normalizar_texto(item.get("status", "")) != status_norm:
                continue

            if search_norm:
                haystack = " ".join(
                    [
                        str(item.get("nucleo", "") or ""),
                        str(item.get("municipio", "") or ""),
                        str(item.get("observacoes", "") or ""),
                        item["aliases_text"],
                        item["logradouros_padrao_text"],
                        item["equipes_padrao_text"],
                    ]
                )
                if search_norm not in normalizar_texto(haystack):
                    continue
            rows.append(item)

        rows.sort(key=lambda item: (normalizar_texto(item.get("nucleo", "")), normalizar_texto(item.get("municipio", ""))))
        return rows

    def get_nucleo_registry_form(self, nucleo_name: str = "", prefill: dict | None = None) -> dict:
        entry = None
        if nucleo_name:
            key = self._normalize_nucleo_key(nucleo_name)
            for row in self.nucleo_reference.get("entries", []):
                if self._normalize_nucleo_key(row.get("nucleo", "")) == key:
                    entry = dict(row)
                    break

        base = {
            "original_nucleo": "",
            "nucleo": "",
            "municipio": "",
            "status": "ativo",
            "aliases": "",
            "observacoes": "",
            "logradouro_principal": "",
            "logradouros_padrao": "",
            "equipes_padrao": "",
        }

        if entry:
            base.update(
                {
                    "original_nucleo": str(entry.get("nucleo", "") or "").strip(),
                    "nucleo": str(entry.get("nucleo", "") or "").strip(),
                    "municipio": str(entry.get("municipio", "") or "").strip(),
                    "status": str(entry.get("status", "ativo") or "ativo").strip() or "ativo",
                    "aliases": "\n".join(entry.get("aliases", [])),
                    "observacoes": str(entry.get("observacoes", "") or "").strip(),
                    "logradouro_principal": str(entry.get("logradouro_principal", "") or "").strip(),
                    "logradouros_padrao": "\n".join(entry.get("logradouros_padrao", [])),
                    "equipes_padrao": "\n".join(entry.get("equipes_padrao", [])),
                }
            )

        if prefill:
            for field in ("nucleo", "municipio", "logradouro_principal", "observacoes"):
                value = str(prefill.get(field, "") or "").strip()
                if value and not str(base.get(field, "") or "").strip():
                    base[field] = value
        return base

    def upsert_nucleo_registry_entry(self, payload: Dict[str, object]) -> tuple[dict | None, str]:
        data = dict(payload or {})
        nucleo = " ".join(str(data.get("nucleo", "") or "").split()).strip()
        municipio = " ".join(str(data.get("municipio", "") or "").split()).strip()
        original_nucleo = " ".join(str(data.get("original_nucleo", "") or "").split()).strip()
        status = str(data.get("status", "ativo") or "ativo").strip().lower()
        observacoes = " ".join(str(data.get("observacoes", "") or "").split()).strip()
        logradouro_principal = " ".join(str(data.get("logradouro_principal", "") or "").split()).strip()

        aliases = split_registry_text(data.get("aliases", ""))
        logradouros_padrao = split_registry_text(data.get("logradouros_padrao", ""))
        equipes_padrao = self._normalize_team_list(split_registry_text(data.get("equipes_padrao", "")))

        if logradouro_principal:
            log_key = normalizar_texto(logradouro_principal)
            if all(normalizar_texto(x) != log_key for x in logradouros_padrao):
                logradouros_padrao.insert(0, logradouro_principal)

        if not nucleo:
            return None, "Informe o nome oficial do nucleo."
        if not municipio:
            return None, "Informe o municipio oficial do nucleo."
        if status not in {"ativo", "inativo"}:
            status = "ativo"

        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        existing_entries = [dict(x) for x in self.nucleo_reference.get("entries", [])]
        original_key = self._normalize_nucleo_key(original_nucleo or nucleo)
        new_key = self._normalize_nucleo_key(nucleo)

        collision_values = {nucleo, *aliases}
        for entry in existing_entries:
            entry_key = self._normalize_nucleo_key(entry.get("nucleo", ""))
            if original_key and entry_key == original_key:
                continue
            for candidate in [entry.get("nucleo", ""), *entry.get("aliases", [])]:
                cand_key = self._normalize_nucleo_key(candidate)
                if not cand_key:
                    continue
                if any(self._normalize_nucleo_key(v) == cand_key for v in collision_values if str(v or "").strip()):
                    return None, "Existe conflito entre nome/alias informado e um nucleo ja cadastrado."

        updated = False
        merged_entries: List[dict] = []
        for entry in existing_entries:
            entry_key = self._normalize_nucleo_key(entry.get("nucleo", ""))
            if original_key and entry_key == original_key:
                created_at = str(entry.get("created_at", "") or "").strip() or now
                merged_entries.append(
                    {
                        "nucleo": nucleo,
                        "municipio": municipio,
                        "status": status,
                        "aliases": aliases,
                        "observacoes": observacoes,
                        "logradouro_principal": logradouro_principal,
                        "logradouros_padrao": logradouros_padrao,
                        "equipes_padrao": equipes_padrao,
                        "created_at": created_at,
                        "updated_at": now,
                    }
                )
                updated = True
            else:
                merged_entries.append(entry)

        if not updated:
            merged_entries.append(
                {
                    "nucleo": nucleo,
                    "municipio": municipio,
                    "status": status,
                    "aliases": aliases,
                    "observacoes": observacoes,
                    "logradouro_principal": logradouro_principal,
                    "logradouros_padrao": logradouros_padrao,
                    "equipes_padrao": equipes_padrao,
                    "created_at": now,
                    "updated_at": now,
                }
            )

        self.nucleo_reference = save_nucleo_registry(self.nucleo_reference_file, merged_entries)
        self._sync_nucleo_reference_to_db(self.nucleo_reference)
        saved_entry = None
        for entry in self.nucleo_reference.get("entries", []):
            if self._normalize_nucleo_key(entry.get("nucleo", "")) == new_key:
                saved_entry = dict(entry)
                break

        return saved_entry, ""


    def _build_service_catalog(self) -> dict:
        catalog_by_servico: Dict[str, dict] = {}

        if self.official_parser and getattr(self.official_parser, "service_dictionary", None):
            entries = getattr(self.official_parser.service_dictionary, "entries", [])
            for entry in entries:
                servico = str(entry.get("servico_oficial", "") or "").strip()
                if not servico:
                    continue
                categoria = str(entry.get("categoria", "") or "").strip() or "servico_nao_mapeado"
                if servico == "hidrometro":
                    categoria = "hidrometro"
                label = servico.replace("_", " ").strip()
                catalog_by_servico[servico] = {
                    "servico": servico,
                    "categoria": categoria,
                    "label": label,
                }

        if not catalog_by_servico and self.dictionary_csv.exists():
            rows: List[dict] = []
            for enc in ("utf-8-sig", "utf-8", "latin-1"):
                try:
                    with self.dictionary_csv.open("r", encoding=enc, newline="") as f:
                        rows = list(csv.DictReader(f))
                    break
                except UnicodeDecodeError:
                    continue
                except Exception:
                    rows = []
                    break

            for row in rows:
                servico = str(
                    row.get("nome_padronizado", "")
                    or row.get("frase_tecnica", "")
                    or row.get("nome_original", "")
                    or ""
                ).strip()
                if not servico:
                    continue
                categoria = str(row.get("categoria", "") or "").strip() or "servico_nao_mapeado"
                if normalizar_texto(servico) == "hidrometro":
                    categoria = "hidrometro"
                catalog_by_servico.setdefault(
                    servico,
                    {
                        "servico": servico,
                        "categoria": categoria,
                        "label": servico,
                    },
                )

        if self.service_mapping_repository is not None:
            try:
                db_services = self.service_mapping_repository.list_services(limit=2000)
            except Exception:
                db_services = []
            for row in db_services:
                if not bool(row.get("ativo", True)):
                    continue
                servico = str(row.get("servico_oficial", "") or "").strip()
                if not servico:
                    continue
                categoria = str(row.get("categoria", "") or "").strip() or "servico_nao_mapeado"
                if normalizar_texto(servico) == "hidrometro":
                    categoria = "hidrometro"
                label = servico.replace("_", " ").strip()
                catalog_by_servico[servico] = {
                    "servico": servico,
                    "categoria": categoria,
                    "label": label,
                }

        options = sorted(
            catalog_by_servico.values(),
            key=lambda x: (str(x.get("categoria", "") or "").lower(), str(x.get("label", "") or "").lower()),
        )

        by_key: Dict[str, dict] = {}
        for item in options:
            servico = str(item.get("servico", "") or "").strip()
            if not servico:
                continue
            by_key[servico] = item

        return {
            "options": options,
            "by_servico": by_key,
        }

    def get_service_catalog_ui(self) -> dict:
        return {
            "options": list(self.service_catalog.get("options", [])),
            "total": len(self.service_catalog.get("options", [])),
        }

    def _manual_service_meta(self, servico: object) -> dict:
        key = str(servico or "").strip()
        entry = self.service_catalog.get("by_servico", {}).get(key)
        if entry:
            categoria = str(entry.get("categoria", "") or "").strip() or "servico_nao_mapeado"
            if key == "hidrometro":
                categoria = "hidrometro"
            return {
                "servico": key,
                "categoria": categoria,
                "label": str(entry.get("label", key) or key),
            }
        categoria = "hidrometro" if key == "hidrometro" else "servico_nao_mapeado"
        return {"servico": key, "categoria": categoria, "label": key}

    def list_registered_services(self, search: str = "", limit: int = 500) -> List[dict]:
        if self.service_mapping_repository is None:
            return []
        try:
            return list(self.service_mapping_repository.list_services(search=search, limit=limit))
        except Exception:
            return []

    def list_registered_aliases(self, search: str = "", limit: int = 500) -> List[dict]:
        if self.service_mapping_repository is None:
            return []
        try:
            return list(self.service_mapping_repository.list_aliases(search=search, limit=limit))
        except Exception:
            return []

    def _suggest_service_for_unmapped_term(self, term: object) -> str:
        raw = str(term or "").strip()
        if not raw:
            return ""
        norm = normalizar_texto(raw)
        norm = re.sub(r"[^a-z0-9]+", " ", norm).strip()
        if not norm:
            return ""

        if norm == "pra":
            return "PRA"
        if norm == "pre":
            return "PRE"
        if "prolongamento" in norm and "rede" in norm:
            return "rede_agua"
        if "ramal" in norm and "esgoto" in norm:
            return "ramal_esgoto"
        if "ramal" in norm and "agua" in norm:
            return "ramal_agua"
        if "rede" in norm and "esgoto" in norm:
            return "rede_esgoto"
        if "rede" in norm and "agua" in norm:
            return "rede_agua"
        if "intradomiciliar" in norm or ("ligacao" in norm and "domic" in norm):
            return "intradomiciliar"
        if "hidrometr" in norm:
            return "hidrometro"
        if "caixa" in norm and "uma" in norm:
            return "caixa_uma"
        if "interlig" in norm and "rede" in norm:
            return "interligacao"
        if "furo" in norm and "direcional" in norm:
            return "furo_direcional"
        if "ades" in norm:
            return "adesao_agua"
        if "economia" in norm:
            return "Economias"
        if "manutencao" in norm and "poco" in norm and "inspec" in norm:
            return "manutencao_poco_inspecao"
        if "poco" in norm and "inspec" in norm:
            return "manutencao_poco_inspecao"
        if "caixa" in norm and "inspec" in norm:
            return "caixa_inspecao"
        if "caixa" in norm and (
            "dagua" in norm
            or ("de" in norm and "agua" in norm)
            or bool(re.search(r"\bd\s+agua\b", norm))
        ):
            return "caixa_dagua"
        if "caps" in norm or "capeamento" in norm:
            return "caps"
        if "serra clip" in norm or ("corte" in norm and "pavimento" in norm):
            return "corte_pavimento"
        if "big bag" in norm or ("retirada" in norm and "entulho" in norm):
            return "retirada_entulho"
        if "luvas" in norm:
            return "luvas"
        if "degraus" in norm or ("retrabalho" in norm and "ramal" in norm):
            return "servico_complementar"
        if "recompos" in norm:
            return "recomposicao_passeio"
        if "concretagem" in norm and "vala" in norm:
            return "concretagem_vala"
        return ""

    def _sanitize_unmapped_alias_term(self, term: object) -> str:
        raw = str(term or "").strip()
        if not raw:
            return ""
        clean = self._sanitize_institutional_text(raw)
        clean = re.sub(r"(?i)^(item|servico|serviço)\s*[:\\-]\s*", "", clean).strip(" -")
        clean = re.sub(r"\s+", " ", clean).strip()
        if not clean:
            return ""
        norm = normalizar_texto(clean)
        if not norm:
            return ""
        generic_noise = {"ok", "na", "n a", "sem", "com", "de", "servico"}
        if norm in generic_noise:
            return ""
        return clean

    def bootstrap_service_aliases(
        self,
        *,
        max_terms: int = 2000,
        min_count: int = 1,
    ) -> dict:
        if self.service_mapping_repository is None:
            raise RuntimeError("Cadastro de aliases indisponivel no momento.")

        repo = self.service_mapping_repository
        created_services = 0
        upserted_aliases = 0

        service_ids: Dict[str, int] = {}
        for servico_oficial, meta in SERVICE_ALIAS_BOOTSTRAP.items():
            row = repo.upsert_service(
                servico_oficial=servico_oficial,
                categoria=str(meta.get("categoria", "") or "").strip() or "servico_nao_mapeado",
                unidade_padrao=str(meta.get("unidade_padrao", "") or "").strip(),
                ativo=True,
            )
            if int(row.get("id", 0) or 0):
                service_ids[servico_oficial] = int(row.get("id", 0) or 0)
                created_services += 1

        for servico_oficial, meta in SERVICE_ALIAS_BOOTSTRAP.items():
            service_id = service_ids.get(servico_oficial, 0)
            if service_id <= 0:
                continue
            seen_alias_norms: set[str] = set()
            for alias in list(meta.get("aliases", []) or []):
                alias_clean = str(alias or "").strip()
                if not alias_clean:
                    continue
                alias_norm = normalizar_texto(alias_clean)
                if not alias_norm or alias_norm in seen_alias_norms:
                    continue
                seen_alias_norms.add(alias_norm)
                repo.upsert_alias(alias_clean, service_id, source="bootstrap_catalog", ativo=True)
                upserted_aliases += 1

        aliases_cleaned = 0
        aliases_deactivated = 0
        try:
            existing_aliases = repo.list_aliases(limit=5000)
        except Exception:
            existing_aliases = []
        for alias_row in existing_aliases:
            if str(alias_row.get("source", "") or "").strip() != "bootstrap_unmapped":
                continue
            alias_id = int(alias_row.get("id", 0) or 0)
            service_id = int(alias_row.get("service_id", 0) or 0)
            alias_text = str(alias_row.get("alias_text", "") or "").strip()
            if alias_id <= 0 or service_id <= 0 or not alias_text:
                continue

            cleaned_text = self._sanitize_unmapped_alias_term(alias_text)
            if not cleaned_text:
                if repo.set_alias_active(alias_id, False):
                    aliases_deactivated += 1
                continue

            if cleaned_text == alias_text:
                continue

            upserted = repo.upsert_alias(
                cleaned_text,
                service_id,
                source="bootstrap_unmapped",
                ativo=True,
            )
            aliases_cleaned += 1
            old_norm = normalizar_texto(alias_text)
            new_norm = normalizar_texto(cleaned_text)
            new_id = int(upserted.get("id", 0) or 0)
            if old_norm != new_norm and new_id and new_id != alias_id:
                if repo.set_alias_active(alias_id, False):
                    aliases_deactivated += 1

        unmapped_terms = repo.list_unmapped_terms_from_management(limit=max_terms, min_count=min_count)
        auto_mapped_terms = 0
        for item in unmapped_terms:
            term = self._sanitize_unmapped_alias_term(item.get("termo", ""))
            if not term:
                continue
            suggested = self._suggest_service_for_unmapped_term(term)
            if not suggested:
                continue
            service_id = service_ids.get(suggested, 0)
            if service_id <= 0:
                continue
            repo.upsert_alias(term, service_id, source="bootstrap_unmapped", ativo=True)
            auto_mapped_terms += 1

        remap_stats = repo.remap_management_execucao_by_aliases(only_unmapped=False)
        self.refresh_service_catalog()
        return {
            "services_upserted": created_services,
            "aliases_upserted": upserted_aliases,
            "unmapped_terms_seen": len(unmapped_terms),
            "unmapped_terms_auto_mapped": auto_mapped_terms,
            "aliases_cleaned": aliases_cleaned,
            "aliases_deactivated": aliases_deactivated,
            "rows_updated": int(remap_stats.get("rows_updated", 0) or 0),
            "rows_scanned": int(remap_stats.get("rows_scanned", 0) or 0),
            "aliases_loaded": int(remap_stats.get("aliases_loaded", 0) or 0),
        }

    def ensure_registered_service(self, servico: object, categoria: object = "", unidade_padrao: object = "") -> dict:
        if self.service_mapping_repository is None:
            raise RuntimeError("Cadastro de servicos indisponivel no momento.")
        meta = self._manual_service_meta(servico)
        servico_oficial = str(meta.get("servico", "") or "").strip()
        if not servico_oficial:
            raise ValueError("Informe um servico oficial valido.")
        categoria_final = str(categoria or "").strip() or str(meta.get("categoria", "") or "").strip() or "servico_nao_mapeado"
        service_row = self.service_mapping_repository.upsert_service(
            servico_oficial,
            categoria_final,
            str(unidade_padrao or "").strip(),
            ativo=True,
        )
        self.refresh_service_catalog()
        return service_row

    def register_service_alias(
        self,
        alias_text: object,
        servico_oficial: object,
        categoria: object = "",
        unidade_padrao: object = "",
        source: object = "manual",
    ) -> dict:
        if self.service_mapping_repository is None:
            raise RuntimeError("Cadastro de aliases indisponivel no momento.")
        service_row = self.ensure_registered_service(servico_oficial, categoria, unidade_padrao)
        alias_clean = str(alias_text or "").strip()
        if not alias_clean:
            raise ValueError("Informe um termo para mapear.")
        alias_row = self.service_mapping_repository.upsert_alias(
            alias_clean,
            int(service_row.get("id", 0) or 0),
            source=str(source or "manual").strip() or "manual",
            ativo=True,
        )
        remap_stats = self.service_mapping_repository.remap_management_execucao_by_aliases(
            only_unmapped=False
        )
        return {"service": service_row, "alias": alias_row, "remap": remap_stats}

    def apply_registered_service_aliases(self, parsed: dict) -> dict:
        if self.service_mapping_repository is None:
            return parsed
        try:
            alias_map = self.service_mapping_repository.list_alias_map()
        except Exception:
            return parsed
        if not alias_map:
            return parsed

        timestamp = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        applied = []

        def _candidate_terms(row_data: dict) -> List[str]:
            values = [
                row_data.get("servico_oficial", ""),
                row_data.get("servico_normalizado", ""),
                row_data.get("servico_bruto", ""),
                row_data.get("item_normalizado", ""),
                row_data.get("item_original", ""),
            ]
            out: List[str] = []
            seen = set()
            for value in values:
                normalized = normalizar_texto(str(value or "")).strip()
                if not normalized or normalized in seen:
                    continue
                seen.add(normalized)
                out.append(normalized)
            return out

        for idx, item in enumerate(parsed.get("execucao", [])):
            if not isinstance(item, dict):
                continue
            match = None
            for term in _candidate_terms(item):
                candidate = alias_map.get(term)
                if candidate:
                    match = candidate
                    break
            if not match:
                continue

            target_service = str(match.get("servico_oficial", "") or "").strip()
            target_category = str(match.get("categoria", "") or "").strip() or "servico_nao_mapeado"
            current_service = str(item.get("servico_oficial", "") or "").strip()
            current_category_item = str(item.get("categoria_item", "") or "").strip()
            current_category = str(item.get("categoria", "") or "").strip()
            if (
                current_service == target_service
                and current_category_item == target_category
                and current_category == target_category
            ):
                continue

            item.setdefault("servico_original_bruto", str(item.get("servico_bruto", item.get("item_original", "")) or "").strip())
            item.setdefault("servico_original_normalizado", str(item.get("servico_normalizado", item.get("item_normalizado", "")) or "").strip())
            item["servico_oficial"] = target_service
            item["item_normalizado"] = target_service
            item["categoria_item"] = target_category
            item["categoria"] = target_category
            item["regra_disparada"] = "alias_cadastrado"
            item["corrigido_alias"] = "sim"
            item["servico_corrigido_alias"] = target_service
            item["categoria_corrigida_alias"] = target_category
            item["alias_usado"] = str(match.get("alias_text", "") or "").strip()
            item["data_correcao_alias"] = timestamp
            applied.append(
                {
                    "id_item": str(item.get("id_item", "") or "").strip(),
                    "servico": item["servico_oficial"],
                    "categoria": item["categoria_item"],
                    "alias": item["alias_usado"],
                }
            )

        for row in parsed.get("servicos_nao_mapeados", []):
            if not isinstance(row, dict):
                continue
            match = None
            for term in _candidate_terms(row):
                candidate = alias_map.get(term)
                if candidate:
                    match = candidate
                    break
            if not match:
                continue
            row["corrigido_alias"] = "sim"
            row["servico_corrigido_alias"] = str(match.get("servico_oficial", "") or "").strip()
            row["categoria_corrigida_alias"] = str(match.get("categoria", "") or "").strip() or "servico_nao_mapeado"
            row["alias_usado"] = str(match.get("alias_text", "") or "").strip()
            row["data_correcao_alias"] = timestamp

        if applied:
            parsed.setdefault("correcoes_aliases", [])
            parsed["correcoes_aliases"].extend(applied)

        return parsed
    def apply_nucleo_defaults(self, fields: Dict[str, object]) -> Tuple[Dict[str, object], dict]:
        raw_fields = fields or {}
        data: Dict[str, object] = {}
        for key, value in raw_fields.items():
            if key in {"data", "nucleo", "logradouro", "municipio", "equipe", "aplicar_todos"}:
                data[key] = str(value or "").strip()
            else:
                data[key] = value
        info = {
            "matched": False,
            "nucleo": "",
            "applied": {},
            "profile": {},
        }

        nucleo = data.get("nucleo", "")
        if not nucleo:
            return data, info

        profile = self._get_nucleo_profile(nucleo)
        if not profile:
            return data, info

        info["matched"] = True
        info["nucleo"] = str(profile.get("nucleo", "") or "").strip()
        info["profile"] = {
            "municipio": str(profile.get("municipio", "") or "").strip(),
            "logradouros_padrao": list(profile.get("logradouros_padrao", [])),
            "equipes_padrao": list(profile.get("equipes_padrao", [])),
        }

        def _apply_if_empty(field: str, value: str) -> None:
            clean = str(value or "").strip()
            if field == "equipe":
                clean = self._first_team(clean)
            if not clean:
                return
            if str(data.get(field, "") or "").strip():
                return
            data[field] = clean
            info["applied"][field] = clean

        _apply_if_empty("municipio", str(profile.get("municipio", "") or ""))

        logradouros_padrao = list(profile.get("logradouros_padrao", []))
        if len(logradouros_padrao) == 1:
            _apply_if_empty("logradouro", logradouros_padrao[0])

        equipes_padrao = list(profile.get("equipes_padrao", []))
        if len(equipes_padrao) == 1:
            _apply_if_empty("equipe", equipes_padrao[0])

        return data, info

    def reconcile_with_nucleo_master(self, parsed: dict) -> dict:
        return reconcile_parsed_with_registry(parsed, self.nucleo_reference)

    def build_main_lists(self, main_fields: dict, context_overview: dict | None = None) -> dict:
        ctx = context_overview or {}

        logradouros = self._split_multi_text(main_fields.get("logradouro", ""))
        if not logradouros:
            logradouros = list(ctx.get("logradouros", []))

        equipes = self._normalize_team_list([main_fields.get("equipe", "")])
        if not equipes:
            equipes = self._normalize_team_list(list(ctx.get("equipes", [])))

        municipios = self._split_multi_text(main_fields.get("municipio", ""))
        if not municipios:
            municipios = list(ctx.get("municipios", []))

        nucleos = self._split_multi_text(main_fields.get("nucleo", ""))
        if not nucleos:
            nucleos = list(ctx.get("nucleos", []))

        return {
            "nucleos": nucleos,
            "municipios": municipios,
            "logradouros": logradouros,
            "equipes": equipes,
            "nucleos_count": len(nucleos),
            "municipios_count": len(municipios),
            "logradouros_count": len(logradouros),
            "equipes_count": len(equipes),
        }

    def extract_main_fields(self, parsed: dict) -> dict:
        fr = parsed.get("frentes", [])
        ex = parsed.get("execucao", [])
        oc = parsed.get("ocorrencias", [])
        ob = parsed.get("observacoes", [])

        first_fr = fr[0] if fr else {}
        first_ex = ex[0] if ex else {}
        first_oc = oc[0] if oc else {}
        first_ob = ob[0] if ob else {}

        def first_non_empty(*values: object) -> str:
            for value in values:
                text = str(value or "").strip()
                if text:
                    return text
            return ""

        return {
            "data": first_non_empty(
                parsed.get("data_referencia", ""),
                first_fr.get("data_referencia", ""),
                first_ex.get("data_referencia", ""),
                first_oc.get("data_referencia", ""),
                first_ob.get("data", ""),
            ),
            "nucleo": first_non_empty(
                parsed.get("nucleo", ""),
                first_fr.get("nucleo", ""),
                first_ex.get("nucleo", ""),
                first_oc.get("nucleo", ""),
                first_ob.get("nucleo", ""),
            ),
            "logradouro": first_non_empty(
                parsed.get("logradouro", ""),
                first_fr.get("logradouro", ""),
                first_ex.get("logradouro", ""),
                first_oc.get("logradouro", ""),
                first_ob.get("logradouro", ""),
            ),
            "municipio": first_non_empty(
                parsed.get("municipio", ""),
                first_fr.get("municipio", ""),
                first_ex.get("municipio", ""),
                first_oc.get("municipio", ""),
                first_ob.get("municipio", ""),
            ),
            "equipe": first_non_empty(
                parsed.get("equipe", ""),
                first_fr.get("equipe", ""),
                first_ex.get("equipe", ""),
                first_oc.get("equipe", ""),
                first_ob.get("equipe", ""),
            ),
        }

    def _missing_main_fields(self, main_fields: dict) -> List[str]:
        obrigatorios = ["data", "nucleo", "logradouro", "municipio", "equipe"]
        return [f for f in obrigatorios if not str(main_fields.get(f, "") or "").strip()]

    def _is_true_flag(self, value: object) -> bool:
        raw = str(value or "").strip().lower()
        return raw in {"1", "true", "on", "yes", "sim", "s"}

    def _row_has_value(self, row: dict, field: str) -> bool:
        return bool(str(row.get(field, "") or "").strip())

    def _apply_value_to_rows(self, rows: List[dict], field: str, value: str, force_all: bool) -> None:
        for row in rows:
            if force_all or not self._row_has_value(row, field):
                row[field] = value

    def collect_context_overview(self, parsed: dict) -> dict:
        def _collect(field: str, split_multi: bool = False) -> List[str]:
            values = set()
            for bucket in ("frentes", "execucao", "ocorrencias", "observacoes"):
                for row in parsed.get(bucket, []):
                    text = str(row.get(field, "") or "").strip()
                    if not text:
                        continue
                    chunks = self._split_multi_text(text) if split_multi else [text]
                    for chunk in chunks:
                        clean = str(chunk or "").strip()
                        if clean:
                            values.add(clean)
            return sorted(values)

        nucleos = _collect("nucleo")
        logradouros = _collect("logradouro", split_multi=True)
        municipios = _collect("municipio")
        equipes_set = {}
        for bucket in ("frentes", "execucao", "ocorrencias", "observacoes"):
            for row in parsed.get(bucket, []):
                equipe = self._first_team(row.get("equipe", ""))
                if not equipe:
                    continue
                key = normalizar_texto(equipe)
                if key and key not in equipes_set:
                    equipes_set[key] = equipe
        equipes = sorted(equipes_set.values())
        return {
            "nucleos": nucleos,
            "logradouros": logradouros,
            "municipios": municipios,
            "equipes": equipes,
            "nucleos_count": len(nucleos),
            "logradouros_count": len(logradouros),
            "municipios_count": len(municipios),
            "equipes_count": len(equipes),
        }

    def build_nucleo_groups(self, parsed: dict) -> List[dict]:
        grouped: Dict[str, dict] = {}

        def _add_unique(target: List[str], value: str) -> None:
            clean = str(value or "").strip()
            if not clean:
                return
            key = normalizar_texto(clean)
            if any(normalizar_texto(existing) == key for existing in target):
                return
            target.append(clean)

        for bucket in ("frentes", "execucao", "ocorrencias", "observacoes"):
            for row in parsed.get(bucket, []):
                nucleo = str(row.get("nucleo_oficial", "") or row.get("nucleo", "") or "").strip()
                if not nucleo:
                    continue

                group = grouped.setdefault(
                    nucleo,
                    {
                        "nucleo": nucleo,
                        "municipio": "",
                        "logradouros": [],
                        "equipes": [],
                        "nucleo_oficial": str(row.get("nucleo_oficial", "") or "").strip(),
                        "municipio_oficial": str(row.get("municipio_oficial", "") or "").strip(),
                        "status_cadastro": str(row.get("nucleo_status_cadastro", "") or "").strip(),
                    },
                )

                municipio = self._resolve_single_municipio_for_nucleo(
                    nucleo,
                    explicit_oficial=row.get("municipio_oficial", ""),
                    detected_values=self._split_multi_text(row.get("municipio", "")),
                )
                if municipio:
                    group["municipio"] = municipio
                for logradouro in self._split_multi_text(row.get("logradouro", "")):
                    _add_unique(group["logradouros"], logradouro)
                equipe = self._first_team(row.get("equipe", ""))
                if equipe:
                    _add_unique(group["equipes"], equipe)

        out: List[dict] = []
        for nucleo in sorted(grouped.keys()):
            item = grouped[nucleo]
            logradouros = item["logradouros"]
            equipes = item["equipes"]
            out.append(
                {
                    "nucleo": nucleo,
                    "municipio": str(item.get("municipio", "") or "").strip(),
                    "logradouro": " / ".join(logradouros),
                    "equipe": " / ".join(equipes),
                    "logradouros": logradouros,
                    "equipes": equipes,
                    "nucleo_oficial": str(item.get("nucleo_oficial", "") or "").strip(),
                    "municipio_oficial": str(item.get("municipio_oficial", "") or "").strip(),
                    "status_cadastro": str(item.get("status_cadastro", "") or "").strip(),
                    "cadastrado": str(item.get("status_cadastro", "") or "").strip() == "cadastrado",
                }
            )

        return out
    def apply_overrides(self, parsed: dict, overrides: Dict[str, object], force_all: bool | None = None) -> dict:
        data = aplicar_regra_primeira_equipe(deepcopy(parsed))
        raw_overrides = overrides or {}
        raw_nucleo_overrides = raw_overrides.get("nucleo_overrides", [])
        raw_manual_corrections = raw_overrides.get("manual_service_corrections", [])
        values = {
            k: str(v or "").strip()
            for k, v in raw_overrides.items()
            if k not in {"nucleo_overrides", "manual_service_corrections"}
        }
        values["equipe"] = self._first_team(values.get("equipe", ""))

        nucleo_overrides: List[dict] = []
        if isinstance(raw_nucleo_overrides, str):
            raw = str(raw_nucleo_overrides or "").strip()
            if raw:
                try:
                    loaded = json.loads(raw)
                    if isinstance(loaded, list):
                        nucleo_overrides = [x for x in loaded if isinstance(x, dict)]
                except Exception:
                    nucleo_overrides = []
        elif isinstance(raw_nucleo_overrides, list):
            nucleo_overrides = [x for x in raw_nucleo_overrides if isinstance(x, dict)]

        manual_corrections: List[dict] = []
        if isinstance(raw_manual_corrections, str):
            raw = str(raw_manual_corrections or "").strip()
            if raw:
                try:
                    loaded = json.loads(raw)
                    if isinstance(loaded, list):
                        manual_corrections = [x for x in loaded if isinstance(x, dict)]
                except Exception:
                    manual_corrections = []
        elif isinstance(raw_manual_corrections, list):
            manual_corrections = [x for x in raw_manual_corrections if isinstance(x, dict)]

        if force_all is None:
            force_all = self._is_true_flag(values.get("aplicar_todos", ""))

        data_value = values.get("data", "")
        if data_value:
            data["data_referencia"] = data_value
            for row in data.get("frentes", []):
                row["data_referencia"] = data_value
                row["data"] = data_value
            for row in data.get("execucao", []):
                row["data_referencia"] = data_value
                row["data"] = data_value
            for row in data.get("ocorrencias", []):
                row["data_referencia"] = data_value
                row["data"] = data_value
            for row in data.get("observacoes", []):
                row["data"] = data_value

        for field in ("nucleo", "logradouro", "municipio", "equipe"):
            value = values.get(field, "")
            if not value:
                continue
            self._apply_value_to_rows(data.get("frentes", []), field, value, force_all)
            self._apply_value_to_rows(data.get("execucao", []), field, value, force_all)
            self._apply_value_to_rows(data.get("ocorrencias", []), field, value, force_all)
            self._apply_value_to_rows(data.get("observacoes", []), field, value, force_all)
            if force_all or not str(data.get(field, "") or "").strip():
                data[field] = value

        if nucleo_overrides:
            def _match_nucleo(row_nucleo: object, target_nucleo: object) -> bool:
                row_key = normalizar_texto(str(row_nucleo or "").strip())
                target_key = normalizar_texto(str(target_nucleo or "").strip())
                return bool(row_key) and bool(target_key) and row_key == target_key

            for row_override in nucleo_overrides:
                nucleo_name = str(row_override.get("nucleo", "") or "").strip()
                if not nucleo_name:
                    continue

                municipio = str(row_override.get("municipio", "") or "").strip()
                logradouro = str(row_override.get("logradouro", "") or "").strip()
                equipe = self._first_team(row_override.get("equipe", ""))

                for bucket in ("frentes", "execucao", "ocorrencias", "observacoes"):
                    for row in data.get(bucket, []):
                        if not _match_nucleo(row.get("nucleo", ""), nucleo_name):
                            continue
                        if municipio:
                            row["municipio"] = municipio
                        if logradouro:
                            row["logradouro"] = logradouro
                        if equipe:
                            row["equipe"] = equipe

        self.apply_registered_service_aliases(data)

        if manual_corrections:
            self.apply_manual_service_corrections(data, manual_corrections)

        return data

    def bind_contract_to_parsed(self, parsed: dict, contract_value: object, force: bool = True) -> dict:
        value = str(contract_value or "").strip()
        if not value:
            return parsed

        parsed["contrato"] = value
        for bucket in ("frentes", "execucao", "ocorrencias", "observacoes", "servicos_nao_mapeados"):
            rows = parsed.get(bucket, [])
            if not isinstance(rows, list):
                continue
            for row in rows:
                if not isinstance(row, dict):
                    continue
                if force or not str(row.get("contrato", "") or "").strip():
                    row["contrato"] = value
        return parsed

    def apply_manual_service_corrections(self, parsed: dict, corrections: List[dict]) -> dict:
        prepared: Dict[str, dict] = {}
        prepared_by_suffix: Dict[str, dict] = {}
        prepared_by_service: Dict[str, dict] = {}
        prepared_by_term: Dict[str, dict] = {}

        def _key_suffix(unmapped_key: str) -> str:
            key = str(unmapped_key or "").strip()
            if not key:
                return ""
            if key.startswith("id:"):
                return key
            if ":" in key:
                return key.split(":", 1)[1]
            return key


        def _key_service_token(unmapped_key: str) -> str:
            suffix = _key_suffix(unmapped_key)
            if not suffix or suffix.startswith("id:"):
                return ""
            return suffix.split("|", 1)[0]

        def _row_service_term(row_data: dict) -> str:
            return normalizar_texto(
                str(
                    row_data.get("servico_normalizado", "")
                    or row_data.get("servico_bruto", "")
                    or row_data.get("item_normalizado", "")
                    or row_data.get("item_original", "")
                    or ""
                )
            )


        for raw in corrections:
            key = str(raw.get("unmapped_key", "") or "").strip()
            servico = str(raw.get("servico", "") or "").strip()
            if not key or not servico:
                continue
            meta = self._manual_service_meta(servico)
            row = {
                "unmapped_key": key,
                "servico": meta["servico"],
                "categoria": meta["categoria"],
                "label": meta["label"],
            }
            prepared[key] = row
            suffix = _key_suffix(key)
            if suffix and suffix not in prepared_by_suffix:
                prepared_by_suffix[suffix] = row
            service_token = _key_service_token(key)
            if service_token and service_token not in prepared_by_service:
                prepared_by_service[service_token] = row
            term_hint = normalizar_texto(str(raw.get("servico_term", "") or ""))
            if term_hint and term_hint not in prepared_by_term:
                prepared_by_term[term_hint] = row

        if not prepared:
            return parsed

        timestamp = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        applied = []

        def _find_correction(row_key: str, row_data: dict) -> dict | None:
            exact = prepared.get(row_key)
            if exact:
                return exact
            suffix = _key_suffix(row_key)
            by_suffix = prepared_by_suffix.get(suffix)
            if by_suffix:
                return by_suffix
            service_token = _key_service_token(row_key)
            if service_token:
                by_service = prepared_by_service.get(service_token)
                if by_service:
                    return by_service
            term_key = _row_service_term(row_data)
            if term_key:
                by_term = prepared_by_term.get(term_key)
                if by_term:
                    return by_term
            return None


        for idx, item in enumerate(parsed.get("execucao", [])):
            key = self._unmapped_row_key(
                {
                    "id_item": item.get("id_item", ""),
                    "servico_normalizado": item.get("servico_normalizado", item.get("item_normalizado", "")),
                    "servico_bruto": item.get("servico_bruto", item.get("item_original", "")),
                    "quantidade": item.get("quantidade", ""),
                    "unidade": item.get("unidade", ""),
                    "nucleo": item.get("nucleo", ""),
                    "logradouro": item.get("logradouro", ""),
                    "equipe": item.get("equipe", ""),
                    "municipio": item.get("municipio", ""),
                },
                fallback_prefix=f"ex{idx}",
            )
            correction = _find_correction(key, item)
            if not correction:
                continue

            item.setdefault("servico_original_bruto", str(item.get("servico_bruto", item.get("item_original", "")) or "").strip())
            item.setdefault("servico_original_normalizado", str(item.get("servico_normalizado", item.get("item_normalizado", "")) or "").strip())
            item["servico_oficial"] = correction["servico"]
            item["item_normalizado"] = correction["servico"]
            item["categoria_item"] = correction["categoria"]
            item["categoria"] = correction["categoria"]
            item["regra_disparada"] = "correcao_manual"
            item["correcao_manual_aplicada"] = "sim"
            item["servico_corrigido_manual"] = correction["servico"]
            item["categoria_corrigida_manual"] = correction["categoria"]
            item["data_correcao_manual"] = timestamp
            applied.append(
                {
                    "unmapped_key": key,
                    "servico": correction["servico"],
                    "categoria": correction["categoria"],
                    "id_item": str(item.get("id_item", "") or "").strip(),
                }
            )

        for idx, row in enumerate(parsed.get("servicos_nao_mapeados", [])):
            key = self._unmapped_row_key(row, fallback_prefix=f"um{idx}")
            correction = _find_correction(key, row)
            if not correction:
                continue
            row["corrigido_manual"] = "sim"
            row["servico_corrigido_manual"] = correction["servico"]
            row["categoria_corrigida_manual"] = correction["categoria"]
            row["data_correcao_manual"] = timestamp

        if applied:
            parsed.setdefault("correcoes_manuais", [])
            parsed["correcoes_manuais"].extend(applied)

        return parsed

    def _bool_true(self, value: object) -> bool:
        return str(value or "").strip().lower() in {"1", "true", "sim", "yes", "on", "s"}

    def _unmapped_row_key(self, row: dict, fallback_prefix: str = "row") -> str:
        id_item = str(row.get("id_item", "") or "").strip()
        if id_item:
            return f"id:{id_item}"

        serv_norm = normalizar_texto(
            str(row.get("servico_normalizado", "") or "")
            or str(row.get("servico_bruto", "") or "")
            or str(row.get("item_normalizado", "") or "")
            or str(row.get("item_original", "") or "")
        )
        quantidade = str(row.get("quantidade", "") or "").strip()
        unidade = normalizar_texto(str(row.get("unidade", "") or ""))
        nucleo = normalizar_texto(str(row.get("nucleo", "") or ""))
        logradouro = normalizar_texto(str(row.get("logradouro", "") or ""))
        equipe = normalizar_texto(self._first_team(row.get("equipe", "")))
        municipio = normalizar_texto(str(row.get("municipio", "") or ""))
        return f"{fallback_prefix}:{serv_norm}|{quantidade}|{unidade}|{nucleo}|{logradouro}|{equipe}|{municipio}"

    def _contexto_linha(self, row: dict) -> str:
        return self._format_context_label(
            {
                "nucleo": row.get("nucleo", ""),
                "municipio": row.get("municipio", ""),
                "equipe": row.get("equipe", ""),
                "logradouro": row.get("logradouro", ""),
            }
        )

    def _recent_unmapped_frequency(self, window_days: int = 45, recent_runs: int = 300) -> Dict[str, int]:
        days = self._safe_int(window_days, default=45, min_value=1, max_value=3650)
        runs_limit = self._safe_int(recent_runs, default=300, min_value=10, max_value=5000)

        all_history = self.read_history(limit=max(runs_limit * 3, 500))
        cutoff = datetime.now() - timedelta(days=days)
        selected = []
        for row in all_history:
            dt = self._parse_history_datetime(row.get("processed_at", ""))
            if dt and dt < cutoff:
                continue
            selected.append(row)
            if len(selected) >= runs_limit:
                break

        counter: Counter = Counter()
        for run in selected:
            output_dir = str(run.get("output_dir", "") or "").strip()
            if not output_dir:
                continue
            rows = self._read_csv_rows(Path(output_dir) / "servico_nao_mapeado.csv")
            for row in rows:
                if self._bool_true(row.get("corrigido_manual", "")):
                    continue
                termo = str(row.get("servico_normalizado", "") or "").strip()
                if not termo:
                    termo = str(row.get("servico_bruto", "") or "").strip()
                if not termo:
                    continue
                counter[normalizar_texto(termo)] += 1

        return dict(counter)

    def collect_unmapped(self, parsed: dict, include_corrigidos: bool = False) -> List[dict]:
        rows: List[dict] = []
        for idx, raw in enumerate(parsed.get("servicos_nao_mapeados", [])):
            row = dict(raw)
            row["equipe"] = self._first_team(row.get("equipe", ""))
            if not include_corrigidos and self._bool_true(row.get("corrigido_manual", "")):
                continue
            if not include_corrigidos and self._bool_true(row.get("corrigido_alias", "")):
                continue
            if not row.get("regra_disparada"):
                row["regra_disparada"] = "nao_mapeado"
            row.setdefault("id_item", "")
            row.setdefault("sugestao_categoria", "")
            row["unmapped_key"] = self._unmapped_row_key(row, fallback_prefix=f"um{idx}")
            row["contexto_linha"] = self._contexto_linha(row)
            rows.append(row)

        extra = []
        for idx, item in enumerate(parsed.get("execucao", [])):
            servico_oficial = str(item.get("servico_oficial", "") or "").strip().lower()
            categoria_item = str(item.get("categoria_item", "") or "").strip().lower()
            categoria = str(item.get("categoria", "") or "").strip().lower()
            is_unmapped = (
                servico_oficial == "servico_nao_mapeado"
                or categoria_item == "servico_nao_mapeado"
                or categoria == "servico_nao_mapeado"
            )
            if not is_unmapped:
                continue
            extra_row = {
                "id_item": item.get("id_item", ""),
                "data": item.get("data_referencia", ""),
                "mensagem_original": item.get("mensagem_origem", ""),
                "servico_bruto": item.get("servico_bruto", item.get("item_original", "")),
                "servico_normalizado": item.get("servico_normalizado", item.get("item_normalizado", "")),
                "quantidade": item.get("quantidade", ""),
                "unidade": item.get("unidade", ""),
                "nucleo": item.get("nucleo", ""),
                "logradouro": item.get("logradouro", ""),
                "municipio": item.get("municipio", ""),
                "equipe": self._first_team(item.get("equipe", "")),
                "sugestao_categoria": item.get("categoria_item", ""),
                "regra_disparada": item.get("regra_disparada", "nao_mapeado"),
            }
            extra_row["unmapped_key"] = self._unmapped_row_key(extra_row, fallback_prefix=f"ex{idx}")
            extra_row["contexto_linha"] = self._contexto_linha(extra_row)
            extra.append(extra_row)

        if not rows:
            return extra

        seen = {str(r.get("unmapped_key", "") or "").strip() for r in rows}
        for r in extra:
            key = str(r.get("unmapped_key", "") or "").strip()
            if key and key not in seen:
                rows.append(r)
                seen.add(key)
        return rows

    def collect_incomplete_execution(self, parsed: dict) -> List[dict]:
        incomplete = []
        for item in parsed.get("execucao", []):
            quantidade = item.get("quantidade", "")
            unidade = str(item.get("unidade", "") or "").strip()
            servico_bruto = str(item.get("servico_bruto", item.get("item_original", "")) or "").strip()
            is_incomplete = (quantidade in ("", None)) or (not unidade) or (not servico_bruto)
            if is_incomplete:
                incomplete.append(
                    {
                        "id_item": item.get("id_item", ""),
                        "mensagem_origem": item.get("mensagem_origem", ""),
                        "quantidade": quantidade,
                        "unidade": unidade,
                        "servico_bruto": servico_bruto,
                        "servico_oficial": item.get("servico_oficial", item.get("item_normalizado", "")),
                    }
                )
        return incomplete

    def collect_unmapped_term_stats(self, unmapped_rows: List[dict], max_items: int = 12) -> List[dict]:
        counter: Counter = Counter()
        label_by_key: Dict[str, str] = {}
        for row in unmapped_rows:
            bruto = str(row.get("servico_bruto", "") or "").strip()
            if not bruto:
                continue
            key = bruto.lower()
            counter[key] += 1
            if key not in label_by_key:
                label_by_key[key] = bruto

        out: List[dict] = []
        for key, qtd in counter.most_common(max_items):
            out.append({"termo": label_by_key[key], "ocorrencias": qtd})
        return out

    def _safe_int(self, value: object, default: int, min_value: int, max_value: int) -> int:
        try:
            parsed = int(str(value or "").strip())
        except Exception:
            parsed = default
        return max(min_value, min(parsed, max_value))

    def _parse_history_datetime(self, value: object) -> datetime | None:
        text = str(value or "").strip()
        if not text:
            return None
        for fmt in ("%d/%m/%Y %H:%M:%S", "%Y-%m-%d %H:%M:%S"):
            try:
                return datetime.strptime(text, fmt)
            except ValueError:
                continue
        return None

    def _format_context_label(self, row: dict) -> str:
        nucleo = str(row.get("nucleo", "") or "").strip()
        municipio = str(row.get("municipio", "") or "").strip()
        equipe = self._first_team(row.get("equipe", ""))
        logradouro = str(row.get("logradouro", "") or "").strip()

        parts = []
        if nucleo:
            parts.append(f"Nucleo: {nucleo}")
        if municipio:
            parts.append(f"Municipio: {municipio}")
        if equipe:
            parts.append(f"Equipe: {equipe}")
        if logradouro:
            parts.append(f"Logradouro: {logradouro}")
        return " | ".join(parts)

    def _write_unmapped_alias_candidates(self, alias_candidates: List[dict]) -> str:
        candidate_headers = [
            "termo_candidato",
            "exemplo_servico_bruto",
            "ocorrencias",
            "ultima_ocorrencia",
            "sugestao_categoria",
            "servico_corrigido_recorrente",
            "corrigidos_manuais",
            "regra_recorrente",
            "contexto_recorrente",
            "ultimo_nucleo",
            "ultimo_municipio",
            "ultimo_equipe",
            "recomendacao",
        ]
        self.unmapped_candidates_file.parent.mkdir(parents=True, exist_ok=True)
        with self.unmapped_candidates_file.open("w", encoding="utf-8-sig", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=candidate_headers, extrasaction="ignore")
            writer.writeheader()
            for row in alias_candidates:
                clean_row = {h: row.get(h, "") for h in candidate_headers}
                writer.writerow(clean_row)

        try:
            return self.unmapped_candidates_file.resolve().as_uri()
        except Exception:
            return ""

    def _build_unmapped_dashboard_from_db(
        self,
        *,
        days: int,
        runs_limit: int,
        top_limit: int,
        recent_limit: int,
    ) -> dict | None:
        if self.db_manager is None:
            return None

        where_sql = """
            (data_referencia IS NULL OR data_referencia >= CURRENT_DATE - (%s * INTERVAL '1 day'))
            AND COALESCE(NULLIF(TRIM(servico_normalizado), ''), NULLIF(TRIM(servico_bruto), ''), NULLIF(TRIM(item_normalizado), ''), NULLIF(TRIM(item_original), '')) IS NOT NULL
            AND (
                NULLIF(TRIM(servico_oficial), '') IS NULL
                OR LOWER(TRIM(servico_oficial)) IN ('servico_nao_mapeado', 'nao_mapeado', '-')
                OR LOWER(TRIM(COALESCE(categoria, ''))) = 'servico_nao_mapeado'
                OR LOWER(TRIM(COALESCE(categoria_item, ''))) = 'servico_nao_mapeado'
            )
        """

        top_rows: List[dict] = []
        recent_rows: List[dict] = []
        total_itens = 0
        termos_distintos = 0
        runs_com_nao_mapeado = 0
        ultima_ocorrencia_global = "-"

        try:
            with self.db_manager.connection() as conn:
                with conn.cursor() as cur:
                    cur.execute(
                        f"""
                        WITH base AS (
                            SELECT
                                id,
                                data_referencia,
                                COALESCE(NULLIF(TRIM(servico_normalizado), ''), NULLIF(TRIM(servico_bruto), ''), NULLIF(TRIM(item_normalizado), ''), NULLIF(TRIM(item_original), '')) AS termo,
                                COALESCE(NULLIF(TRIM(servico_bruto), ''), NULLIF(TRIM(item_original), ''), '') AS servico_bruto_ex,
                                COALESCE(NULLIF(TRIM(nucleo_oficial), ''), NULLIF(TRIM(nucleo), ''), '') AS nucleo,
                                COALESCE(NULLIF(TRIM(municipio_oficial), ''), NULLIF(TRIM(municipio), ''), '') AS municipio,
                                COALESCE(NULLIF(TRIM(equipe), ''), '') AS equipe,
                                COALESCE(NULLIF(TRIM(logradouro), ''), '') AS logradouro,
                                COALESCE(NULLIF(TRIM(categoria), ''), NULLIF(TRIM(categoria_item), ''), '') AS sugestao_categoria
                            FROM management_execucao
                            WHERE {where_sql}
                        ),
                        agg AS (
                            SELECT
                                termo,
                                COUNT(*) AS ocorrencias,
                                MAX(data_referencia) AS ultima_data
                            FROM base
                            GROUP BY termo
                        ),
                        sample AS (
                            SELECT DISTINCT ON (termo)
                                termo,
                                nucleo,
                                municipio,
                                equipe,
                                logradouro,
                                servico_bruto_ex,
                                sugestao_categoria,
                                data_referencia
                            FROM base
                            ORDER BY termo, data_referencia DESC NULLS LAST, id DESC
                        )
                        SELECT
                            agg.termo,
                            agg.ocorrencias,
                            agg.ultima_data,
                            sample.nucleo,
                            sample.municipio,
                            sample.equipe,
                            sample.logradouro,
                            sample.servico_bruto_ex,
                            sample.sugestao_categoria
                        FROM agg
                        LEFT JOIN sample ON sample.termo = agg.termo
                        ORDER BY agg.ocorrencias DESC, agg.ultima_data DESC NULLS LAST, agg.termo ASC
                        LIMIT %s
                        """,
                        (days, top_limit),
                    )
                    top_fetched = cur.fetchall() or []

                    cur.execute(
                        f"""
                        SELECT
                            COUNT(*) AS total_itens,
                            COUNT(DISTINCT termo) AS termos_distintos,
                            COUNT(DISTINCT data_referencia) AS runs_com_nao_mapeado,
                            MAX(data_referencia) AS ultima_data
                        FROM (
                            SELECT
                                data_referencia,
                                COALESCE(NULLIF(TRIM(servico_normalizado), ''), NULLIF(TRIM(servico_bruto), ''), NULLIF(TRIM(item_normalizado), ''), NULLIF(TRIM(item_original), '')) AS termo
                            FROM management_execucao
                            WHERE {where_sql}
                        ) t
                        """,
                        (days,),
                    )
                    summary_row = cur.fetchone()

                    cur.execute(
                        f"""
                        SELECT
                            data_referencia,
                            COALESCE(NULLIF(TRIM(nucleo_oficial), ''), NULLIF(TRIM(nucleo), ''), '') AS nucleo,
                            COALESCE(NULLIF(TRIM(municipio_oficial), ''), NULLIF(TRIM(municipio), ''), '') AS municipio,
                            COALESCE(NULLIF(TRIM(equipe), ''), '') AS equipe,
                            COALESCE(NULLIF(TRIM(logradouro), ''), '') AS logradouro,
                            COALESCE(NULLIF(TRIM(servico_normalizado), ''), NULLIF(TRIM(servico_bruto), ''), NULLIF(TRIM(item_normalizado), ''), NULLIF(TRIM(item_original), '')) AS termo,
                            COALESCE(NULLIF(TRIM(servico_bruto), ''), '') AS servico_bruto,
                            COALESCE(NULLIF(TRIM(categoria), ''), NULLIF(TRIM(categoria_item), ''), '') AS sugestao_categoria,
                            quantidade,
                            unidade
                        FROM management_execucao
                        WHERE {where_sql}
                        ORDER BY data_referencia DESC NULLS LAST, id DESC
                        LIMIT %s
                        """,
                        (days, recent_limit),
                    )
                    recent_fetched = cur.fetchall() or []
        except Exception:
            return None

        for item in top_fetched:
            termo, ocorrencias, ultima_data, nucleo, municipio, equipe, logradouro, servico_bruto_ex, sugestao_categoria = item
            ultima_text = "-"
            if isinstance(ultima_data, (date, datetime)):
                ultima_text = (
                    ultima_data.strftime("%d/%m/%Y")
                    if isinstance(ultima_data, date) and not isinstance(ultima_data, datetime)
                    else ultima_data.strftime("%d/%m/%Y %H:%M")
                )
            context_parts = []
            if str(nucleo or "").strip():
                context_parts.append(f"Nucleo: {str(nucleo).strip()}")
            if str(municipio or "").strip():
                context_parts.append(f"Municipio: {str(municipio).strip()}")
            if str(equipe or "").strip():
                context_parts.append(f"Equipe: {self._first_team(equipe)}")
            if str(logradouro or "").strip():
                context_parts.append(f"Logradouro: {str(logradouro).strip()}")

            top_rows.append(
                {
                    "termo": str(termo or "").strip(),
                    "ocorrencias": int(ocorrencias or 0),
                    "ultima_ocorrencia": ultima_text,
                    "ultimo_nucleo": str(nucleo or "").strip(),
                    "ultimo_municipio": str(municipio or "").strip(),
                    "ultimo_equipe": self._first_team(equipe),
                    "ultimo_logradouro": str(logradouro or "").strip(),
                    "contexto_recorrente": " | ".join(context_parts),
                    "exemplo_servico_bruto": str(servico_bruto_ex or "").strip(),
                    "sugestao_categoria_recorrente": str(sugestao_categoria or "").strip(),
                    "regra_recorrente": "db_master_execucao",
                    "corrigidos_manuais": 0,
                    "nao_corrigidos": int(ocorrencias or 0),
                    "servico_corrigido_recorrente": "",
                }
            )

        if summary_row:
            total_itens = int(summary_row[0] or 0)
            termos_distintos = int(summary_row[1] or 0)
            runs_com_nao_mapeado = int(summary_row[2] or 0)
            last_dt = summary_row[3]
            if isinstance(last_dt, (date, datetime)):
                ultima_ocorrencia_global = (
                    last_dt.strftime("%d/%m/%Y")
                    if isinstance(last_dt, date) and not isinstance(last_dt, datetime)
                    else last_dt.strftime("%d/%m/%Y %H:%M")
                )

        for item in recent_fetched:
            data_ref, nucleo, municipio, equipe, logradouro, termo, servico_bruto, sugestao_categoria, quantidade, unidade = item
            processed_at = "-"
            if isinstance(data_ref, (date, datetime)):
                processed_at = (
                    data_ref.strftime("%d/%m/%Y")
                    if isinstance(data_ref, date) and not isinstance(data_ref, datetime)
                    else data_ref.strftime("%d/%m/%Y %H:%M")
                )
            ctx = self._format_context_label(
                {
                    "nucleo": nucleo,
                    "municipio": municipio,
                    "equipe": equipe,
                    "logradouro": logradouro,
                }
            )
            recent_rows.append(
                {
                    "termo": str(termo or "").strip(),
                    "servico_bruto": str(servico_bruto or "").strip(),
                    "servico_normalizado": str(termo or "").strip(),
                    "mensagem_original": "",
                    "quantidade": str(quantidade or "").strip(),
                    "unidade": str(unidade or "").strip(),
                    "nucleo": str(nucleo or "").strip(),
                    "municipio": str(municipio or "").strip(),
                    "equipe": self._first_team(equipe),
                    "logradouro": str(logradouro or "").strip(),
                    "sugestao_categoria": str(sugestao_categoria or "").strip(),
                    "regra_disparada": "db_master_execucao",
                    "corrigido_manual": "nao",
                    "servico_corrigido_manual": "",
                    "data_obra": processed_at,
                    "processed_at": processed_at,
                    "contexto": ctx,
                }
            )

        alias_candidates: List[dict] = []
        for row in top_rows:
            recommendation = "revisar_manual"
            if row.get("sugestao_categoria_recorrente"):
                recommendation = f"avaliar_categoria:{row['sugestao_categoria_recorrente']}"
            alias_candidates.append(
                {
                    "termo_candidato": row.get("termo", ""),
                    "exemplo_servico_bruto": row.get("exemplo_servico_bruto", ""),
                    "ocorrencias": row.get("ocorrencias", 0),
                    "ultima_ocorrencia": row.get("ultima_ocorrencia", "-"),
                    "sugestao_categoria": row.get("sugestao_categoria_recorrente", ""),
                    "servico_corrigido_recorrente": "",
                    "corrigidos_manuais": 0,
                    "regra_recorrente": "db_master_execucao",
                    "contexto_recorrente": row.get("contexto_recorrente", ""),
                    "ultimo_nucleo": row.get("ultimo_nucleo", ""),
                    "ultimo_municipio": row.get("ultimo_municipio", ""),
                    "ultimo_equipe": row.get("ultimo_equipe", ""),
                    "recomendacao": recommendation,
                }
            )

        candidates_uri = self._write_unmapped_alias_candidates(alias_candidates)

        return {
            "window_days": days,
            "recent_runs": runs_limit,
            "runs_consideradas": runs_com_nao_mapeado,
            "runs_com_nao_mapeado": runs_com_nao_mapeado,
            "total_itens": total_itens,
            "termos_distintos": termos_distintos,
            "ultima_ocorrencia_global": ultima_ocorrencia_global,
            "top_terms": top_rows,
            "recent_items": recent_rows,
            "alias_candidates": alias_candidates[:top_limit],
            "alias_candidates_count": len(alias_candidates),
            "alias_candidates_path": str(self.unmapped_candidates_file.resolve()),
            "alias_candidates_uri": candidates_uri,
        }

    def build_unmapped_dashboard(
        self,
        window_days: int = 30,
        recent_runs: int = 120,
        top_terms: int = 20,
        recent_items_limit: int = 20,
    ) -> dict:
        days = self._safe_int(window_days, default=30, min_value=1, max_value=3650)
        runs_limit = self._safe_int(recent_runs, default=120, min_value=1, max_value=5000)
        top_limit = self._safe_int(top_terms, default=20, min_value=5, max_value=100)
        recent_limit = self._safe_int(recent_items_limit, default=20, min_value=5, max_value=200)

        db_dashboard = self._build_unmapped_dashboard_from_db(
            days=days,
            runs_limit=runs_limit,
            top_limit=top_limit,
            recent_limit=recent_limit,
        )
        if db_dashboard and (
            int(db_dashboard.get("total_itens", 0) or 0) > 0
            or int(db_dashboard.get("termos_distintos", 0) or 0) > 0
        ):
            return db_dashboard

        all_history = self.read_history(limit=max(runs_limit * 3, 500))
        cutoff = datetime.now() - timedelta(days=days)
        selected_runs: List[dict] = []
        for row in all_history:
            dt = self._parse_history_datetime(row.get("processed_at", ""))
            if dt and dt < cutoff:
                continue
            selected_runs.append({**row, "_processed_dt": dt})
            if len(selected_runs) >= runs_limit:
                break

        aggregated: Dict[str, dict] = {}
        recent_items: List[dict] = []
        runs_with_unmapped = 0
        last_seen_dt: datetime | None = None

        for run in selected_runs:
            output_dir = str(run.get("output_dir", "") or "").strip()
            if not output_dir:
                continue
            unmapped_csv = Path(output_dir) / "servico_nao_mapeado.csv"
            rows = self._read_csv_rows(unmapped_csv)
            if not rows:
                continue

            runs_with_unmapped += 1
            run_dt = run.get("_processed_dt")
            if run_dt and (last_seen_dt is None or run_dt > last_seen_dt):
                last_seen_dt = run_dt

            for row in rows:
                termo_base = str(row.get("servico_normalizado", "") or "").strip()
                if not termo_base:
                    termo_base = str(row.get("servico_bruto", "") or "").strip()
                if not termo_base:
                    continue

                termo_key = termo_base.lower()
                item_context = self._format_context_label(row)
                sugestao_categoria = str(row.get("sugestao_categoria", "") or "").strip()
                regra_disparada = str(row.get("regra_disparada", "") or "").strip()
                corrigido_manual = self._bool_true(row.get("corrigido_manual", ""))
                servico_corrigido = str(row.get("servico_corrigido_manual", "") or "").strip()
                item = {
                    "termo": termo_base,
                    "servico_bruto": str(row.get("servico_bruto", "") or "").strip(),
                    "servico_normalizado": str(row.get("servico_normalizado", "") or "").strip(),
                    "mensagem_original": str(row.get("mensagem_original", "") or "").strip(),
                    "quantidade": str(row.get("quantidade", "") or "").strip(),
                    "unidade": str(row.get("unidade", "") or "").strip(),
                    "nucleo": str(row.get("nucleo", "") or "").strip(),
                    "municipio": str(row.get("municipio", "") or "").strip(),
                    "equipe": self._first_team(row.get("equipe", "")),
                    "logradouro": str(row.get("logradouro", "") or "").strip(),
                    "sugestao_categoria": sugestao_categoria,
                    "regra_disparada": regra_disparada,
                    "corrigido_manual": "sim" if corrigido_manual else "nao",
                    "servico_corrigido_manual": servico_corrigido,
                    "data_obra": str(row.get("data", "") or "").strip(),
                    "processed_at": str(run.get("processed_at", "") or "").strip(),
                    "contexto": item_context,
                    "_processed_dt": run_dt,
                }
                recent_items.append(item)

                agg = aggregated.get(termo_key)
                if not agg:
                    agg = {
                        "termo": termo_base,
                        "ocorrencias": 0,
                        "ultima_ocorrencia": "",
                        "_ultima_dt": None,
                        "ultimo_nucleo": "",
                        "ultimo_municipio": "",
                        "ultimo_equipe": "",
                        "ultimo_logradouro": "",
                        "exemplo_servico_bruto": "",
                        "corrigidos_manuais": 0,
                        "nao_corrigidos": 0,
                        "context_counter": Counter(),
                        "categoria_counter": Counter(),
                        "regra_counter": Counter(),
                        "servico_corrigido_counter": Counter(),
                    }
                    aggregated[termo_key] = agg

                agg["ocorrencias"] += 1
                if item_context:
                    agg["context_counter"][item_context] += 1
                if sugestao_categoria:
                    agg["categoria_counter"][sugestao_categoria] += 1
                if regra_disparada:
                    agg["regra_counter"][regra_disparada] += 1
                if servico_corrigido:
                    agg["servico_corrigido_counter"][servico_corrigido] += 1
                if corrigido_manual:
                    agg["corrigidos_manuais"] += 1
                else:
                    agg["nao_corrigidos"] += 1
                if not agg["exemplo_servico_bruto"] and item.get("servico_bruto"):
                    agg["exemplo_servico_bruto"] = item["servico_bruto"]

                if run_dt and (agg["_ultima_dt"] is None or run_dt > agg["_ultima_dt"]):
                    agg["_ultima_dt"] = run_dt
                    agg["ultima_ocorrencia"] = run_dt.strftime("%d/%m/%Y %H:%M")
                    agg["ultimo_nucleo"] = item["nucleo"]
                    agg["ultimo_municipio"] = item["municipio"]
                    agg["ultimo_equipe"] = item["equipe"]
                    agg["ultimo_logradouro"] = item["logradouro"]

        top_rows: List[dict] = []
        for row in aggregated.values():
            context_ranked = row["context_counter"].most_common(3)
            context_text = "; ".join(f"{ctx} ({qtd})" for ctx, qtd in context_ranked if ctx)
            categoria_recorrente = ""
            if row["categoria_counter"]:
                categoria_recorrente = row["categoria_counter"].most_common(1)[0][0]
            regra_recorrente = ""
            if row["regra_counter"]:
                regra_recorrente = row["regra_counter"].most_common(1)[0][0]
            servico_corrigido_recorrente = ""
            if row["servico_corrigido_counter"]:
                servico_corrigido_recorrente = row["servico_corrigido_counter"].most_common(1)[0][0]
            top_rows.append(
                {
                    "termo": row["termo"],
                    "ocorrencias": row["ocorrencias"],
                    "ultima_ocorrencia": row["ultima_ocorrencia"] or "-",
                    "ultimo_nucleo": row["ultimo_nucleo"],
                    "ultimo_municipio": row["ultimo_municipio"],
                    "ultimo_equipe": row["ultimo_equipe"],
                    "ultimo_logradouro": row["ultimo_logradouro"],
                    "contexto_recorrente": context_text,
                    "exemplo_servico_bruto": row["exemplo_servico_bruto"],
                    "sugestao_categoria_recorrente": categoria_recorrente,
                    "regra_recorrente": regra_recorrente,
                    "corrigidos_manuais": int(row["corrigidos_manuais"]),
                    "nao_corrigidos": int(row["nao_corrigidos"]),
                    "servico_corrigido_recorrente": servico_corrigido_recorrente,
                    "_ultima_dt": row["_ultima_dt"],
                }
            )

        top_rows.sort(
            key=lambda r: (
                int(r["ocorrencias"]),
                r.get("_ultima_dt") if r.get("_ultima_dt") else datetime.min,
            ),
            reverse=True,
        )
        top_rows = top_rows[:top_limit]
        for row in top_rows:
            row.pop("_ultima_dt", None)

        recent_items.sort(
            key=lambda x: x.get("_processed_dt") if x.get("_processed_dt") else datetime.min,
            reverse=True,
        )
        recent_rows: List[dict] = []
        for item in recent_items[:recent_limit]:
            clean = dict(item)
            clean.pop("_processed_dt", None)
            recent_rows.append(clean)

        alias_candidates: List[dict] = []
        for row in top_rows:
            recommendation = "revisar_manual"
            if row.get("servico_corrigido_recorrente"):
                recommendation = f"avaliar_alias_para:{row['servico_corrigido_recorrente']}"
            elif row.get("sugestao_categoria_recorrente"):
                recommendation = f"avaliar_categoria:{row['sugestao_categoria_recorrente']}"

            alias_candidates.append(
                {
                    "termo_candidato": row.get("termo", ""),
                    "exemplo_servico_bruto": row.get("exemplo_servico_bruto", ""),
                    "ocorrencias": row.get("ocorrencias", 0),
                    "ultima_ocorrencia": row.get("ultima_ocorrencia", "-"),
                    "sugestao_categoria": row.get("sugestao_categoria_recorrente", ""),
                    "servico_corrigido_recorrente": row.get("servico_corrigido_recorrente", ""),
                    "corrigidos_manuais": row.get("corrigidos_manuais", 0),
                    "regra_recorrente": row.get("regra_recorrente", ""),
                    "contexto_recorrente": row.get("contexto_recorrente", ""),
                    "ultimo_nucleo": row.get("ultimo_nucleo", ""),
                    "ultimo_municipio": row.get("ultimo_municipio", ""),
                    "ultimo_equipe": row.get("ultimo_equipe", ""),
                    "recomendacao": recommendation,
                }
            )

        candidates_uri = self._write_unmapped_alias_candidates(alias_candidates)

        return {
            "window_days": days,
            "recent_runs": runs_limit,
            "runs_consideradas": len(selected_runs),
            "runs_com_nao_mapeado": runs_with_unmapped,
            "total_itens": len(recent_items),
            "termos_distintos": len(aggregated),
            "ultima_ocorrencia_global": last_seen_dt.strftime("%d/%m/%Y %H:%M") if last_seen_dt else "-",
            "top_terms": top_rows,
            "recent_items": recent_rows,
            "alias_candidates": alias_candidates[:top_limit],
            "alias_candidates_count": len(alias_candidates),
            "alias_candidates_path": str(self.unmapped_candidates_file.resolve()),
            "alias_candidates_uri": candidates_uri,
        }

    def enrich_unmapped_for_review(self, unmapped_rows: List[dict]) -> List[dict]:
        freq_map = self._recent_unmapped_frequency(window_days=45, recent_runs=300)
        enriched: List[dict] = []
        for row in unmapped_rows:
            item = dict(row)
            norm = str(item.get("servico_normalizado", "") or "").strip()
            if not norm:
                norm = str(item.get("servico_bruto", "") or "").strip()
            freq = freq_map.get(normalizar_texto(norm), 0)
            item["frequencia_recente"] = int(freq)
            item.setdefault("contexto_linha", self._contexto_linha(item))
            item.setdefault("sugestao_categoria", "")
            item.setdefault("unmapped_key", self._unmapped_row_key(item, fallback_prefix="rv"))
            enriched.append(item)
        return enriched

    def _municipio_profile_status(self, main_fields: dict) -> dict:
        nucleo = str(main_fields.get("nucleo", "") or "").strip()
        municipio = str(main_fields.get("municipio", "") or "").strip()

        status = {
            "has_profile": False,
            "expected_municipio": "",
            "is_divergent": False,
            "is_missing": not bool(municipio),
        }

        if not nucleo:
            return status

        profile = self._get_nucleo_profile(nucleo)
        if not profile:
            return status

        expected = str(profile.get("municipio", "") or "").strip()
        status["has_profile"] = True
        status["expected_municipio"] = expected

        if expected and municipio and normalizar_texto(expected) != normalizar_texto(municipio):
            status["is_divergent"] = True

        return status

    def split_alerts_by_priority(self, alert_items: List[dict]) -> dict:
        groups = {"critical": [], "warning": [], "info": []}
        for item in alert_items:
            level = str(item.get("level", "info") or "info").strip().lower()
            if level not in groups:
                level = "info"
            groups[level].append(item)
        return groups

    def summarize_preview(self, parsed: dict) -> dict:
        total_exec = len(parsed.get("execucao", []))
        total_frentes = len(parsed.get("frentes", []))
        total_ocorrencias = len(parsed.get("ocorrencias", []))
        total_obs = len(parsed.get("observacoes", []))
        total_nao_mapeado = len(self.collect_unmapped(parsed))
        total_incompleto = len(self.collect_incomplete_execution(parsed))
        total_mapeado = max(total_exec - total_nao_mapeado, 0)
        percentual_mapeado = round((total_mapeado / total_exec) * 100, 1) if total_exec else 100.0
        return {
            "frentes": total_frentes,
            "execucao": total_exec,
            "ocorrencias": total_ocorrencias,
            "observacoes": total_obs,
            "nao_mapeados": total_nao_mapeado,
            "incompletos": total_incompleto,
            "percentual_mapeado": percentual_mapeado,
        }

    def build_alert_items(
        self,
        parsed: dict,
        main_fields: dict | None = None,
        main_lists: dict | None = None,
        context_overview: dict | None = None,
        unmapped_rows: List[dict] | None = None,
        incomplete_rows: List[dict] | None = None,
        manual_municipio_status: dict | None = None,
    ) -> List[dict]:
        alert_items: List[dict] = []
        main = dict(main_fields or self.extract_main_fields(parsed))
        context = context_overview or self.collect_context_overview(parsed)
        lists = main_lists or self.build_main_lists(main, context)

        missing_main = self._missing_main_fields(main)
        missing_text = ", ".join(missing_main)

        if "nucleo" in missing_main:
            alert_items.append(
                {
                    "level": "critical",
                    "title": "Nucleo ausente",
                    "message": "Informe o nucleo explicitamente. O sistema nao infere nucleo por viela/logradouro.",
                    "code": "missing_nucleo",
                }
            )

        municipio_profile = self._municipio_profile_status(main)
        nucleo_master_summary = dict(parsed.get("nucleo_master_summary", {}) or {})
        if "municipio" in missing_main:
            expected = str(municipio_profile.get("expected_municipio", "") or "").strip()
            if expected:
                msg = (
                    f"Municipio ausente para o nucleo informado. Cadastro de referencia indica '{expected}'. "
                    "Preencha antes da geracao final."
                )
            else:
                msg = "Municipio ausente. Preencha para evitar consolidacao com contexto incompleto."
            alert_items.append(
                {
                    "level": "critical",
                    "title": "Municipio ausente",
                    "message": msg,
                    "code": "missing_municipio",
                }
            )

        if municipio_profile.get("is_divergent"):
            expected = str(municipio_profile.get("expected_municipio", "") or "").strip()
            atual = str(main.get("municipio", "") or "").strip()
            alert_items.append(
                {
                    "level": "warning",
                    "title": "Municipio divergente do cadastro do nucleo",
                    "message": f"Nucleo com cadastro conhecido: municipio esperado '{expected}', valor atual '{atual}'.",
                    "code": "municipio_divergente",
                }
            )

        manual_municipio_status = dict(manual_municipio_status or {})
        if manual_municipio_status.get("is_divergent"):
            expected = str(manual_municipio_status.get("expected_municipio", "") or "").strip()
            atual = str(manual_municipio_status.get("provided_municipio", "") or "").strip()
            if expected and atual:
                alert_items.append(
                    {
                        "level": "warning",
                        "title": "Municipio divergente do cadastro do nucleo",
                        "message": (
                            f"O valor informado manualmente para municipio foi '{atual}', "
                            f"mas o cadastro do nucleo indica '{expected}'. O municipio oficial sera preservado."
                        ),
                        "code": "municipio_divergente_manual",
                    }
                )

        divergencias_master = list(nucleo_master_summary.get("divergencias_municipio", []) or [])
        if divergencias_master:
            amostra = divergencias_master[0]
            alert_items.append(
                {
                    "level": "warning",
                    "title": "Municipio reconciliado pelo cadastro mestre",
                    "message": (
                        f"O cadastro mestre prevaleceu para o nucleo "
                        f"'{amostra.get('nucleo_oficial', amostra.get('nucleo_detectado_texto', '-'))}'. "
                        f"Municipio oficial: '{amostra.get('municipio_oficial', '-')}'."
                    ),
                    "code": "municipio_reconciliado_cadastro",
                }
            )

        nucleos_nao_cadastrados = list(nucleo_master_summary.get("nucleos_nao_cadastrados", []) or [])
        if nucleos_nao_cadastrados:
            preview = ", ".join(nucleos_nao_cadastrados[:3])
            if len(nucleos_nao_cadastrados) > 3:
                preview += ", ..."
            alert_items.append(
                {
                    "level": "warning",
                    "title": "Nucleos sem cadastro mestre",
                    "message": (
                        f"Foram detectados {len(nucleos_nao_cadastrados)} nucleo(s) sem cadastro: {preview}. "
                        "Cadastre-os para consolidar municipio oficial."
                    ),
                    "code": "nucleo_sem_cadastro",
                }
            )

        conflitos_sem_cadastro = list(nucleo_master_summary.get("conflitos_nucleo_sem_cadastro", []) or [])
        if conflitos_sem_cadastro:
            conflito = conflitos_sem_cadastro[0]
            municipios = ", ".join(list(conflito.get("municipios_detectados", []) or [])[:3])
            alert_items.append(
                {
                    "level": "critical",
                    "title": "Nucleo novo com municipios conflitantes",
                    "message": (
                        f"O nucleo '{conflito.get('nucleo', '-')}' apareceu com municipios divergentes "
                        f"({municipios}). O municipio final foi limpo ate o cadastro oficial."
                    ),
                    "code": "nucleo_sem_cadastro_conflitante",
                }
            )

        missing_secondary = [f for f in missing_main if f not in {"nucleo", "municipio"}]
        if missing_secondary:
            alert_items.append(
                {
                    "level": "warning",
                    "title": "Campos principais faltantes",
                    "message": f"Campos ainda vazios: {', '.join(missing_secondary)}.",
                    "code": "missing_main_fields",
                }
            )

        if missing_main and not missing_secondary:
            alert_items.append(
                {
                    "level": "info",
                    "title": "Campos principais com pendencias",
                    "message": f"Campos faltantes detectados: {missing_text}.",
                    "code": "missing_fields_info",
                }
            )

        unmapped = list(unmapped_rows) if unmapped_rows is not None else self.collect_unmapped(parsed)
        unmapped_count = len(unmapped)
        if unmapped_count > 0:
            level = "critical" if unmapped_count >= 5 else "warning"
            term_stats = self.collect_unmapped_term_stats(unmapped, max_items=3)
            recurring_hint = ""
            if term_stats:
                recurring_hint = " Termos recorrentes na previa: " + ", ".join(
                    f"{str(t.get('termo', '')).strip()} ({t.get('ocorrencias', 0)})" for t in term_stats
                ) + "."
            alert_items.append(
                {
                    "level": level,
                    "title": "Servicos nao mapeados",
                    "message": f"Foram detectados {unmapped_count} item(ns) com servico_nao_mapeado.{recurring_hint}",
                    "code": "servico_nao_mapeado",
                }
            )

        incomplete = list(incomplete_rows) if incomplete_rows is not None else self.collect_incomplete_execution(parsed)
        incomplete_count = len(incomplete)
        if incomplete_count > 0:
            level = "critical" if incomplete_count >= 5 else "warning"
            alert_items.append(
                {
                    "level": level,
                    "title": "Linhas de execucao incompletas",
                    "message": f"Foram detectadas {incomplete_count} linha(s) de execucao incompleta(s).",
                    "code": "execucao_incompleta",
                }
            )

        if int(lists.get("logradouros_count", 0) or 0) > 1:
            alert_items.append(
                {
                    "level": "info",
                    "title": "Multiplos logradouros detectados",
                    "message": f"Foram identificados {lists.get('logradouros_count', 0)} logradouros nesta previa.",
                    "code": "multiplos_logradouros",
                }
            )

        if int(lists.get("equipes_count", 0) or 0) > 1:
            alert_items.append(
                {
                    "level": "info",
                    "title": "Multiplas equipes detectadas",
                    "message": f"Foram identificadas {lists.get('equipes_count', 0)} equipes nesta previa.",
                    "code": "multiplas_equipes",
                }
            )

        if int(context.get("nucleos_count", 0) or 0) > 1:
            alert_items.append(
                {
                    "level": "info",
                    "title": "Multiplos nucleos detectados",
                    "message": f"A mensagem possui {context.get('nucleos_count', 0)} nucleos. Revise o ajuste por nucleo antes de gerar.",
                    "code": "multiplos_nucleos",
                }
            )

        if not parsed.get("execucao", []):
            alert_items.append(
                {
                    "level": "info",
                    "title": "Sem itens de execucao",
                    "message": "Bloco EXECUCAO vazio: os arquivos serao gerados sem itens de execucao.",
                    "code": "sem_execucao",
                }
            )

        if not alert_items:
            alert_items.append(
                {
                    "level": "info",
                    "title": "Sem pendencias criticas",
                    "message": "A previa esta consistente para gerar os arquivos.",
                    "code": "sem_pendencias",
                }
            )

        return alert_items

    def validate_preview(self, parsed: dict) -> List[str]:
        return [a["message"] for a in self.build_alert_items(parsed)]

    def build_preview(self, raw_message: str, manual_fields: Dict[str, object]) -> dict:
        parsed, parser_mode = self.parse_message(raw_message)
        parsed = self.reconcile_with_nucleo_master(parsed)
        parsed = self.apply_overrides(parsed, manual_fields, force_all=False)
        parsed = self.reconcile_with_nucleo_master(parsed)

        main_fields = self.extract_main_fields(parsed)
        context_overview = self.collect_context_overview(parsed)
        main_lists = self.build_main_lists(main_fields, context_overview)
        nucleo_groups = self.build_nucleo_groups(parsed)

        unmapped_rows = self.enrich_unmapped_for_review(self.collect_unmapped(parsed))
        incomplete_rows = self.collect_incomplete_execution(parsed)
        municipio_profile_status = self._municipio_profile_status(main_fields)
        manual_municipio_status = {"is_divergent": False, "expected_municipio": "", "provided_municipio": ""}

        manual_nucleo = str(manual_fields.get("nucleo", "") or "").strip() or str(
            parsed.get("nucleo_detectado_texto", "") or parsed.get("nucleo", "")
        ).strip()
        manual_municipio = str(manual_fields.get("municipio", "") or "").strip()
        if manual_nucleo and manual_municipio:
            manual_municipio_status = self._municipio_profile_status(
                {"nucleo": manual_nucleo, "municipio": manual_municipio}
            )
            manual_municipio_status["provided_municipio"] = manual_municipio

        alert_items = self.build_alert_items(
            parsed,
            main_fields=main_fields,
            main_lists=main_lists,
            context_overview=context_overview,
            unmapped_rows=unmapped_rows,
            incomplete_rows=incomplete_rows,
            manual_municipio_status=manual_municipio_status,
        )

        return {
            "raw_message": raw_message,
            "parser_mode": parser_mode,
            "parsed": parsed,
            "main_fields": main_fields,
            "main_lists": main_lists,
            "missing_main_fields": self._missing_main_fields(main_fields),
            "alerts": [a["message"] for a in alert_items],
            "alert_items": alert_items,
            "alerts_by_priority": self.split_alerts_by_priority(alert_items),
            "summary": self.summarize_preview(parsed),
            "unmapped_rows": unmapped_rows,
            "unmapped_term_stats": self.collect_unmapped_term_stats(unmapped_rows),
            "incomplete_rows": incomplete_rows,
            "context_overview": context_overview,
            "nucleo_groups": nucleo_groups,
            "nucleo_reference": self.get_nucleo_reference_ui(),
            "municipio_profile_status": municipio_profile_status,
            "service_catalog": self.get_service_catalog_ui(),
        }

    def save_draft(self, draft_payload: dict) -> str:
        draft_id = f"{datetime.now():%Y%m%d%H%M%S}_{uuid.uuid4().hex[:8]}"
        draft_path = self.draft_dir / f"{draft_id}.json"
        with draft_path.open("w", encoding="utf-8") as f:
            json.dump(draft_payload, f, ensure_ascii=False, indent=2)
        return draft_id

    def load_draft(self, draft_id: str) -> dict:
        draft_path = self.draft_dir / f"{draft_id}.json"
        if not draft_path.exists():
            raise FileNotFoundError(f"Rascunho nao encontrado: {draft_id}")
        with draft_path.open("r", encoding="utf-8") as f:
            return json.load(f)

    def _create_output_dir(self) -> Path:
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
        output_dir = self.outputs_root / f"saida_web_{stamp}"
        output_dir.mkdir(parents=True, exist_ok=True)
        return output_dir

    def _history_headers(self) -> List[str]:
        return [
            "processed_at",
            "obra_data",
            "nucleo",
            "nucleo_detectado_texto",
            "nucleo_oficial",
            "logradouro",
            "municipio",
            "municipio_detectado_texto",
            "municipio_oficial",
            "nucleo_status_cadastro",
            "equipe",
            "status",
            "contract_id",
            "contract_label",
            "output_dir",
            "base_gerencial_path",
            "master_dir",
            "nao_mapeados",
            "alertas",
            "mensagem",
        ]

    def _append_history_to_db(self, row_data: dict, generated_files: List[dict]) -> bool:
        if self.db_manager is None:
            return False
        headers = self._history_headers()
        payload = {h: str(row_data.get(h, "") or "") for h in headers}
        try:
            with self.db_manager.connection() as conn:
                try:
                    with conn.cursor() as cur:
                        cur.execute(
                            """
                            INSERT INTO processing_history (
                                processed_at, obra_data, nucleo, nucleo_detectado_texto, nucleo_oficial,
                                logradouro, municipio, municipio_detectado_texto, municipio_oficial,
                                nucleo_status_cadastro, equipe, status, contract_id, contract_label,
                                output_dir, base_gerencial_path, master_dir, nao_mapeados, alertas, mensagem,
                                generated_files_json
                            ) VALUES (
                                %s, %s, %s, %s, %s,
                                %s, %s, %s, %s,
                                %s, %s, %s, %s, %s,
                                %s, %s, %s, %s, %s, %s,
                                %s::jsonb
                            )
                            """,
                            (
                                payload["processed_at"],
                                payload["obra_data"],
                                payload["nucleo"],
                                payload["nucleo_detectado_texto"],
                                payload["nucleo_oficial"],
                                payload["logradouro"],
                                payload["municipio"],
                                payload["municipio_detectado_texto"],
                                payload["municipio_oficial"],
                                payload["nucleo_status_cadastro"],
                                payload["equipe"],
                                payload["status"],
                                payload["contract_id"],
                                payload["contract_label"],
                                payload["output_dir"],
                                payload["base_gerencial_path"],
                                payload["master_dir"],
                                payload["nao_mapeados"],
                                payload["alertas"],
                                payload["mensagem"],
                                json.dumps(generated_files or [], ensure_ascii=False),
                            ),
                        )
                    conn.commit()
                    return True
                except Exception:
                    conn.rollback()
                    raise
        except Exception:
            return False

    def _read_history_from_db(self, limit: int) -> List[dict]:
        if self.db_manager is None:
            return []

        safe_limit = max(1, min(int(limit), 10000))
        try:
            with self.db_manager.connection() as conn:
                with conn.cursor() as cur:
                    cur.execute(
                        """
                        SELECT
                            processed_at, obra_data, nucleo, nucleo_detectado_texto, nucleo_oficial,
                            logradouro, municipio, municipio_detectado_texto, municipio_oficial,
                            nucleo_status_cadastro, equipe, status, contract_id, contract_label,
                            output_dir, base_gerencial_path, master_dir, nao_mapeados, alertas, mensagem,
                            COALESCE(generated_files_json::text, '[]') AS generated_files_json
                        FROM processing_history
                        ORDER BY id DESC
                        LIMIT %s
                        """,
                        (safe_limit,),
                    )
                    fetched = cur.fetchall() or []
                    cols = [str(d[0]) for d in (cur.description or [])]
        except Exception:
            return []

        rows: List[dict] = []
        for item in fetched:
            row = {cols[i]: item[i] for i in range(min(len(cols), len(item)))}
            generated_raw = row.get("generated_files_json", "[]")
            parsed_generated: List[dict] = []
            if isinstance(generated_raw, (list, tuple)):
                parsed_generated = [dict(x) for x in generated_raw if isinstance(x, dict)]
            else:
                try:
                    maybe = json.loads(str(generated_raw or "[]"))
                    if isinstance(maybe, list):
                        parsed_generated = [dict(x) for x in maybe if isinstance(x, dict)]
                except Exception:
                    parsed_generated = []
            row["generated_files"] = parsed_generated
            rows.append({k: ("" if v is None else str(v)) for k, v in row.items() if k != "generated_files"})
            rows[-1]["generated_files"] = parsed_generated
        return rows

    def _ensure_history_schema(self, headers: List[str]) -> None:
        if not self.history_file.exists() or self.history_file.stat().st_size == 0:
            return

        try:
            with self.history_file.open("r", encoding="utf-8-sig", newline="") as f:
                reader = csv.reader(f)
                existing_headers = next(reader, [])
        except Exception:
            return

        existing_headers = [str(h or "").strip() for h in existing_headers]
        if existing_headers == headers:
            return

        try:
            with self.history_file.open("r", encoding="utf-8-sig", newline="") as f:
                old_rows = list(csv.DictReader(f))
        except Exception:
            old_rows = []

        migrated: List[dict] = []
        for old in old_rows:
            clean_row = {}
            for h in headers:
                clean_row[h] = str(old.get(h, "") or "")
            migrated.append(clean_row)

        tmp_file = self.history_file.with_suffix(self.history_file.suffix + ".tmp")
        try:
            with tmp_file.open("w", encoding="utf-8-sig", newline="") as f:
                writer = csv.DictWriter(f, fieldnames=headers)
                writer.writeheader()
                if migrated:
                    writer.writerows(migrated)
            tmp_file.replace(self.history_file)
        finally:
            if tmp_file.exists():
                try:
                    tmp_file.unlink()
                except Exception:
                    pass

    def _to_file_uri(self, raw_path: object) -> str:
        text = str(raw_path or "").strip()
        if not text:
            return ""
        try:
            return Path(text).resolve().as_uri()
        except Exception:
            return ""

    def _path_within_outputs_root(self, path: Path) -> bool:
        try:
            path.resolve().relative_to(self.outputs_root.resolve())
            return True
        except Exception:
            return False

    def _relative_result_path(self, path: Path) -> str:
        try:
            relative = path.resolve().relative_to(self.outputs_root.resolve())
        except Exception:
            return ""
        return relative.as_posix()

    def _classify_generated_file(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        if suffix in {".xlsx", ".xls", ".csv"}:
            return "spreadsheet"
        if suffix in {".pdf", ".doc", ".docx", ".txt"}:
            return "report"
        return "other"

    def _infer_municipio_from_output(self, output_dir: Path) -> str:
        rows = reconcile_rows_with_registry(
            self._read_csv_rows(output_dir / "execucao.csv"),
            self.nucleo_reference,
        )
        municipios: List[str] = []
        seen = set()
        for row in rows:
            municipio = self._resolve_single_municipio_for_nucleo(
                row.get("nucleo", ""),
                explicit_oficial=row.get("municipio_oficial", ""),
                detected_values=self._split_multi_text(row.get("municipio", "")),
            )
            if not municipio:
                continue
            key = normalizar_texto(municipio)
            if not key or key in seen:
                continue
            seen.add(key)
            municipios.append(municipio)

        if not municipios:
            return ""
        if len(municipios) == 1:
            return municipios[0]
        return ""

    def _list_generated_files(self, output_dir: Path, base_gerencial_path: Path | None = None) -> List[dict]:
        generated: List[dict] = []
        seen = set()

        def _append_file(file_path: Path) -> None:
            if not file_path.exists() or not file_path.is_file():
                return
            abs_path = str(file_path.resolve())
            if abs_path in seen:
                return
            seen.add(abs_path)
            generated.append(
                {
                    "name": file_path.name,
                    "path": abs_path,
                    "uri": self._to_file_uri(abs_path),
                    "relative_path": self._relative_result_path(file_path),
                    "category": self._classify_generated_file(file_path),
                }
            )

        if output_dir.exists() and output_dir.is_dir():
            expected = [
                "execucao.csv",
                "frentes.csv",
                "ocorrencias.csv",
                "observacoes.csv",
                "servico_nao_mapeado.csv",
                "base_gerencial.xlsx",
                "entrada_original.txt",
            ]
            for name in expected:
                _append_file(output_dir / name)

            reports_dir = output_dir / "relatorios_nucleos"
            if reports_dir.exists() and reports_dir.is_dir():
                for file_path in sorted(reports_dir.rglob("*")):
                    _append_file(file_path)

        if base_gerencial_path and base_gerencial_path.exists() and base_gerencial_path.is_file():
            _append_file(base_gerencial_path)

        return generated

    def _append_history(self, row: dict) -> None:
        headers = self._history_headers()
        self.history_file.parent.mkdir(parents=True, exist_ok=True)
        self._ensure_history_schema(headers)
        row_data = dict(row or {})
        row_data["equipe"] = self._first_team(row_data.get("equipe", ""))

        output_dir_raw = str(row_data.get("output_dir", "") or "").strip()
        base_gerencial_raw = str(row_data.get("base_gerencial_path", "") or "").strip()
        output_dir_path = Path(output_dir_raw) if output_dir_raw else None
        base_gerencial_path = Path(base_gerencial_raw) if base_gerencial_raw else None
        generated_files = self._list_generated_files(
            output_dir_path if output_dir_path else Path("__sem_saida__"),
            base_gerencial_path,
        )
        self._append_history_to_db(row_data, generated_files)

        exists = self.history_file.exists() and self.history_file.stat().st_size > 0
        with self.history_file.open("a", encoding="utf-8-sig", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=headers)
            if not exists:
                writer.writeheader()
            clean = {}
            for h in headers:
                val = str(row_data.get(h, "") or "")
                clean[h] = " ".join(val.split())
            writer.writerow(clean)

    def resolve_result_file(self, relative_path: str) -> Path | None:
        safe_relative = str(relative_path or "").strip().replace("\\", "/").lstrip("/")
        if not safe_relative:
            return None
        candidate = (self.outputs_root / safe_relative).resolve()
        if not self._path_within_outputs_root(candidate):
            return None
        if not candidate.exists() or not candidate.is_file():
            return None
        return candidate

    def read_history(self, limit: int = 100) -> List[dict]:
        rows = self._read_history_from_db(limit)
        if not rows:
            if not self.history_file.exists():
                return []
            with self.history_file.open("r", encoding="utf-8-sig", newline="") as f:
                rows = list(csv.DictReader(f))
            rows.reverse()

        out: List[dict] = []
        municipio_cache: Dict[str, str] = {}
        generated_files_cache: Dict[str, List[dict]] = {}

        for raw_row in rows[:limit]:
            row = {str(k): str(v or "") for k, v in raw_row.items() if k is not None}
            row["equipe"] = self._first_team(row.get("equipe", ""))
            row["nucleo_detectado_texto"] = str(row.get("nucleo_detectado_texto", "") or "").strip()
            row["municipio_detectado_texto"] = str(row.get("municipio_detectado_texto", "") or "").strip()
            row["nucleo_oficial"] = str(row.get("nucleo_oficial", "") or "").strip()
            row["municipio_oficial"] = str(row.get("municipio_oficial", "") or "").strip()
            row["nucleo_status_cadastro"] = str(row.get("nucleo_status_cadastro", "") or "").strip()
            if row["nucleo_oficial"]:
                row["nucleo"] = row["nucleo_oficial"]
            status = str(row.get("status", "") or "").strip()
            status_level = "erro" if status.lower().startswith("erro") else "sucesso"
            alert_list = [x.strip() for x in str(row.get("alertas", "") or "").split("|") if x.strip()]
            alert_count = len(alert_list)
            alert_preview = ""
            if alert_list:
                alert_preview = " | ".join(alert_list[:2])
                if alert_count > 2:
                    alert_preview += f" (+{alert_count - 2})"

            output_dir_raw = str(row.get("output_dir", "") or "").strip()
            base_gerencial_raw = str(row.get("base_gerencial_path", "") or "").strip()
            output_dir_path = Path(output_dir_raw) if output_dir_raw else None
            base_gerencial_path = Path(base_gerencial_raw) if base_gerencial_raw else None
            inferred_municipio = ""

            if output_dir_path:
                cache_key = str(output_dir_path)
                if cache_key not in municipio_cache:
                    municipio_cache[cache_key] = self._infer_municipio_from_output(output_dir_path)
                inferred_municipio = municipio_cache.get(cache_key, "")

            municipio = self._resolve_history_final_municipio(
                row,
                output_dir=output_dir_path,
                inferred_municipio=inferred_municipio,
            )

            row["nucleo"] = self._join_history_display_values(
                self._split_nucleo_text_for_aggregation(row.get("nucleo", "")),
                canonicalizer=self._canonicalize_nucleo_for_aggregation,
            ) or str(row.get("nucleo", "") or "").strip()
            row["logradouro"] = self._join_history_display_values(
                self._explode_history_values(row.get("logradouro", ""))
            ) or str(row.get("logradouro", "") or "").strip()
            row["equipe"] = self._join_history_display_values(
                self._explode_history_values(row.get("equipe", "")),
                canonicalizer=self._canonicalize_equipe_for_aggregation,
            ) or self._first_team(row.get("equipe", ""))

            generated_files = []
            raw_generated = raw_row.get("generated_files", []) if isinstance(raw_row, dict) else []
            if isinstance(raw_generated, list):
                generated_files = [dict(item) for item in raw_generated if isinstance(item, dict)]

            if not generated_files:
                cache_key = f"{output_dir_raw}|{base_gerencial_raw}"
                if cache_key not in generated_files_cache:
                    generated_files_cache[cache_key] = self._list_generated_files(
                        output_dir_path if output_dir_path else Path("__sem_saida__"),
                        base_gerencial_path,
                    )
                generated_files = generated_files_cache[cache_key]
            names_preview = [str(item.get("name", "") or "").strip() for item in generated_files if str(item.get("name", "") or "").strip()]
            generated_preview = ", ".join(names_preview[:3])
            if len(names_preview) > 3:
                generated_preview += f", +{len(names_preview) - 3}"

            out.append(
                {
                    **row,
                    "municipio": municipio,
                    "status_level": status_level,
                    "alert_count": str(alert_count),
                    "has_alerts": "sim" if alert_count > 0 else "nao",
                    "alert_list": alert_list,
                    "alert_preview": alert_preview,
                    "nao_mapeados": str(row.get("nao_mapeados", "") or "0"),
                    "output_name": Path(output_dir_raw).name if output_dir_raw else "",
                    "base_gerencial_name": Path(base_gerencial_raw).name if base_gerencial_raw else "",
                    "output_dir_uri": self._to_file_uri(output_dir_raw),
                    "base_gerencial_uri": self._to_file_uri(base_gerencial_raw),
                    "generated_files": generated_files,
                    "generated_files_preview": generated_preview,
                }
            )
        return out

    def list_processing_results(self, limit: int = 200) -> List[dict]:
        try:
            safe_limit = int(limit)
        except Exception:
            safe_limit = 200
        safe_limit = max(1, min(safe_limit, 1000))

        results: List[dict] = []
        entities_cache: Dict[str, Dict[str, List[str]]] = {}
        contract_cache: Dict[str, str] = {}
        for row in self.read_history(limit=safe_limit):
            row_data = dict(row)
            output_dir_raw = str(row_data.get("output_dir", "") or "").strip()
            entities: Dict[str, List[str]] = {"nucleos": [], "equipes": [], "municipios": []}
            if output_dir_raw:
                if output_dir_raw not in entities_cache:
                    entities_cache[output_dir_raw] = self._collect_management_entities_for_row(row_data)
                entities = entities_cache[output_dir_raw]
            else:
                entities = self._collect_management_entities_for_row(row_data)

            nucleos_reais = list(entities.get("nucleos", []) or [])
            municipios_reais = list(entities.get("municipios", []) or [])

            nucleo_display = self._join_history_display_values(
                nucleos_reais,
                canonicalizer=self._canonicalize_nucleo_for_aggregation,
            )
            if nucleo_display:
                row_data["nucleo"] = nucleo_display

            municipio_display = self._join_history_display_values(
                municipios_reais,
                canonicalizer=self._canonicalize_municipio_for_aggregation,
            )
            if municipio_display:
                row_data["municipio"] = municipio_display

            generated_files = list(row.get("generated_files", []) or [])
            report_files = [item for item in generated_files if item.get("category") == "report"]
            spreadsheet_files = [item for item in generated_files if item.get("category") == "spreadsheet"]

            contract_id_text = str(row.get("contract_id", "") or "").strip()
            contract_label_text = str(row.get("contract_label", "") or "").strip()
            if not contract_label_text and output_dir_raw:
                if output_dir_raw not in contract_cache:
                    extracted = ""
                    try:
                        consolidated = Path(output_dir_raw) / "relatorio_consolidado.json"
                        if consolidated.exists():
                            payload = json.loads(consolidated.read_text(encoding="utf-8"))
                            extracted = str((payload or {}).get("contrato", "") or "").strip()
                    except Exception:
                        extracted = ""
                    contract_cache[output_dir_raw] = extracted
                contract_label_text = contract_cache.get(output_dir_raw, "")

            if contract_label_text:
                row_data["contract_label"] = contract_label_text
            if contract_id_text:
                row_data["contract_id"] = contract_id_text

            primary_download = None
            if spreadsheet_files:
                primary_download = spreadsheet_files[0]
            elif report_files:
                primary_download = report_files[0]
            elif generated_files:
                primary_download = generated_files[0]

            results.append(
                {
                    **row_data,
                    "result_id": str(row.get("output_name", "") or "").strip(),
                    "process_type": "Processamento de entrada",
                    "contract_id": contract_id_text,
                    "contract_label": contract_label_text,
                    "nucleos_reais": nucleos_reais,
                    "municipios_reais": municipios_reais,
                    "report_files": report_files,
                    "spreadsheet_files": spreadsheet_files,
                    "primary_download_relative_path": str((primary_download or {}).get("relative_path", "") or ""),
                    "primary_download_name": str((primary_download or {}).get("name", "") or ""),
                    "has_files": bool(generated_files),
                }
            )
        return results

    def get_processing_result(self, output_name: str) -> dict | None:
        target = str(output_name or "").strip()
        if not target:
            return None
        for row in self.list_processing_results(limit=1000):
            if str(row.get("result_id", "") or "").strip() == target:
                return row
        return None

    def _parse_date_flexible(self, value: object) -> date | None:
        text = str(value or "").strip()
        if not text:
            return None
        for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
            try:
                return datetime.strptime(text, fmt).date()
            except ValueError:
                continue
        return None

    def _parse_quantity(self, value: object) -> float:
        text = str(value or "").strip()
        if not text:
            return 0.0
        cleaned = text.replace(" ", "")
        # Keep only number separators to avoid "10m" and similar noise.
        cleaned = re.sub(r"[^0-9,.\-]", "", cleaned)
        if not cleaned:
            return 0.0

        if "," in cleaned and "." in cleaned:
            if cleaned.rfind(",") > cleaned.rfind("."):
                cleaned = cleaned.replace(".", "").replace(",", ".")
            else:
                cleaned = cleaned.replace(",", "")
        elif "," in cleaned:
            cleaned = cleaned.replace(".", "").replace(",", ".")

        try:
            return float(cleaned)
        except Exception:
            return 0.0

    def _read_csv_rows(self, csv_path: Path) -> List[dict]:
        if not csv_path.exists() or csv_path.stat().st_size == 0:
            return []

        for enc in ("utf-8-sig", "utf-8", "latin-1"):
            try:
                with csv_path.open("r", encoding=enc, newline="") as f:
                    return list(csv.DictReader(f))
            except UnicodeDecodeError:
                continue
            except Exception:
                return []
        return []

    def _match_filter_text(self, value: object, normalized_filter: str) -> bool:
        if not normalized_filter:
            return True
        text_norm = normalizar_texto(str(value or ""))
        return normalized_filter in text_norm

    def _nucleo_nome_base(self, value: object) -> str:
        raw = " ".join(str(value or "").split()).strip()
        if not raw:
            return ""

        # Collapse operational suffixes like:
        # "Mississipi - Esgoto", "Mississipi – Esgoto", "Mississipi — Esgoto".
        # Also tolerates mojibake variants after unsafe decoding.
        split_match = re.split(r"\s*(?:-|–|—|â€“|â€”)\s*", raw, maxsplit=1)
        if len(split_match) < 2:
            return raw
        base = str(split_match[0] or "").strip()
        return base or raw

    def _looks_like_logradouro(self, value: object) -> bool:
        norm = normalizar_texto(str(value or "")).strip()
        if not norm:
            return False
        prefixes = (
            "viela",
            "rua",
            "travessa",
            "tv",
            "avenida",
            "av",
            "estrada",
            "rodovia",
            "acesso",
            "alameda",
            "beco",
            "praca",
            "largo",
        )
        return any(norm == p or norm.startswith(f"{p} ") for p in prefixes)

    def _canonicalize_nucleo_for_aggregation(
        self, value: object, fallback_nucleo: object = ""
    ) -> str:
        raw = " ".join(str(value or "").split()).strip()
        base = self._nucleo_nome_base(raw)

        if base and self._looks_like_logradouro(base):
            base = self._nucleo_nome_base(fallback_nucleo)

        if not base:
            return ""

        profile = self._get_nucleo_profile(base)
        if profile:
            if str(profile.get("status", "") or "").strip().lower() == "inativo":
                return ""
            return str(profile.get("nucleo", "") or "").strip() or base

        return base

    def _split_nucleo_text_for_aggregation(self, value: object) -> List[str]:
        text = str(value or "").replace("\r\n", "\n").replace("\r", "\n").strip()
        if not text:
            return []

        if text.lower().startswith("multiplos") and ":" in text:
            text = text.split(":", 1)[1]

        text = text.replace("|", "/")
        parts = re.split(r"\s*(?:/|;|,|\n)+\s*", text)

        out: List[str] = []
        seen = set()
        for part in parts:
            clean = re.sub(r"^[\u2022\-\*]+\s*", "", str(part or "")).strip(" -")
            clean = " ".join(clean.split()).strip()
            if not clean:
                continue
            key = normalizar_texto(clean)
            if not key or key in seen:
                continue
            seen.add(key)
            out.append(clean)
        return out

    def _first_team(self, value: object) -> str:
        return extrair_primeira_equipe(value)

    def _normalize_team_list(self, values: List[object]) -> List[str]:
        out: List[str] = []
        seen = set()
        for value in values:
            team = self._first_team(value)
            if not team:
                continue
            key = normalizar_texto(team)
            if not key or key in seen:
                continue
            seen.add(key)
            out.append(team)
        return out

    def _canonicalize_municipio_for_aggregation(self, value: object) -> str:
        raw = " ".join(str(value or "").split()).strip()
        if not raw:
            return ""

        key = normalizar_texto(raw).strip()
        if not key:
            return ""

        ref_name = self.nucleo_reference.get("municipio_by_key", {}).get(key)
        if ref_name:
            return str(ref_name).strip()

        canonical_map = {
            "carapicuiba": "Carapicuíba",
            "sao paulo": "São Paulo",
        }
        if key in canonical_map:
            return canonical_map[key]

        words = [w for w in key.split(" ") if w]
        return " ".join(w.capitalize() for w in words)

    def _resolve_single_municipio_for_nucleo(
        self,
        nucleo: object,
        *,
        explicit_oficial: object = "",
        detected_values: List[object] | None = None,
    ) -> str:
        nucleo_final = self._canonicalize_nucleo_for_aggregation(nucleo)
        profile = self._get_nucleo_profile(nucleo_final) if nucleo_final else None

        municipio = self._canonicalize_municipio_for_aggregation(
            str((profile or {}).get("municipio", "") or "")
        )
        if municipio:
            return municipio

        municipio = self._canonicalize_municipio_for_aggregation(explicit_oficial)
        if municipio:
            return municipio

        for raw in list(detected_values or []):
            municipio = self._canonicalize_municipio_for_aggregation(raw)
            if municipio:
                return municipio
        return ""

    def _resolve_final_municipio_for_nucleos(
        self,
        nucleos: List[object] | None,
        *,
        explicit_oficial: object = "",
        detected_values: List[object] | None = None,
    ) -> str:
        canonical_nucleos: List[str] = []
        seen = set()
        for raw in list(nucleos or []):
            nucleo = self._canonicalize_nucleo_for_aggregation(raw)
            key = normalizar_texto(nucleo)
            if not key or key in seen:
                continue
            seen.add(key)
            canonical_nucleos.append(nucleo)

        if len(canonical_nucleos) != 1:
            return ""

        nucleo = canonical_nucleos[0]
        oficial = self._canonicalize_municipio_for_aggregation(explicit_oficial)
        if oficial:
            return self._resolve_single_municipio_for_nucleo(
                nucleo,
                explicit_oficial=oficial,
                detected_values=[],
            )

        profile = self._get_nucleo_profile(nucleo)
        profile_municipio = self._canonicalize_municipio_for_aggregation(
            str((profile or {}).get("municipio", "") or "")
        )
        if profile_municipio:
            return profile_municipio

        detected_map: Dict[str, str] = {}
        for raw in list(detected_values or []):
            municipio = self._canonicalize_municipio_for_aggregation(raw)
            key = normalizar_texto(municipio)
            if key and key not in detected_map:
                detected_map[key] = municipio

        if len(detected_map) == 1:
            return next(iter(detected_map.values()))
        return ""

    def _resolve_history_final_municipio(
        self,
        row: dict,
        *,
        output_dir: Path | None = None,
        inferred_municipio: object = "",
    ) -> str:
        detected_values = self._explode_history_values(
            row.get("municipio_oficial", "") or row.get("municipio", "")
        )
        inferred = self._canonicalize_municipio_for_aggregation(inferred_municipio)
        if not inferred and not detected_values and output_dir:
            inferred = self._infer_municipio_from_output(output_dir)
        if not detected_values and inferred:
            detected_values = [inferred]

        row_nucleos = self._split_nucleo_text_for_aggregation(
            row.get("nucleo_oficial", "") or row.get("nucleo", "")
        )
        return self._resolve_final_municipio_for_nucleos(
            row_nucleos,
            explicit_oficial=row.get("municipio_oficial", ""),
            detected_values=detected_values,
        )

    def _resolve_institutional_row_municipio(
        self,
        row: dict,
        row_nucleos: List[str],
        entities: dict,
    ) -> str:
        return self._resolve_final_municipio_for_nucleos(
            row_nucleos,
            explicit_oficial=row.get("municipio_oficial", ""),
            detected_values=list(entities.get("municipios", []) or []),
        )

    def _join_history_display_values(
        self, values: List[object], *, canonicalizer=None
    ) -> str:
        out: List[str] = []
        seen = set()
        for raw in values:
            value = canonicalizer(raw) if canonicalizer else str(raw or "").strip()
            value = str(value or "").strip()
            if not value:
                continue
            key = normalizar_texto(value)
            if not key or key in seen:
                continue
            seen.add(key)
            out.append(value)
        return " / ".join(out)

    def _canonicalize_equipe_for_aggregation(self, value: object) -> str:
        raw = " ".join(self._first_team(value).split()).strip(" -")
        if not raw:
            return ""

        text = re.sub(r"^[\u2022\-\*]+\s*", "", raw).strip()
        text = re.sub(
            r"^(?:resp(?:onsavel)?\.?|responsavel(?:\s+tecnico)?)\s*[:\-]?\s*",
            "",
            text,
            flags=re.IGNORECASE,
        )

        code_match = re.match(
            r"^equipes?\s*[:\-]?\s*(\d{1,3})$",
            text,
            flags=re.IGNORECASE,
        )
        if code_match:
            code = code_match.group(1)
            return f"Equipe {code}"

        text = re.sub(r"^equipes?\b(?:\s*[:\-]\s*|\s+)", "", text, flags=re.IGNORECASE)
        text = " ".join(text.split()).strip(" -")
        if not text:
            return ""

        typo_aliases = {
            "wesley": "Weslyn",
        }
        norm_key = normalizar_texto(text)
        if norm_key in typo_aliases:
            return typo_aliases[norm_key]

        return text

    def _fix_mojibake_text(self, value: object) -> str:
        text = str(value or "").strip()
        if not text:
            return ""
        if "Ã" not in text and "Â" not in text:
            return text
        try:
            repaired = text.encode("latin-1").decode("utf-8")
            if repaired:
                return repaired
        except Exception:
            return text
        return text

    def _sanitize_institutional_text(self, value: object) -> str:
        text = self._fix_mojibake_text(value)
        if not text:
            return ""
        text = re.sub(r"[\u2705\u26A0\uFE0F]", " ", text)
        text = re.sub(r"[\U0001F300-\U0001FAFF]", " ", text)
        text = re.sub(r"[^0-9A-Za-zÀ-ÿ\s,.;:/()\-+%_'’]", " ", text)
        text = re.sub(r"\s+", " ", text).strip()
        return text

    def _institutional_text_or_fallback(self, value: object, fallback: str = "Não informado") -> str:
        text = self._sanitize_institutional_text(value)
        if not text or text in {"-", "--"}:
            return fallback
        return text

    def _normalize_institutional_note(self, value: object) -> str:
        text = self._sanitize_institutional_text(value)
        if not text or text in {"-", "--"}:
            return "Não informado"
        text = re.sub(r"(?i)^\s*obs\.?\s*:?\s*", "Observação: ", text)
        text = re.sub(r"(?i)^\s*observacao\s*:?\s*", "Observação: ", text)
        text = re.sub(r"(?i)^\s*observação\s*:?\s*", "Observação: ", text)
        return text

    def _title_case_label(self, value: object) -> str:
        text = " ".join(str(value or "").split()).strip()
        if not text:
            return "-"
        connectors = {"de", "da", "do", "das", "dos", "e", "em", "para", "com"}
        out: List[str] = []
        for idx, word in enumerate(text.split(" ")):
            lower = word.lower()
            if lower == "uma":
                out.append("UMA")
                continue
            if idx > 0 and lower in connectors:
                out.append(lower)
                continue
            out.append(lower.capitalize())
        return " ".join(out)

    def _institutional_label(self, value: object, domain: str = "geral") -> str:
        raw = self._sanitize_institutional_text(value)
        if not raw:
            return "Não informado"
        text = re.sub(r"\s+", " ", raw.replace("_", " ")).strip(" -")
        norm = normalizar_texto(text)
        if not norm:
            return "Não informado"

        common_map = {
            "-": "Não informado",
            "--": "Não informado",
            "n a": "Não informado",
            "n/a": "Não informado",
            "nao_informado": "Não informado",
            "nao informado": "Não informado",
            "servico_nao_informado": "Serviço não informado",
            "servico nao informado": "Serviço não informado",
            "servico_nao_mapeado": "Serviço não mapeado",
            "servico nao mapeado": "Serviço não mapeado",
            "nao_mapeado": "Serviço não mapeado",
            "nao mapeado": "Serviço não mapeado",
        }
        service_map = {
            "pra": "PRA",
            "pre": "PRE",
            "hidrometro": "Instalação de hidrômetro",
            "hidrometros": "Instalação de hidrômetros",
            "hidrometros instalados": "Instalação de hidrômetros",
            "instalacao de hidrometro": "Instalação de hidrômetro",
            "instalacao de hidrometros": "Instalação de hidrômetros",
            "intradomiciliares": "Ligações intradomiciliares",
            "ligacoes intradomiciliares": "Ligações intradomiciliares",
            "ligacao intradomiciliar": "Ligação intradomiciliar",
            "ligacao intradomiciliares": "Ligações intradomiciliares",
            "rede agua": "Rede de Água",
            "rede_agua": "Rede de Água",
            "rede de agua": "Rede de Água",
            "prolongamento rede agua": "Prolongamento de Rede de Água",
            "prolongamento de rede": "Prolongamento de Rede de Água",
            "prolongamento de rede agua": "Prolongamento de Rede de Água",
            "rede esgoto": "Rede de Esgoto",
            "rede_esgoto": "Rede de Esgoto",
            "ramal agua": "Ramal de Água",
            "ramais agua": "Ramais de Água",
            "ramal esgoto": "Ramal de esgoto",
            "ramais esgoto": "Ramais de esgoto",
            "interligacao": "Interligação de rede",
            "interligacao de rede": "Interligação de rede",
            "caixa uma": "Instalação de Caixas UMA",
            "caixas uma": "Instalação de Caixas UMA",
            "caixas uma instaladas": "Instalação de Caixas UMA",
            "instalacao de caixas uma": "Instalação de Caixas UMA",
            "instalacao de caixas umas": "Instalação de Caixas UMA",
            "instalacao de caixa dagua": "Instalação de Caixa d'água",
            "instalacao de caixas dagua": "Instalação de Caixas d'água",
            "instalacao de caixa d agua": "Instalação de Caixa d'água",
            "instalacao de caixas d agua": "Instalação de Caixas d'água",
            "adesoes realizadas": "Adesões realizadas",
            "caixas dagua instaladas": "Instalação de Caixas d'água",
            "economias": "Economias",
            "luvas executadas": "Luvas executadas",
            "embutida": "Instalação de caixas UMA (embutida)",
            "mureta": "Instalação de caixas UMA (mureta)",
            "servico complementar": "Serviço complementar",
            "recomposicao de valas": "Recomposição de valas",
            "recomposicao asfaltica": "Recomposição asfáltica",
            "atividade realizada": "Atividade realizada",
            "atividades realizadas": "Atividades realizadas",
        }
        category_map = {
            "hidrometro": "Hidrômetro",
            "intradomiciliar": "Intradomiciliar",
            "rede": "Rede",
            "rede agua": "Rede de Água",
            "rede_agua": "Rede de Água",
            "ramal_agua": "Ramal de Água",
            "ramal_esgoto": "Ramal de esgoto",
            "interligacao": "Interligação",
            "caixa_uma": "Caixa UMA",
            "caixa uma": "Caixa UMA",
            "ligacao_agua": "Ligação de Água",
            "adesao_agua": "Adesão de Água",
            "caixa_dagua": "Caixa d'água",
            "economias": "Economias",
            "rede_esgoto": "Rede de Esgoto",
            "caixa_inspecao": "Caixa de inspeção",
            "recomposicao": "Recomposição",
            "apoio_operacional": "Apoio operacional",
            "insumo": "Insumo",
            "acabamento": "Acabamento",
            "servico_nao_mapeado": "Serviço não mapeado",
        }
        occurrence_map = {
            "interferencia": "Interferência",
            "interferencia com rede de esgoto": "Interferência com rede de esgoto",
            "restricao_operacional": "Restrição operacional",
            "restricao operacional": "Restrição operacional",
            "ocorrencia operacional": "Ocorrência operacional",
            "ocorrencia_operacional": "Ocorrência operacional",
            "imovel fechado": "Imóvel fechado",
            "vistoria": "Vistoria",
            "chuva": "Chuvas",
            "sem_producao": "Sem produção",
            "sem producao": "Sem produção",
        }

        if norm in common_map:
            return common_map[norm]
        if domain == "servico" and norm in service_map:
            return service_map[norm]
        if domain == "categoria" and norm in category_map:
            return category_map[norm]
        if domain == "ocorrencia" and norm in occurrence_map:
            return occurrence_map[norm]
        label = self._title_case_label(text)
        if not label or label in {"-", "--"}:
            return "Não informado"
        return label

    def _summarize_values(self, values: List[str], limit: int = 3) -> str:
        clean = [str(v or "").strip() for v in values if str(v or "").strip()]
        if not clean:
            return "-"
        if len(clean) <= limit:
            return "; ".join(clean)
        return f"{'; '.join(clean[:limit])} (+{len(clean) - limit})"

    def _count_label(self, value: int, singular: str, plural: str) -> str:
        return f"{int(value)} {singular if int(value) == 1 else plural}"

    def _join_natural_phrases(self, items: List[str]) -> str:
        clean = [str(item or "").strip() for item in items if str(item or "").strip()]
        if not clean:
            return ""
        if len(clean) == 1:
            return clean[0]
        if len(clean) == 2:
            return f"{clean[0]} e {clean[1]}"
        return f"{', '.join(clean[:-1])} e {clean[-1]}"

    def _build_dashboard_bar_chart(
        self,
        rows: List[dict],
        *,
        label_field: str,
        value_field: str,
        limit: int = 10,
        domain: str = "",
        value_mode: str = "count",
        unit_field: str = "",
    ) -> dict:
        items: List[dict] = []
        max_value = 0.0
        for idx, row in enumerate(list(rows or [])[:limit], start=1):
            label = str(row.get(label_field, "") or "").strip() or "-"
            if domain:
                label = self._institutional_label(label, domain=domain)
            raw_value = row.get(value_field, 0)
            unit_raw = str(row.get(unit_field, "") or "").strip() if unit_field else ""
            unit = "" if unit_raw in {"-", "--", "n/a", "N/A"} else unit_raw
            value = (
                float(raw_value or 0.0)
                if value_mode == "number"
                else float(self._safe_int(raw_value, default=0, min_value=0, max_value=1000000))
            )
            max_value = max(max_value, value)
            display = (
                self._display_number(value)
                if value_mode == "number"
                else str(int(value))
            )
            value_display = f"{display} {unit}".strip()
            items.append(
                {
                    "rank": idx,
                    "label": label,
                    "value": value,
                    "value_fmt": display,
                    "value_display": value_display,
                    "unit": unit,
                    "meta": str(row.get("categoria", "") or "").strip(),
                }
            )

        if max_value <= 0:
            max_value = 1.0

        for item in items:
            if item["value"] <= 0:
                item["width_pct"] = 0
            else:
                item["width_pct"] = max(10.0, round((item["value"] / max_value) * 100.0, 1))

        return {
            "items": items,
            "has_data": bool(items),
            "max_value_fmt": self._display_number(max_value)
            if value_mode == "number"
            else str(int(max_value)),
        }

    def _build_dashboard_period_chart(self, rows: List[dict]) -> dict:
        items: List[dict] = []
        for row in list(rows or []):
            label = str(row.get("periodo", "") or "").strip() or "-"
            dt = self._parse_date_flexible(label)
            value = self._safe_int(
                row.get("processamentos", 0),
                default=0,
                min_value=0,
                max_value=1000000,
            )
            items.append(
                {
                    "label": label,
                    "short_label": dt.strftime("%d/%m") if dt else label,
                    "value": value,
                    "sort_key": dt or date.min,
                }
            )

        items.sort(key=lambda item: item["sort_key"])
        max_value = max((item["value"] for item in items), default=0)
        scale = max_value or 1
        for item in items:
            item["height_pct"] = max(8.0, round((item["value"] / scale) * 100.0, 1)) if item["value"] else 0
            item["value_fmt"] = str(int(item["value"]))

        return {
            "items": items,
            "has_data": bool(items),
            "max_value_fmt": str(int(max_value)),
        }

    def _build_dashboard_mapping_chart(self, kpis: dict) -> dict:
        mapped = self._safe_int(
            kpis.get("total_mapeados", 0),
            default=0,
            min_value=0,
            max_value=1000000,
        )
        unmapped = self._safe_int(
            kpis.get("total_nao_mapeados", 0),
            default=0,
            min_value=0,
            max_value=1000000,
        )
        total = mapped + unmapped
        mapped_pct = (100.0 * mapped / total) if total else 100.0
        unmapped_pct = 100.0 - mapped_pct if total else 0.0
        return {
            "mapped": mapped,
            "unmapped": unmapped,
            "mapped_pct": mapped_pct,
            "mapped_pct_fmt": self._display_number(mapped_pct),
            "unmapped_pct": unmapped_pct,
            "unmapped_pct_fmt": self._display_number(unmapped_pct),
            "has_data": total > 0,
        }

    def _extract_management_entities_from_output(
        self, output_dir: Path
    ) -> Dict[str, List[str]]:
        nucleos_map: Dict[str, str] = {}
        equipes_map: Dict[str, str] = {}
        municipios_map: Dict[str, str] = {}
        fallback_nucleo = ""

        for csv_name in ("execucao.csv", "frentes.csv", "ocorrencias.csv"):
            rows = reconcile_rows_with_registry(
                self._read_csv_rows(output_dir / csv_name),
                self.nucleo_reference,
            )
            for row in rows:
                row_profile = None
                row_primary_nucleo = str(row.get("nucleo_oficial", "") or row.get("nucleo", "") or "").strip()
                for nucleo_raw in self._split_nucleo_text_for_aggregation(
                    row_primary_nucleo
                ):
                    nucleo = self._canonicalize_nucleo_for_aggregation(
                        nucleo_raw, fallback_nucleo=fallback_nucleo
                    )
                    if nucleo:
                        row_profile = self._get_nucleo_profile(nucleo) or row_profile
                        nuc_key = normalizar_texto(nucleo)
                        if nuc_key and nuc_key not in nucleos_map:
                            nucleos_map[nuc_key] = nucleo
                        if not fallback_nucleo:
                            fallback_nucleo = nucleo

                equipe_raw = self._first_team(row.get("equipe", ""))
                equipe = self._canonicalize_equipe_for_aggregation(equipe_raw)
                if equipe:
                    eqp_key = normalizar_texto(equipe)
                    if eqp_key and eqp_key not in equipes_map:
                        equipes_map[eqp_key] = equipe

                municipio = self._resolve_single_municipio_for_nucleo(
                    row_primary_nucleo,
                    explicit_oficial=(row_profile or {}).get("municipio", "")
                    or row.get("municipio_oficial", ""),
                    detected_values=self._split_multi_text(row.get("municipio", "")),
                )
                if municipio:
                    mun_key = normalizar_texto(municipio)
                    if mun_key and mun_key not in municipios_map:
                        municipios_map[mun_key] = municipio

        return {
            "nucleos": list(nucleos_map.values()),
            "equipes": list(equipes_map.values()),
            "municipios": list(municipios_map.values()),
        }

    def _collect_management_entities_for_row(self, row: dict) -> Dict[str, List[str]]:
        entities = {
            "nucleos": [],
            "equipes": [],
            "municipios": [],
        }

        status_level = str(row.get("status_level", "") or "").strip().lower()
        output_dir_raw = str(row.get("output_dir", "") or "").strip()
        if status_level == "sucesso" and output_dir_raw:
            output_dir = Path(output_dir_raw)
            if output_dir.exists():
                entities = self._extract_management_entities_from_output(output_dir)

        if not entities["nucleos"]:
            fallback_nucleo = ""
            nuclei_map: Dict[str, str] = {}
            source_nucleo = str(row.get("nucleo_oficial", "") or row.get("nucleo", "") or "")
            for raw in self._split_nucleo_text_for_aggregation(source_nucleo):
                nucleo = self._canonicalize_nucleo_for_aggregation(
                    raw, fallback_nucleo=fallback_nucleo
                )
                if not nucleo:
                    continue
                key = normalizar_texto(nucleo)
                if key and key not in nuclei_map:
                    nuclei_map[key] = nucleo
                if not fallback_nucleo:
                    fallback_nucleo = nucleo
            entities["nucleos"] = list(nuclei_map.values())

        if not entities["equipes"]:
            equipes_map: Dict[str, str] = {}
            raw = self._first_team(row.get("equipe", ""))
            equipe = self._canonicalize_equipe_for_aggregation(raw)
            if equipe:
                key = normalizar_texto(equipe)
                if key and key not in equipes_map:
                    equipes_map[key] = equipe
            entities["equipes"] = list(equipes_map.values())

        if not entities["municipios"]:
            municipios_map: Dict[str, str] = {}
            detected_values = self._explode_history_values(row.get("municipio", ""))
            if len(entities.get("nucleos", [])) == 1:
                municipio = self._resolve_single_municipio_for_nucleo(
                    entities["nucleos"][0],
                    explicit_oficial=row.get("municipio_oficial", ""),
                    detected_values=detected_values,
                )
                if municipio:
                    municipios_map[normalizar_texto(municipio)] = municipio
            else:
                official_municipio = str(row.get("municipio_oficial", "") or "").strip()
                source_values = [official_municipio] if official_municipio else detected_values
                for raw in source_values:
                    municipio = self._canonicalize_municipio_for_aggregation(raw)
                    if not municipio:
                        continue
                    key = normalizar_texto(municipio)
                    if key and key not in municipios_map:
                        municipios_map[key] = municipio
            entities["municipios"] = list(municipios_map.values())

        return entities

    def _format_date_range(self, start: date | None, end: date | None) -> str:
        if not start and not end:
            return "-"
        if start and end:
            if start == end:
                return start.strftime("%d/%m/%Y")
            return f"{start:%d/%m/%Y} a {end:%d/%m/%Y}"
        if start:
            return f"A partir de {start:%d/%m/%Y}"
        return f"Ate {end:%d/%m/%Y}"

    def _format_datetime_range(self, start: datetime | None, end: datetime | None) -> str:
        if not start and not end:
            return "-"
        if start and end:
            if start == end:
                return start.strftime("%d/%m/%Y %H:%M")
            return f"{start:%d/%m/%Y %H:%M} a {end:%d/%m/%Y %H:%M}"
        if start:
            return f"A partir de {start:%d/%m/%Y %H:%M}"
        return f"Ate {end:%d/%m/%Y %H:%M}"

    def _display_number(self, value: float) -> str:
        if abs(value - round(value)) < 1e-9:
            return str(int(round(value)))
        return f"{value:.2f}".rstrip("0").rstrip(".")

    def build_management_dashboard(self, filters: Dict[str, object] | None = None) -> dict:
        raw_filters = filters or {}
        obra_from = self._parse_date_flexible(raw_filters.get("obra_from", ""))
        obra_to = self._parse_date_flexible(raw_filters.get("obra_to", ""))
        processed_from = self._parse_date_flexible(raw_filters.get("processed_from", ""))
        processed_to = self._parse_date_flexible(raw_filters.get("processed_to", ""))
        contract_filter = normalizar_texto(str(raw_filters.get("contrato", "") or "").strip())
        nucleo_filter = normalizar_texto(str(raw_filters.get("nucleo", "") or "").strip())
        equipe_filter = normalizar_texto(str(raw_filters.get("equipe", "") or "").strip())
        top_n = self._safe_int(raw_filters.get("top_n", 10), default=10, min_value=3, max_value=50)
        history_limit = self._safe_int(raw_filters.get("history_limit", 1000), default=1000, min_value=20, max_value=10000)

        history_rows = self.read_history(limit=history_limit)
        candidate_runs: List[dict] = []
        for row in history_rows:
            if str(row.get("status_level", "")).strip().lower() != "sucesso":
                continue
            output_dir_raw = str(row.get("output_dir", "") or "").strip()
            if not output_dir_raw:
                continue
            output_dir = Path(output_dir_raw)
            if not output_dir.exists():
                continue

            processed_dt = self._parse_history_datetime(row.get("processed_at", ""))
            obra_dt = self._parse_date_flexible(row.get("obra_data", ""))

            if processed_from and (not processed_dt or processed_dt.date() < processed_from):
                continue
            if processed_to and (not processed_dt or processed_dt.date() > processed_to):
                continue
            if obra_from and (not obra_dt or obra_dt < obra_from):
                continue
            if obra_to and (not obra_dt or obra_dt > obra_to):
                continue
            if contract_filter:
                contract_probe = " ".join(
                    [
                        str(row.get("contract_label", "") or "").strip(),
                        str(row.get("contract_id", "") or "").strip(),
                    ]
                )
                if not self._match_filter_text(contract_probe, contract_filter):
                    continue

            candidate_runs.append(
                {
                    **row,
                    "_output_dir_path": output_dir,
                    "_processed_dt": processed_dt,
                    "_obra_dt": obra_dt,
                }
            )

        nucleo_stats: Dict[str, dict] = {}
        equipe_stats: Dict[str, dict] = {}
        service_stats: Dict[str, dict] = {}
        occurrence_stats: Dict[str, dict] = {}

        total_execucao = 0
        total_volume = 0.0
        total_frentes = 0
        frentes_com_producao = 0
        frentes_sem_producao = 0
        total_ocorrencias = 0

        nucleos_ativos = set()
        equipes_ativas = set()
        servicos_distintos = set()
        runs_com_dados = set()
        obra_dates_detected: List[date] = []
        processed_dates_detected: List[datetime] = []

        def _ensure_nucleo(nome: str) -> dict:
            key = normalizar_texto(nome) or "__sem_nucleo__"
            item = nucleo_stats.get(key)
            if item:
                return item
            item = {
                "nucleo": nome or "(sem nucleo)",
                "registros_execucao": 0,
                "volume_total": 0.0,
                "frentes": 0,
                "ocorrencias": 0,
                "equipes": set(),
                "servicos": set(),
                "runs": set(),
            }
            nucleo_stats[key] = item
            return item

        def _ensure_equipe(nome: str) -> dict:
            key = normalizar_texto(nome) or "__sem_equipe__"
            item = equipe_stats.get(key)
            if item:
                return item
            item = {
                "equipe": nome or "(sem equipe)",
                "registros_execucao": 0,
                "volume_total": 0.0,
                "frentes": 0,
                "ocorrencias": 0,
                "nucleos": set(),
                "servicos": set(),
                "runs": set(),
            }
            equipe_stats[key] = item
            return item

        for run in candidate_runs:
            output_dir = run["_output_dir_path"]
            run_key = str(output_dir.resolve())
            processed_dt = run.get("_processed_dt")
            run_had_rows = False
            run_fallback_nucleo = ""

            exec_rows = self._read_csv_rows(output_dir / "execucao.csv")
            frentes_rows = self._read_csv_rows(output_dir / "frentes.csv")
            ocorr_rows = self._read_csv_rows(output_dir / "ocorrencias.csv")

            for row in exec_rows:
                nucleo_raw = str(row.get("nucleo", "") or "").strip()
                nucleo = self._canonicalize_nucleo_for_aggregation(
                    nucleo_raw, fallback_nucleo=run_fallback_nucleo
                )
                if nucleo and not run_fallback_nucleo:
                    run_fallback_nucleo = nucleo

                eq_single = self._canonicalize_equipe_for_aggregation(
                    row.get("equipe", "")
                )
                equipes = [eq_single] if eq_single else []

                if not (
                    self._match_filter_text(nucleo_raw, nucleo_filter)
                    or self._match_filter_text(nucleo, nucleo_filter)
                ):
                    continue
                if equipe_filter and not any(
                    self._match_filter_text(eq, equipe_filter) for eq in equipes
                ):
                    continue

                quantidade = self._parse_quantity(row.get("quantidade", ""))
                servico = (
                    str(row.get("servico_oficial", "") or "").strip()
                    or str(row.get("item_normalizado", "") or "").strip()
                    or str(row.get("servico_bruto", "") or "").strip()
                    or str(row.get("item_original", "") or "").strip()
                    or "servico_nao_informado"
                )
                categoria = (
                    str(row.get("categoria_item", "") or "").strip()
                    or str(row.get("categoria", "") or "").strip()
                    or "-"
                )
                unidade = str(row.get("unidade", "") or "").strip()
                data_ref = self._parse_date_flexible(row.get("data_referencia", "")) or run.get("_obra_dt")

                total_execucao += 1
                total_volume += quantidade
                run_had_rows = True

                if data_ref:
                    obra_dates_detected.append(data_ref)
                if processed_dt:
                    processed_dates_detected.append(processed_dt)

                if nucleo:
                    nucleos_ativos.add(nucleo)
                for equipe in equipes:
                    equipes_ativas.add(equipe)
                servicos_distintos.add(servico)

                nuc = _ensure_nucleo(nucleo)
                nuc["registros_execucao"] += 1
                nuc["volume_total"] += quantidade
                nuc["runs"].add(run_key)
                for equipe in equipes:
                    nuc["equipes"].add(equipe)
                nuc["servicos"].add(servico)

                if equipes:
                    for equipe in equipes:
                        eqp = _ensure_equipe(equipe)
                        eqp["registros_execucao"] += 1
                        eqp["volume_total"] += quantidade
                        eqp["runs"].add(run_key)
                        if nucleo:
                            eqp["nucleos"].add(nucleo)
                        eqp["servicos"].add(servico)
                else:
                    eqp = _ensure_equipe("")
                    eqp["registros_execucao"] += 1
                    eqp["volume_total"] += quantidade
                    eqp["runs"].add(run_key)
                    if nucleo:
                        eqp["nucleos"].add(nucleo)
                    eqp["servicos"].add(servico)

                service_key = normalizar_texto(servico) or "__servico__"
                service_item = service_stats.get(service_key)
                if not service_item:
                    service_item = {
                        "servico": servico,
                        "categoria": categoria,
                        "registros": 0,
                        "volume_total": 0.0,
                        "unidades": set(),
                        "nucleos": set(),
                        "equipes": set(),
                    }
                    service_stats[service_key] = service_item
                service_item["registros"] += 1
                service_item["volume_total"] += quantidade
                if unidade:
                    service_item["unidades"].add(unidade)
                if nucleo:
                    service_item["nucleos"].add(nucleo)
                for equipe in equipes:
                    service_item["equipes"].add(equipe)

            for row in frentes_rows:
                nucleo_raw = str(row.get("nucleo", "") or "").strip()
                nucleo = self._canonicalize_nucleo_for_aggregation(
                    nucleo_raw, fallback_nucleo=run_fallback_nucleo
                )
                if nucleo and not run_fallback_nucleo:
                    run_fallback_nucleo = nucleo

                eq_single = self._canonicalize_equipe_for_aggregation(
                    row.get("equipe", "")
                )
                equipes = [eq_single] if eq_single else []
                if not (
                    self._match_filter_text(nucleo_raw, nucleo_filter)
                    or self._match_filter_text(nucleo, nucleo_filter)
                ):
                    continue
                if equipe_filter and not any(
                    self._match_filter_text(eq, equipe_filter) for eq in equipes
                ):
                    continue

                status_frente = normalizar_texto(str(row.get("status_frente", "") or ""))
                data_ref = self._parse_date_flexible(row.get("data_referencia", "")) or run.get("_obra_dt")

                total_frentes += 1
                run_had_rows = True
                if data_ref:
                    obra_dates_detected.append(data_ref)
                if processed_dt:
                    processed_dates_detected.append(processed_dt)

                if "sem_producao" in status_frente or "sem producao" in status_frente:
                    frentes_sem_producao += 1
                else:
                    frentes_com_producao += 1

                if nucleo:
                    nucleos_ativos.add(nucleo)
                    _ensure_nucleo(nucleo)["frentes"] += 1
                    _ensure_nucleo(nucleo)["runs"].add(run_key)
                for equipe in equipes:
                    equipes_ativas.add(equipe)
                    _ensure_equipe(equipe)["frentes"] += 1
                    _ensure_equipe(equipe)["runs"].add(run_key)

            for row in ocorr_rows:
                nucleo_raw = str(row.get("nucleo", "") or "").strip()
                nucleo = self._canonicalize_nucleo_for_aggregation(
                    nucleo_raw, fallback_nucleo=run_fallback_nucleo
                )
                if nucleo and not run_fallback_nucleo:
                    run_fallback_nucleo = nucleo

                eq_single = self._canonicalize_equipe_for_aggregation(
                    row.get("equipe", "")
                )
                equipes = [eq_single] if eq_single else []
                if not (
                    self._match_filter_text(nucleo_raw, nucleo_filter)
                    or self._match_filter_text(nucleo, nucleo_filter)
                ):
                    continue
                if equipe_filter and not any(
                    self._match_filter_text(eq, equipe_filter) for eq in equipes
                ):
                    continue

                tipo = str(row.get("tipo_ocorrencia", "") or "").strip() or "nao_informado"
                data_ref = self._parse_date_flexible(row.get("data_referencia", "")) or run.get("_obra_dt")

                total_ocorrencias += 1
                run_had_rows = True
                if data_ref:
                    obra_dates_detected.append(data_ref)
                if processed_dt:
                    processed_dates_detected.append(processed_dt)

                if nucleo:
                    nucleos_ativos.add(nucleo)
                    _ensure_nucleo(nucleo)["ocorrencias"] += 1
                    _ensure_nucleo(nucleo)["runs"].add(run_key)
                for equipe in equipes:
                    equipes_ativas.add(equipe)
                    _ensure_equipe(equipe)["ocorrencias"] += 1
                    _ensure_equipe(equipe)["runs"].add(run_key)

                occ_key = normalizar_texto(tipo) or "__ocorrencia__"
                occ_item = occurrence_stats.get(occ_key)
                if not occ_item:
                    occ_item = {
                        "tipo_ocorrencia": tipo,
                        "ocorrencias": 0,
                        "nucleos": set(),
                        "equipes": set(),
                    }
                    occurrence_stats[occ_key] = occ_item
                occ_item["ocorrencias"] += 1
                if nucleo:
                    occ_item["nucleos"].add(nucleo)
                for equipe in equipes:
                    occ_item["equipes"].add(equipe)

            if run_had_rows:
                runs_com_dados.add(run_key)

        by_nucleo = []
        for row in nucleo_stats.values():
            by_nucleo.append(
                {
                    "nucleo": row["nucleo"],
                    "registros_execucao": row["registros_execucao"],
                    "volume_total": row["volume_total"],
                    "volume_total_fmt": self._display_number(row["volume_total"]),
                    "frentes": row["frentes"],
                    "ocorrencias": row["ocorrencias"],
                    "equipes_ativas": len(row["equipes"]),
                    "servicos_distintos": len(row["servicos"]),
                    "runs": len(row["runs"]),
                }
            )
        by_nucleo.sort(
            key=lambda r: (float(r["volume_total"]), int(r["registros_execucao"]), int(r["frentes"])),
            reverse=True,
        )

        by_equipe = []
        for row in equipe_stats.values():
            by_equipe.append(
                {
                    "equipe": row["equipe"],
                    "registros_execucao": row["registros_execucao"],
                    "volume_total": row["volume_total"],
                    "volume_total_fmt": self._display_number(row["volume_total"]),
                    "frentes": row["frentes"],
                    "ocorrencias": row["ocorrencias"],
                    "nucleos_ativos": len(row["nucleos"]),
                    "servicos_distintos": len(row["servicos"]),
                    "runs": len(row["runs"]),
                }
            )
        by_equipe.sort(
            key=lambda r: (float(r["volume_total"]), int(r["registros_execucao"]), int(r["frentes"])),
            reverse=True,
        )

        ranking_servicos = []
        for row in service_stats.values():
            ranking_servicos.append(
                {
                    "servico": row["servico"],
                    "categoria": row["categoria"],
                    "registros": row["registros"],
                    "volume_total": row["volume_total"],
                    "volume_total_fmt": self._display_number(row["volume_total"]),
                    "unidades": ", ".join(sorted(row["unidades"])) if row["unidades"] else "-",
                    "nucleos_ativos": len(row["nucleos"]),
                    "equipes_ativas": len(row["equipes"]),
                }
            )
        ranking_servicos.sort(
            key=lambda r: (float(r["volume_total"]), int(r["registros"])),
            reverse=True,
        )

        ranking_ocorrencias = []
        for row in occurrence_stats.values():
            ranking_ocorrencias.append(
                {
                    "tipo_ocorrencia": row["tipo_ocorrencia"],
                    "ocorrencias": row["ocorrencias"],
                    "nucleos_ativos": len(row["nucleos"]),
                    "equipes_ativas": len(row["equipes"]),
                }
            )
        ranking_ocorrencias.sort(key=lambda r: int(r["ocorrencias"]), reverse=True)

        obra_start = min(obra_dates_detected) if obra_dates_detected else None
        obra_end = max(obra_dates_detected) if obra_dates_detected else None
        processed_start = min(processed_dates_detected) if processed_dates_detected else None
        processed_end = max(processed_dates_detected) if processed_dates_detected else None

        media_volume_por_run = (total_volume / len(runs_com_dados)) if runs_com_dados else 0.0
        media_volume_por_item = (total_volume / total_execucao) if total_execucao else 0.0
        itens_por_frente = (total_execucao / total_frentes) if total_frentes else 0.0
        ocorrencias_por_frente = (total_ocorrencias / total_frentes) if total_frentes else 0.0

        top_nucleo = by_nucleo[0] if by_nucleo else None
        top_equipe = by_equipe[0] if by_equipe else None
        top_servico = ranking_servicos[0] if ranking_servicos else None
        top_ocorrencia = ranking_ocorrencias[0] if ranking_ocorrencias else None

        return {
            "filters_applied": {
                "obra_from": str(raw_filters.get("obra_from", "") or ""),
                "obra_to": str(raw_filters.get("obra_to", "") or ""),
                "processed_from": str(raw_filters.get("processed_from", "") or ""),
                "processed_to": str(raw_filters.get("processed_to", "") or ""),
                "nucleo": str(raw_filters.get("nucleo", "") or ""),
                "equipe": str(raw_filters.get("equipe", "") or ""),
                "top_n": top_n,
                "history_limit": history_limit,
            },
            "consolidado_periodo": {
                "runs_consideradas": len(candidate_runs),
                "runs_com_dados": len(runs_com_dados),
                "periodo_obra": self._format_date_range(obra_start, obra_end),
                "periodo_processamento": self._format_datetime_range(processed_start, processed_end),
                "execucao_registros": total_execucao,
                "frentes": total_frentes,
                "frentes_com_producao": frentes_com_producao,
                "frentes_sem_producao": frentes_sem_producao,
                "ocorrencias": total_ocorrencias,
                "volume_total": total_volume,
                "volume_total_fmt": self._display_number(total_volume),
                "nucleos_ativos": len(nucleos_ativos),
                "equipes_ativas": len(equipes_ativas),
                "servicos_distintos": len(servicos_distintos),
            },
            "indicadores_produtividade": {
                "media_volume_por_run": media_volume_por_run,
                "media_volume_por_run_fmt": self._display_number(media_volume_por_run),
                "media_volume_por_item": media_volume_por_item,
                "media_volume_por_item_fmt": self._display_number(media_volume_por_item),
                "itens_por_frente": itens_por_frente,
                "itens_por_frente_fmt": self._display_number(itens_por_frente),
                "ocorrencias_por_frente": ocorrencias_por_frente,
                "ocorrencias_por_frente_fmt": self._display_number(ocorrencias_por_frente),
            },
            "visao_resumida": {
                "maior_nucleo_volume": top_nucleo,
                "maior_equipe_volume": top_equipe,
                "servico_lider": top_servico,
                "ocorrencia_lider": top_ocorrencia,
            },
            "indicadores_por_nucleo": by_nucleo[:top_n],
            "indicadores_por_equipe": by_equipe[:top_n],
            "ranking_servicos": ranking_servicos[:top_n],
            "ranking_ocorrencias": ranking_ocorrencias[:top_n],
            "has_data": bool(runs_com_dados),
        }
    def build_management_layer(self, filters: Dict[str, object] | None = None) -> dict:
        raw_filters = filters or {}
        top_n = self._safe_int(raw_filters.get("top_n", 10), default=10, min_value=3, max_value=50)
        history_limit = self._safe_int(raw_filters.get("history_limit", 3000), default=3000, min_value=20, max_value=10000)

        status_filter = str(raw_filters.get("status", "") or "").strip().lower()
        if status_filter not in {"", "sucesso", "erro"}:
            status_filter = ""

        alertas_filter = str(raw_filters.get("alertas", "") or "").strip().lower()
        if alertas_filter not in {"", "com_alerta", "sem_alerta"}:
            alertas_filter = ""

        base = self.build_management_dashboard(
            {
                "obra_from": str(raw_filters.get("obra_from", "") or ""),
                "obra_to": str(raw_filters.get("obra_to", "") or ""),
                "processed_from": str(raw_filters.get("processed_from", "") or ""),
                "processed_to": str(raw_filters.get("processed_to", "") or ""),
                "nucleo": str(raw_filters.get("nucleo", "") or ""),
                "equipe": str(raw_filters.get("equipe", "") or ""),
                "top_n": max(top_n, 50),
                "history_limit": history_limit,
            }
        )

        obra_from = self._parse_date_flexible(raw_filters.get("obra_from", ""))
        obra_to = self._parse_date_flexible(raw_filters.get("obra_to", ""))
        processed_from = self._parse_date_flexible(raw_filters.get("processed_from", ""))
        processed_to = self._parse_date_flexible(raw_filters.get("processed_to", ""))
        nucleo_filter = normalizar_texto(str(raw_filters.get("nucleo", "") or "").strip())
        municipio_filter = normalizar_texto(str(raw_filters.get("municipio", "") or "").strip())
        equipe_filter = normalizar_texto(str(raw_filters.get("equipe", "") or "").strip())

        history_rows = self.read_history(limit=history_limit)
        filtered_history: List[dict] = []

        for row in history_rows:
            processed_dt = self._parse_history_datetime(row.get("processed_at", ""))
            obra_dt = self._parse_date_flexible(row.get("obra_data", ""))

            if processed_from and (not processed_dt or processed_dt.date() < processed_from):
                continue
            if processed_to and (not processed_dt or processed_dt.date() > processed_to):
                continue
            if obra_from and (not obra_dt or obra_dt < obra_from):
                continue
            if obra_to and (not obra_dt or obra_dt > obra_to):
                continue

            row_status = str(row.get("status_level", "") or "").strip().lower()
            if status_filter and row_status != status_filter:
                continue

            has_alert = str(row.get("has_alerts", "") or "").strip().lower() == "sim"
            if alertas_filter == "com_alerta" and not has_alert:
                continue
            if alertas_filter == "sem_alerta" and has_alert:
                continue

            entities = self._collect_management_entities_for_row(row)

            if nucleo_filter and not any(
                self._match_filter_text(v, nucleo_filter)
                for v in entities.get("nucleos", [])
            ):
                continue
            if municipio_filter and not any(
                self._match_filter_text(v, municipio_filter)
                for v in entities.get("municipios", [])
            ):
                continue
            if equipe_filter and not any(
                self._match_filter_text(v, equipe_filter)
                for v in entities.get("equipes", [])
            ):
                continue

            filtered_history.append(
                {
                    **row,
                    "_processed_dt": processed_dt,
                    "_obra_dt": obra_dt,
                    "_entities": entities,
                }
            )

        total_processamentos = len(filtered_history)
        processamentos_sucesso = sum(1 for r in filtered_history if str(r.get("status_level", "") or "") == "sucesso")
        processamentos_erro = sum(1 for r in filtered_history if str(r.get("status_level", "") or "") == "erro")
        processamentos_com_alerta = sum(1 for r in filtered_history if str(r.get("has_alerts", "") or "") == "sim")

        ranking_nucleo_counter = Counter()
        ranking_equipe_counter = Counter()
        ranking_municipio_counter = Counter()
        ranking_nucleo_labels: Dict[str, str] = {}
        ranking_equipe_labels: Dict[str, str] = {}
        ranking_municipio_labels: Dict[str, str] = {}
        resumo_periodo: Dict[str, dict] = {}
        alert_counter = Counter()
        municipio_stats: Dict[str, dict] = {}

        def _count_label(counter: Counter, labels: Dict[str, str], value: object) -> None:
            label = str(value or "").strip()
            key = normalizar_texto(label)
            if not key:
                return
            counter[key] += 1
            if key not in labels:
                labels[key] = label

        def _ensure_period(periodo: str) -> dict:
            key = periodo or "(sem data)"
            item = resumo_periodo.get(key)
            if item:
                return item
            item = {
                "periodo": key,
                "processamentos": 0,
                "processamentos_alerta": 0,
                "sucesso": 0,
                "erro": 0,
                "nao_mapeados": 0,
            }
            resumo_periodo[key] = item
            return item

        def _ensure_municipio(municipio: str) -> dict:
            key = normalizar_texto(municipio) or "__sem_municipio__"
            item = municipio_stats.get(key)
            if item:
                return item
            item = {
                "municipio": municipio or "(sem municipio)",
                "processamentos": 0,
                "processamentos_alerta": 0,
                "sucesso": 0,
                "erro": 0,
                "nao_mapeados": 0,
                "nucleos": set(),
                "equipes": set(),
            }
            municipio_stats[key] = item
            return item

        for row in filtered_history:
            entities = dict(row.get("_entities", {}) or {})
            if not entities:
                entities = self._collect_management_entities_for_row(row)
            row_nucleos = list(entities.get("nucleos", []) or [])
            row_equipes = list(entities.get("equipes", []) or [])
            row_municipios = list(entities.get("municipios", []) or [])

            for nuc in row_nucleos:
                _count_label(ranking_nucleo_counter, ranking_nucleo_labels, nuc)
            for eqp in row_equipes:
                _count_label(ranking_equipe_counter, ranking_equipe_labels, eqp)
            for mun in row_municipios:
                _count_label(
                    ranking_municipio_counter, ranking_municipio_labels, mun
                )

            obra_data = str(row.get("obra_data", "") or "").strip() or "(sem data)"
            per = _ensure_period(obra_data)
            per["processamentos"] += 1
            if str(row.get("has_alerts", "") or "") == "sim":
                per["processamentos_alerta"] += 1
            if str(row.get("status_level", "") or "") == "sucesso":
                per["sucesso"] += 1
            else:
                per["erro"] += 1
            per["nao_mapeados"] += self._safe_int(row.get("nao_mapeados", 0), default=0, min_value=0, max_value=1000000)

            for mun in row_municipios:
                mun_item = _ensure_municipio(mun)
                mun_item["processamentos"] += 1
                if str(row.get("has_alerts", "") or "") == "sim":
                    mun_item["processamentos_alerta"] += 1
                if str(row.get("status_level", "") or "") == "sucesso":
                    mun_item["sucesso"] += 1
                else:
                    mun_item["erro"] += 1
                mun_item["nao_mapeados"] += self._safe_int(row.get("nao_mapeados", 0), default=0, min_value=0, max_value=1000000)
                for nuc in row_nucleos:
                    if nuc:
                        mun_item["nucleos"].add(nuc)
                for eqp in row_equipes:
                    if eqp:
                        mun_item["equipes"].add(eqp)

            for alert in row.get("alert_list", []) or []:
                alert_txt = str(alert or "").strip()
                if alert_txt:
                    alert_counter[alert_txt] += 1

        volume_categoria_map: Dict[str, dict] = {}
        total_mapeados = 0
        total_nao_mapeados = 0

        for row in base.get("ranking_servicos", []):
            categoria = str(row.get("categoria", "") or "").strip() or "(sem categoria)"
            registros = self._safe_int(row.get("registros", 0), default=0, min_value=0, max_value=1000000)
            volume = float(row.get("volume_total", 0.0) or 0.0)

            categoria_norm = normalizar_texto(categoria)
            servico_norm = normalizar_texto(str(row.get("servico", "") or ""))
            if categoria_norm == "servico_nao_mapeado" or servico_norm == "servico_nao_mapeado":
                total_nao_mapeados += registros
            else:
                total_mapeados += registros

            item = volume_categoria_map.get(categoria)
            if not item:
                item = {
                    "categoria": categoria,
                    "volume_total": 0.0,
                    "registros": 0,
                }
                volume_categoria_map[categoria] = item
            item["volume_total"] += volume
            item["registros"] += registros

        volume_por_categoria = []
        volume_total_base = float(base.get("consolidado_periodo", {}).get("volume_total", 0.0) or 0.0)
        total_execucoes_base = int(base.get("consolidado_periodo", {}).get("execucao_registros", 0) or 0)
        for item in volume_categoria_map.values():
            pct = (100.0 * item["volume_total"] / volume_total_base) if volume_total_base > 0 else 0.0
            volume_por_categoria.append(
                {
                    "categoria": item["categoria"],
                    "registros": item["registros"],
                    "volume_total": item["volume_total"],
                    "volume_total_fmt": self._display_number(item["volume_total"]),
                    "percentual_volume": pct,
                    "percentual_volume_fmt": self._display_number(pct),
                }
            )
        volume_por_categoria.sort(key=lambda r: float(r.get("volume_total", 0.0)), reverse=True)

        ranking_nao_mapeados = []
        nm_counter = Counter()
        nm_context = {}

        for row in filtered_history:
            if str(row.get("status_level", "") or "") != "sucesso":
                continue
            output_dir_raw = str(row.get("output_dir", "") or "").strip()
            if not output_dir_raw:
                continue
            output_dir = Path(output_dir_raw)
            if not output_dir.exists():
                continue

            nm_rows = self._read_csv_rows(output_dir / "servico_nao_mapeado.csv")
            for nm in nm_rows:
                nm_nucleo = self._canonicalize_nucleo_for_aggregation(nm.get("nucleo", ""))
                nm_equipe = self._canonicalize_equipe_for_aggregation(nm.get("equipe", ""))
                nm_municipio = self._canonicalize_municipio_for_aggregation(
                    nm.get("municipio", "")
                )

                if nucleo_filter and not self._match_filter_text(nm_nucleo, nucleo_filter):
                    continue
                if equipe_filter and not self._match_filter_text(nm_equipe, equipe_filter):
                    continue
                if municipio_filter and not self._match_filter_text(nm_municipio, municipio_filter):
                    continue

                termo = (
                    str(nm.get("servico_normalizado", "") or "").strip()
                    or str(nm.get("servico_bruto", "") or "").strip()
                    or "(sem termo)"
                )
                termo_key = normalizar_texto(termo) or termo
                nm_counter[termo_key] += 1
                nm_context[termo_key] = {
                    "termo": termo,
                    "nucleo": nm_nucleo,
                    "municipio": nm_municipio,
                    "equipe": nm_equipe,
                    "regra": str(nm.get("regra_disparada", "") or "").strip(),
                }

        for termo_key, qtd in nm_counter.most_common(top_n):
            ctx = nm_context.get(termo_key, {})
            ranking_nao_mapeados.append(
                {
                    "termo": ctx.get("termo", termo_key),
                    "ocorrencias": qtd,
                    "ultimo_nucleo": ctx.get("nucleo", ""),
                    "ultimo_municipio": ctx.get("municipio", ""),
                    "ultimo_equipe": ctx.get("equipe", ""),
                    "regra": ctx.get("regra", ""),
                }
            )

        total_itens_mapeamento = total_mapeados + total_nao_mapeados
        percentual_mapeado = (100.0 * total_mapeados / total_itens_mapeamento) if total_itens_mapeamento else 100.0
        percentual_nao_mapeado = 100.0 - percentual_mapeado if total_itens_mapeamento else 0.0

        ranking_nucleos_processamentos = [
            {
                "nucleo": ranking_nucleo_labels.get(k, k),
                "processamentos": v,
            }
            for k, v in ranking_nucleo_counter.most_common(top_n)
        ]
        ranking_equipes_processamentos = [
            {
                "equipe": ranking_equipe_labels.get(k, k),
                "processamentos": v,
            }
            for k, v in ranking_equipe_counter.most_common(top_n)
        ]
        ranking_nucleos_volume = [
            {
                "nucleo": str(row.get("nucleo", "") or "").strip() or "-",
                "volume_total": float(row.get("volume_total", 0.0) or 0.0),
                "volume_total_fmt": str(row.get("volume_total_fmt", "") or self._display_number(float(row.get("volume_total", 0.0) or 0.0))),
            }
            for row in list(base.get("indicadores_por_nucleo", []) or [])[:top_n]
        ]
        ranking_equipes_volume = [
            {
                "equipe": str(row.get("equipe", "") or "").strip() or "-",
                "volume_total": float(row.get("volume_total", 0.0) or 0.0),
                "volume_total_fmt": str(row.get("volume_total_fmt", "") or self._display_number(float(row.get("volume_total", 0.0) or 0.0))),
            }
            for row in list(base.get("indicadores_por_equipe", []) or [])[:top_n]
        ]
        ranking_municipios_processamentos = [
            {
                "municipio": ranking_municipio_labels.get(k, k),
                "processamentos": v,
            }
            for k, v in ranking_municipio_counter.most_common(top_n)
        ]
        alertas_recorrentes = [
            {"alerta": k, "ocorrencias": v}
            for k, v in alert_counter.most_common(top_n)
        ]

        indicadores_por_municipio = []
        for item in municipio_stats.values():
            indicadores_por_municipio.append(
                {
                    "municipio": item["municipio"],
                    "processamentos": item["processamentos"],
                    "processamentos_alerta": item["processamentos_alerta"],
                    "sucesso": item["sucesso"],
                    "erro": item["erro"],
                    "nao_mapeados": item["nao_mapeados"],
                    "nucleos_ativos": len(item["nucleos"]),
                    "equipes_ativas": len(item["equipes"]),
                }
            )
        indicadores_por_municipio.sort(key=lambda r: int(r["processamentos"]), reverse=True)

        resumo_por_periodo = list(resumo_periodo.values())
        resumo_por_periodo.sort(
            key=lambda r: self._parse_date_flexible(r.get("periodo", "")) or date.min,
            reverse=True,
        )

        def _window_stats(rows: List[dict], start_dt: date, end_dt: date) -> dict:
            subset = [r for r in rows if r.get("_obra_dt") and start_dt <= r.get("_obra_dt") <= end_dt]
            return {
                "processamentos": len(subset),
                "com_alerta": sum(1 for r in subset if str(r.get("has_alerts", "") or "") == "sim"),
                "nao_mapeados": sum(self._safe_int(r.get("nao_mapeados", 0), default=0, min_value=0, max_value=1000000) for r in subset),
                "sucesso": sum(1 for r in subset if str(r.get("status_level", "") or "") == "sucesso"),
                "erro": sum(1 for r in subset if str(r.get("status_level", "") or "") == "erro"),
            }

        comparativos_basicos = {
            "has_comparativo": False,
            "label_atual": "-",
            "label_anterior": "-",
            "atual": {},
            "anterior": {},
            "delta": {},
        }

        obra_dates_filtered = [r.get("_obra_dt") for r in filtered_history if r.get("_obra_dt")]
        curr_start = None
        curr_end = None

        if obra_from or obra_to:
            if obra_dates_filtered:
                curr_start = obra_from or min(obra_dates_filtered)
                curr_end = obra_to or max(obra_dates_filtered)
        elif obra_dates_filtered:
            curr_end = max(obra_dates_filtered)
            curr_start = curr_end - timedelta(days=6)

        if curr_start and curr_end and curr_end >= curr_start:
            span_days = (curr_end - curr_start).days + 1
            prev_end = curr_start - timedelta(days=1)
            prev_start = prev_end - timedelta(days=span_days - 1)

            atual_stats = _window_stats(filtered_history, curr_start, curr_end)
            anterior_stats = _window_stats(filtered_history, prev_start, prev_end)

            def _delta_pack(current_value: int, previous_value: int) -> dict:
                absolute = int(current_value) - int(previous_value)
                pct = (100.0 * absolute / previous_value) if previous_value else 0.0
                return {
                    "absolute": absolute,
                    "percent": pct,
                    "percent_fmt": self._display_number(pct),
                }

            comparativos_basicos = {
                "has_comparativo": True,
                "label_atual": self._format_date_range(curr_start, curr_end),
                "label_anterior": self._format_date_range(prev_start, prev_end),
                "atual": atual_stats,
                "anterior": anterior_stats,
                "delta": {
                    "processamentos": _delta_pack(atual_stats["processamentos"], anterior_stats["processamentos"]),
                    "com_alerta": _delta_pack(atual_stats["com_alerta"], anterior_stats["com_alerta"]),
                    "nao_mapeados": _delta_pack(atual_stats["nao_mapeados"], anterior_stats["nao_mapeados"]),
                },
            }

        top_nucleo = (
            ranking_nucleos_volume[0]["nucleo"]
            if ranking_nucleos_volume
            else (ranking_nucleos_processamentos[0]["nucleo"] if ranking_nucleos_processamentos else "-")
        )
        top_equipe = (
            ranking_equipes_volume[0]["equipe"]
            if ranking_equipes_volume
            else (ranking_equipes_processamentos[0]["equipe"] if ranking_equipes_processamentos else "-")
        )
        top_alerta = alertas_recorrentes[0]["alerta"] if alertas_recorrentes else "-"
        top_nao_mapeado = ranking_nao_mapeados[0]["termo"] if ranking_nao_mapeados else "-"

        resumo_executivo = {
            "periodo": base.get("consolidado_periodo", {}).get("periodo_obra", "-"),
            "linhas": [
                f"Quantitativo consolidado do período: {self._display_number(volume_total_base)}, com {total_execucoes_base} execucao(oes) e {total_processamentos} processamento(s).",
                f"Nucleo mais recorrente: {top_nucleo}. Equipe mais recorrente: {top_equipe}.",
                f"Mapeamento: {self._display_number(percentual_mapeado)}% mapeado e {self._display_number(percentual_nao_mapeado)}% nao mapeado.",
                f"Alerta mais recorrente: {top_alerta}. Nao mapeado mais recorrente: {top_nao_mapeado}.",
            ],
        }

        base["filters_applied"]["municipio"] = str(raw_filters.get("municipio", "") or "")
        base["filters_applied"]["status"] = status_filter
        base["filters_applied"]["alertas"] = alertas_filter

        base["kpis_principais"] = {
            "total_processamentos": total_processamentos,
            "total_execucoes": int(base.get("consolidado_periodo", {}).get("execucao_registros", 0) or 0),
            "total_volume": float(base.get("consolidado_periodo", {}).get("volume_total", 0.0) or 0.0),
            "total_volume_fmt": str(base.get("consolidado_periodo", {}).get("volume_total_fmt", "0") or "0"),
            "total_frentes": int(base.get("consolidado_periodo", {}).get("frentes", 0) or 0),
            "total_ocorrencias": int(base.get("consolidado_periodo", {}).get("ocorrencias", 0) or 0),
            "total_nucleos": len(ranking_nucleo_counter) if ranking_nucleo_counter else int(base.get("consolidado_periodo", {}).get("nucleos_ativos", 0) or 0),
            "total_equipes": len(ranking_equipe_counter) if ranking_equipe_counter else int(base.get("consolidado_periodo", {}).get("equipes_ativas", 0) or 0),
            "total_municipios": len(ranking_municipio_counter),
            "total_mapeados": total_mapeados,
            "total_nao_mapeados": total_nao_mapeados,
            "percentual_mapeado": percentual_mapeado,
            "percentual_mapeado_fmt": self._display_number(percentual_mapeado),
            "percentual_nao_mapeado": percentual_nao_mapeado,
            "percentual_nao_mapeado_fmt": self._display_number(percentual_nao_mapeado),
            "processamentos_com_alerta": processamentos_com_alerta,
            "processamentos_sem_alerta": max(total_processamentos - processamentos_com_alerta, 0),
            "processamentos_sucesso": processamentos_sucesso,
            "processamentos_erro": processamentos_erro,
        }

        base["volume_por_categoria"] = volume_por_categoria[:top_n]
        base["indicadores_por_municipio"] = indicadores_por_municipio[:top_n]
        base["resumo_por_periodo"] = resumo_por_periodo[: max(top_n, 10)]
        base["ranking_nucleos_processamentos"] = ranking_nucleos_processamentos
        base["ranking_equipes_processamentos"] = ranking_equipes_processamentos
        base["ranking_nucleos_volume"] = ranking_nucleos_volume
        base["ranking_equipes_volume"] = ranking_equipes_volume
        base["ranking_municipios_processamentos"] = ranking_municipios_processamentos
        base["ranking_nao_mapeados"] = ranking_nao_mapeados[:top_n]
        base["alertas_recorrentes"] = alertas_recorrentes
        base["comparativos_basicos"] = comparativos_basicos
        base["resumo_executivo"] = resumo_executivo
        chart_limit = min(max(top_n, 5), 10)
        base["cards_executivos"] = [
            {
                "title": "Quantitativo total",
                "value": base["kpis_principais"]["total_volume_fmt"],
                "subtitle": "volume consolidado",
                "tone": "info",
            },
            {
                "title": "Execuções",
                "value": base["kpis_principais"]["total_execucoes"],
                "subtitle": "registros de apoio",
                "tone": "",
            },
            {
                "title": "Frentes",
                "value": base["kpis_principais"]["total_frentes"],
                "subtitle": "frentes registradas",
                "tone": "",
            },
            {
                "title": "Ocorrências",
                "value": base["kpis_principais"]["total_ocorrencias"],
                "subtitle": "impactos registrados",
                "tone": "",
            },
            {
                "title": "Núcleos",
                "value": base["kpis_principais"]["total_nucleos"],
                "subtitle": "núcleos envolvidos",
                "tone": "",
            },
            {
                "title": "Municípios",
                "value": base["kpis_principais"]["total_municipios"],
                "subtitle": "municípios envolvidos",
                "tone": "",
            },
            {
                "title": "Equipes",
                "value": base["kpis_principais"]["total_equipes"],
                "subtitle": "equipes envolvidas",
                "tone": "",
            },
            {
                "title": "Com alerta",
                "value": base["kpis_principais"]["processamentos_com_alerta"],
                "subtitle": "processamentos com atenção",
                "tone": "warning",
            },
            {
                "title": "Mapeado",
                "value": f"{base['kpis_principais']['percentual_mapeado_fmt']}%",
                "subtitle": "itens classificados",
                "tone": "ok",
            },
            {
                "title": "Não mapeado",
                "value": f"{base['kpis_principais']['percentual_nao_mapeado_fmt']}%",
                "subtitle": "itens sem classificação",
                "tone": "critical" if float(base["kpis_principais"]["percentual_nao_mapeado"] or 0) > 0 else "",
            },
        ]
        base["graficos_executivos"] = {
            "periodo": self._build_dashboard_period_chart(base["resumo_por_periodo"]),
            "servicos": self._build_dashboard_bar_chart(
                base.get("ranking_servicos", []),
                label_field="servico",
                value_field="volume_total",
                limit=chart_limit,
                domain="servico",
                value_mode="number",
                unit_field="unidades",
            ),
            "ocorrencias": self._build_dashboard_bar_chart(
                base.get("ranking_ocorrencias", []),
                label_field="tipo_ocorrencia",
                value_field="ocorrencias",
                limit=chart_limit,
                domain="ocorrencia",
            ),
            "nucleos": self._build_dashboard_bar_chart(
                base.get("ranking_nucleos_volume", []),
                label_field="nucleo",
                value_field="volume_total",
                limit=chart_limit,
                value_mode="number",
            ),
            "equipes": self._build_dashboard_bar_chart(
                base.get("ranking_equipes_volume", []),
                label_field="equipe",
                value_field="volume_total",
                limit=chart_limit,
                value_mode="number",
            ),
            "categorias": self._build_dashboard_bar_chart(
                base.get("volume_por_categoria", []),
                label_field="categoria",
                value_field="volume_total",
                limit=chart_limit,
                domain="categoria",
                value_mode="number",
            ),
            "mapeamento": self._build_dashboard_mapping_chart(base["kpis_principais"]),
        }

        if status_filter == "erro":
            base["has_data"] = False
            base["consolidado_periodo"] = {
                "runs_consideradas": 0,
                "runs_com_dados": 0,
                "periodo_obra": "-",
                "periodo_processamento": "-",
                "execucao_registros": 0,
                "frentes": 0,
                "frentes_com_producao": 0,
                "frentes_sem_producao": 0,
                "ocorrencias": 0,
                "volume_total": 0.0,
                "volume_total_fmt": "0",
                "nucleos_ativos": 0,
                "equipes_ativas": 0,
                "servicos_distintos": 0,
            }
            base["indicadores_por_nucleo"] = []
            base["indicadores_por_equipe"] = []
            base["ranking_servicos"] = []
            base["ranking_ocorrencias"] = []
            base["volume_por_categoria"] = []
            base["cards_executivos"] = []
            base["graficos_executivos"] = {
                "periodo": {"items": [], "has_data": False, "max_value_fmt": "0"},
                "servicos": {"items": [], "has_data": False, "max_value_fmt": "0"},
                "ocorrencias": {"items": [], "has_data": False, "max_value_fmt": "0"},
                "nucleos": {"items": [], "has_data": False, "max_value_fmt": "0"},
                "equipes": {"items": [], "has_data": False, "max_value_fmt": "0"},
                "categorias": {"items": [], "has_data": False, "max_value_fmt": "0"},
                "mapeamento": self._build_dashboard_mapping_chart(base["kpis_principais"]),
            }

        return base

    def _explode_history_values(self, value: object) -> List[str]:
        text = str(value or "").replace("\r\n", "\n").replace("\r", "\n").strip()
        if not text:
            return []
        split_on_comma = False
        if text.lower().startswith("multiplos") and ":" in text:
            text = text.split(":", 1)[1]
            split_on_comma = True
        text = text.replace("|", "/")
        if split_on_comma:
            text = text.replace(",", "/")
        return self._split_multi_text(text)

    def _filter_history_for_institutional(self, raw_filters: Dict[str, object], history_limit: int) -> List[dict]:
        obra_from = self._parse_date_flexible(raw_filters.get("obra_from", ""))
        obra_to = self._parse_date_flexible(raw_filters.get("obra_to", ""))
        processed_from = self._parse_date_flexible(raw_filters.get("processed_from", ""))
        processed_to = self._parse_date_flexible(raw_filters.get("processed_to", ""))

        nucleo_filter = normalizar_texto(str(raw_filters.get("nucleo", "") or "").strip())
        municipio_filter = normalizar_texto(str(raw_filters.get("municipio", "") or "").strip())
        equipe_filter = normalizar_texto(str(raw_filters.get("equipe", "") or "").strip())

        status_filter = str(raw_filters.get("status", "") or "").strip().lower()
        if status_filter not in {"", "sucesso", "erro"}:
            status_filter = ""

        alertas_filter = str(raw_filters.get("alertas", "") or "").strip().lower()
        if alertas_filter not in {"", "com_alerta", "sem_alerta"}:
            alertas_filter = ""

        out: List[dict] = []
        for row in self.read_history(limit=history_limit):
            processed_dt = self._parse_history_datetime(row.get("processed_at", ""))
            obra_dt = self._parse_date_flexible(row.get("obra_data", ""))

            if processed_from and (not processed_dt or processed_dt.date() < processed_from):
                continue
            if processed_to and (not processed_dt or processed_dt.date() > processed_to):
                continue
            if obra_from and (not obra_dt or obra_dt < obra_from):
                continue
            if obra_to and (not obra_dt or obra_dt > obra_to):
                continue

            row_status = str(row.get("status_level", "") or "").strip().lower()
            if status_filter and row_status != status_filter:
                continue

            has_alert = str(row.get("has_alerts", "") or "").strip().lower() == "sim"
            if alertas_filter == "com_alerta" and not has_alert:
                continue
            if alertas_filter == "sem_alerta" and has_alert:
                continue

            entities = self._collect_management_entities_for_row(row)

            if nucleo_filter and not any(
                self._match_filter_text(v, nucleo_filter)
                for v in entities.get("nucleos", [])
            ):
                continue
            if municipio_filter and not any(
                self._match_filter_text(v, municipio_filter)
                for v in entities.get("municipios", [])
            ):
                continue
            if equipe_filter and not any(
                self._match_filter_text(v, equipe_filter)
                for v in entities.get("equipes", [])
            ):
                continue

            out.append(
                {
                    **row,
                    "_processed_dt": processed_dt,
                    "_obra_dt": obra_dt,
                    "_entities": entities,
                }
            )
        return out

    def _build_institutional_nucleo_analysis(
        self,
        filtered_history: List[dict],
        top_n: int,
        nucleo_filter: str = "",
        municipio_filter: str = "",
        equipe_filter: str = "",
        include_history_alerts: bool = False,
    ) -> List[dict]:
        nucleo_map: Dict[str, dict] = {}

        def _ensure_item(nucleo: str) -> dict:
            key = normalizar_texto(nucleo) or "__sem_nucleo__"
            item = nucleo_map.get(key)
            if item:
                return item
            item = {
                "nucleo": nucleo or "(sem nucleo)",
                "processamentos": 0,
                "sucesso": 0,
                "erro": 0,
                "processamentos_alerta": 0,
                "volume_total": 0.0,
                "municipio": "",
                "equipes": set(),
                "logradouros": set(),
                "service_counter": Counter(),
                "service_volume_map": defaultdict(float),
                "service_unit_map": {},
                "occurrence_counter": Counter(),
                "observation_counter": Counter(),
            }
            nucleo_map[key] = item
            return item

        def _apply_single_municipio(
            item: dict,
            nucleo: str,
            *,
            explicit_oficial: object = "",
            detected_values: List[str] | None = None,
        ) -> str:
            municipio = self._resolve_single_municipio_for_nucleo(
                nucleo,
                explicit_oficial=explicit_oficial,
                detected_values=detected_values or [],
            )
            if municipio:
                item["municipio"] = municipio
            return str(item.get("municipio", "") or "")

        for row in filtered_history:
            entities = dict(row.get("_entities", {}) or {})
            if not entities:
                entities = self._collect_management_entities_for_row(row)

            row_nucleos_raw = list(entities.get("nucleos", []) or [])
            if not row_nucleos_raw:
                row_nucleos_raw = ["(sem nucleo)"]
            row_nucleos_map: Dict[str, str] = {}
            fallback_nucleo = ""
            for nucleo_raw in row_nucleos_raw:
                nucleo = self._canonicalize_nucleo_for_aggregation(
                    nucleo_raw, fallback_nucleo=fallback_nucleo
                )
                if not nucleo:
                    continue
                key = normalizar_texto(nucleo)
                if not key:
                    continue
                if key not in row_nucleos_map:
                    row_nucleos_map[key] = nucleo
                if not fallback_nucleo:
                    fallback_nucleo = nucleo
            row_nucleos = list(row_nucleos_map.values()) or ["(sem núcleo)"]

            row_municipio = self._resolve_institutional_row_municipio(
                row,
                row_nucleos,
                entities,
            )

            row_equipes_map: Dict[str, str] = {}
            for equipe_raw in list(entities.get("equipes", []) or []):
                equipe = self._canonicalize_equipe_for_aggregation(equipe_raw)
                if not equipe:
                    continue
                key = normalizar_texto(equipe)
                if key and key not in row_equipes_map:
                    row_equipes_map[key] = equipe
            row_equipes = list(row_equipes_map.values())

            row_logradouros_map: Dict[str, str] = {}
            for logradouro_raw in self._explode_history_values(row.get("logradouro", "")):
                logradouro = " ".join(str(logradouro_raw or "").split()).strip()
                if not logradouro:
                    continue
                key = normalizar_texto(logradouro)
                if key and key not in row_logradouros_map:
                    row_logradouros_map[key] = logradouro
            row_logradouros = list(row_logradouros_map.values())
            row_alerts = list(row.get("alert_list", []) or []) if include_history_alerts else []

            has_alert = str(row.get("has_alerts", "") or "").strip().lower() == "sim"
            status_level = str(row.get("status_level", "") or "").strip().lower()

            for nucleo in row_nucleos:
                if nucleo_filter and not self._match_filter_text(nucleo, nucleo_filter):
                    continue
                item = _ensure_item(nucleo)
                item["processamentos"] += 1
                if status_level == "erro":
                    item["erro"] += 1
                else:
                    item["sucesso"] += 1
                if has_alert:
                    item["processamentos_alerta"] += 1

                _apply_single_municipio(
                    item,
                    nucleo,
                    explicit_oficial=row.get("municipio_oficial", ""),
                    detected_values=[row_municipio] if row_municipio else [],
                )
                for equipe in row_equipes:
                    if equipe:
                        item["equipes"].add(equipe)
                for logradouro in row_logradouros:
                    if logradouro:
                        item["logradouros"].add(logradouro)
                for alert in row_alerts:
                    alert_text = str(alert or "").strip()
                    if alert_text:
                        item["observation_counter"][alert_text] += 1

        for row in filtered_history:
            if str(row.get("status_level", "") or "").strip().lower() != "sucesso":
                continue
            output_dir_raw = str(row.get("output_dir", "") or "").strip()
            if not output_dir_raw:
                continue

            output_dir = Path(output_dir_raw)
            if not output_dir.exists():
                continue

            exec_rows = reconcile_rows_with_registry(
                self._read_csv_rows(output_dir / "execucao.csv"),
                self.nucleo_reference,
            )
            frentes_rows = reconcile_rows_with_registry(
                self._read_csv_rows(output_dir / "frentes.csv"),
                self.nucleo_reference,
            )
            occ_rows = reconcile_rows_with_registry(
                self._read_csv_rows(output_dir / "ocorrencias.csv"),
                self.nucleo_reference,
            )

            for exec_row in exec_rows:
                nucleo_raw = str(exec_row.get("nucleo", "") or "").strip()
                nucleo = self._canonicalize_nucleo_for_aggregation(nucleo_raw)
                if not nucleo:
                    continue
                if nucleo_filter and not (self._match_filter_text(nucleo_raw, nucleo_filter) or self._match_filter_text(nucleo, nucleo_filter)):
                    continue

                municipio = self._canonicalize_municipio_for_aggregation(
                    exec_row.get("municipio", "")
                )
                equipe = self._canonicalize_equipe_for_aggregation(
                    exec_row.get("equipe", "")
                )
                if municipio_filter and not self._match_filter_text(municipio, municipio_filter):
                    continue
                if equipe_filter and not self._match_filter_text(equipe, equipe_filter):
                    continue

                item = _ensure_item(nucleo)
                servico_raw = (
                    str(exec_row.get("servico_oficial", "") or "").strip()
                    or str(exec_row.get("item_normalizado", "") or "").strip()
                    or str(exec_row.get("servico_bruto", "") or "").strip()
                    or str(exec_row.get("item_original", "") or "").strip()
                    or "servico_nao_informado"
                )
                servico = self._institutional_label(servico_raw, domain="servico")
                quantidade = self._parse_quantity(exec_row.get("quantidade", ""))
                unidade = str(exec_row.get("unidade", "") or "").strip()
                item["service_counter"][servico] += 1
                item["service_volume_map"][servico] += quantidade
                if unidade and not item["service_unit_map"].get(servico):
                    item["service_unit_map"][servico] = unidade
                item["volume_total"] += quantidade
                _apply_single_municipio(
                    item,
                    nucleo,
                    explicit_oficial=exec_row.get("municipio_oficial", ""),
                    detected_values=[exec_row.get("municipio", "")],
                )
                if equipe:
                    item["equipes"].add(equipe)
                logradouro = str(exec_row.get("logradouro", "") or "").strip()
                if logradouro:
                    item["logradouros"].add(logradouro)

            for frente_row in frentes_rows:
                nucleo_raw = str(frente_row.get("nucleo", "") or "").strip()
                nucleo = self._canonicalize_nucleo_for_aggregation(nucleo_raw)
                if not nucleo:
                    continue
                if nucleo_filter and not (
                    self._match_filter_text(nucleo_raw, nucleo_filter)
                    or self._match_filter_text(nucleo, nucleo_filter)
                ):
                    continue

                municipio = self._canonicalize_municipio_for_aggregation(
                    frente_row.get("municipio", "")
                )
                equipe = self._canonicalize_equipe_for_aggregation(
                    frente_row.get("equipe", "")
                )
                if municipio_filter and not self._match_filter_text(municipio, municipio_filter):
                    continue
                if equipe_filter and not self._match_filter_text(equipe, equipe_filter):
                    continue

                item = _ensure_item(nucleo)
                _apply_single_municipio(
                    item,
                    nucleo,
                    explicit_oficial=frente_row.get("municipio_oficial", ""),
                    detected_values=[frente_row.get("municipio", "")],
                )
                if equipe:
                    item["equipes"].add(equipe)
                logradouro = str(frente_row.get("logradouro", "") or "").strip()
                if logradouro:
                    item["logradouros"].add(logradouro)

            for occ_row in occ_rows:
                nucleo_raw = str(occ_row.get("nucleo", "") or "").strip()
                nucleo = self._canonicalize_nucleo_for_aggregation(nucleo_raw)
                if not nucleo:
                    continue
                if nucleo_filter and not (self._match_filter_text(nucleo_raw, nucleo_filter) or self._match_filter_text(nucleo, nucleo_filter)):
                    continue

                municipio = self._canonicalize_municipio_for_aggregation(
                    occ_row.get("municipio", "")
                )
                equipe = self._canonicalize_equipe_for_aggregation(
                    occ_row.get("equipe", "")
                )
                if municipio_filter and not self._match_filter_text(municipio, municipio_filter):
                    continue
                if equipe_filter and not self._match_filter_text(equipe, equipe_filter):
                    continue

                item = _ensure_item(nucleo)
                tipo_raw = str(occ_row.get("tipo_ocorrencia", "") or "").strip() or "nao_informado"
                tipo = self._institutional_label(tipo_raw, domain="ocorrencia")
                item["occurrence_counter"][tipo] += 1

                descricao = self._normalize_institutional_note(occ_row.get("descricao", ""))
                if descricao and descricao != "Não informado":
                    item["observation_counter"][descricao] += 1
                _apply_single_municipio(
                    item,
                    nucleo,
                    explicit_oficial=occ_row.get("municipio_oficial", ""),
                    detected_values=[occ_row.get("municipio", "")],
                )
                if equipe:
                    item["equipes"].add(equipe)
                logradouro = str(occ_row.get("logradouro", "") or "").strip()
                if logradouro:
                    item["logradouros"].add(logradouro)

        analysis: List[dict] = []
        for item in nucleo_map.values():
            profile = self._get_nucleo_profile(item["nucleo"])
            municipio_oficial = self._canonicalize_municipio_for_aggregation(
                str((profile or {}).get("municipio", "") or item.get("municipio", "") or "")
            )
            equipes = sorted(item["equipes"], key=lambda x: normalizar_texto(x))
            logradouros = sorted(item["logradouros"], key=lambda x: normalizar_texto(x))

            service_volume_rows = sorted(
                item["service_volume_map"].items(),
                key=lambda kv: float(kv[1] or 0.0),
                reverse=True,
            )
            principais_servicos = []
            if service_volume_rows:
                for nome, volume in service_volume_rows[:5]:
                    ocorrencias = int(item["service_counter"].get(nome, 0) or 0)
                    unidade = str(item["service_unit_map"].get(nome, "") or "").strip()
                    volume_fmt = self._display_number(float(volume or 0.0))
                    unit_suffix = f" {unidade}" if unidade else ""
                    descricao = (
                        f"{nome}, com quantitativo consolidado de {volume_fmt}{unit_suffix} no período"
                    )
                    if ocorrencias > 0:
                        descricao += (
                            f" ({self._count_label(ocorrencias, 'registro', 'registros')})."
                        )
                    else:
                        descricao += "."
                    principais_servicos.append(
                        {
                            "nome": nome,
                            "ocorrencias": ocorrencias,
                            "volume_total": float(volume or 0.0),
                            "volume_total_fmt": volume_fmt,
                            "unidade": unidade,
                            "descricao_institucional": descricao,
                        }
                    )
            else:
                principais_servicos = [
                    {
                        "nome": nome,
                        "ocorrencias": qtd,
                        "volume_total": 0.0,
                        "volume_total_fmt": "0",
                        "unidade": "",
                        "descricao_institucional": f"{nome}, com {self._count_label(qtd, 'ocorrência', 'ocorrências')} no período.",
                    }
                    for nome, qtd in item["service_counter"].most_common(5)
                ]
            principais_ocorrencias = [
                {
                    "nome": nome,
                    "ocorrencias": qtd,
                    "descricao_institucional": f"{nome}, com {self._count_label(qtd, 'ocorrência', 'ocorrências')} no período.",
                }
                for nome, qtd in item["occurrence_counter"].most_common(5)
            ]
            observacoes_relevantes = []
            for nome, _ in item["observation_counter"].most_common(3):
                clean_obs = self._normalize_institutional_note(nome)
                if clean_obs and clean_obs != "Não informado":
                    observacoes_relevantes.append(clean_obs)

            top_servico = principais_servicos[0]["nome"] if principais_servicos else "-"
            top_ocorrencia = principais_ocorrencias[0]["nome"] if principais_ocorrencias else "-"
            municipio_texto = municipio_oficial or "-"
            logradouro_texto = self._summarize_values(logradouros, limit=3)
            if item["processamentos"] > 0:
                processos_texto = self._count_label(
                    item["processamentos"], "registro operacional", "registros operacionais"
                )
                volume_texto = self._display_number(float(item.get("volume_total", 0.0) or 0.0))
                if top_ocorrencia != "-":
                    observacao_analitica = (
                        f"No período, o núcleo consolidou quantitativo de {volume_texto}, "
                        f"com destaque para {top_servico.lower()} e recorrência de {top_ocorrencia.lower()} "
                        f"(apoio: {processos_texto})."
                    )
                else:
                    observacao_analitica = (
                        f"No período, o núcleo consolidou quantitativo de {volume_texto}, "
                        f"com destaque para {top_servico.lower()} e sem recorrência relevante de ocorrências "
                        f"(apoio: {processos_texto})."
                    )
            else:
                observacao_analitica = "Sem atividade consolidada no período filtrado."

            analysis.append(
                {
                    "nucleo": item["nucleo"],
                    "processamentos": item["processamentos"],
                    "volume_total": float(item.get("volume_total", 0.0) or 0.0),
                    "volume_total_fmt": self._display_number(float(item.get("volume_total", 0.0) or 0.0)),
                    "sucesso": item["sucesso"],
                    "erro": item["erro"],
                    "processamentos_alerta": item["processamentos_alerta"],
                    "municipio": municipio_texto,
                    "equipes": equipes,
                    "logradouro": logradouro_texto,
                    "logradouros": logradouros,
                    "principais_servicos": principais_servicos,
                    "principais_ocorrencias": principais_ocorrencias,
                    "observacoes_relevantes": observacoes_relevantes,
                    "observacao_analitica": observacao_analitica,
                }
            )

        analysis.sort(
            key=lambda row: (
                float(row.get("volume_total", 0.0) or 0.0),
                int(row.get("processamentos", 0)),
                int(row.get("processamentos_alerta", 0)),
                len(row.get("principais_servicos", [])),
            ),
            reverse=True,
        )
        return analysis[:top_n]

    def build_institutional_report(self, filters: Dict[str, object] | None = None) -> dict:
        raw_filters = filters or {}
        top_n = self._safe_int(raw_filters.get("top_n", 5), default=5, min_value=3, max_value=20)
        history_limit = self._safe_int(raw_filters.get("history_limit", 3000), default=3000, min_value=20, max_value=10000)

        status_filter = str(raw_filters.get("status", "") or "").strip().lower()
        if status_filter not in {"", "sucesso", "erro"}:
            status_filter = ""

        alertas_filter = str(raw_filters.get("alertas", "") or "").strip().lower()
        if alertas_filter not in {"", "com_alerta", "sem_alerta"}:
            alertas_filter = ""

        dashboard_filters = {
            "obra_from": str(raw_filters.get("obra_from", "") or ""),
            "obra_to": str(raw_filters.get("obra_to", "") or ""),
            "processed_from": str(raw_filters.get("processed_from", "") or ""),
            "processed_to": str(raw_filters.get("processed_to", "") or ""),
            "nucleo": str(raw_filters.get("nucleo", "") or ""),
            "municipio": str(raw_filters.get("municipio", "") or ""),
            "equipe": str(raw_filters.get("equipe", "") or ""),
            "status": status_filter,
            "alertas": alertas_filter,
            "top_n": max(top_n, 15),
            "history_limit": history_limit,
        }

        dashboard = self.build_management_layer(dashboard_filters)
        filtered_history = self._filter_history_for_institutional(raw_filters, history_limit)

        nucleo_filter = normalizar_texto(str(raw_filters.get("nucleo", "") or "").strip())
        municipio_filter = normalizar_texto(str(raw_filters.get("municipio", "") or "").strip())
        equipe_filter = normalizar_texto(str(raw_filters.get("equipe", "") or "").strip())

        analise_limite_completo = max(top_n * 3, 30)
        analise_por_nucleo_completa = self._build_institutional_nucleo_analysis(
            filtered_history,
            top_n=analise_limite_completo,
            nucleo_filter=nucleo_filter,
            municipio_filter=municipio_filter,
            equipe_filter=equipe_filter,
            include_history_alerts=False,
        )
        analise_por_nucleo = list(analise_por_nucleo_completa[:top_n])
        analise_por_nucleo_apendice = list(analise_por_nucleo_completa[top_n:])
        analise_tecnica_por_nucleo = self._build_institutional_nucleo_analysis(
            filtered_history,
            top_n=analise_limite_completo,
            nucleo_filter=nucleo_filter,
            municipio_filter=municipio_filter,
            equipe_filter=equipe_filter,
            include_history_alerts=True,
        )

        kpis = dict(dashboard.get("kpis_principais", {}) or {})
        total_processamentos = int(kpis.get("total_processamentos", 0) or 0)
        total_execucoes = int(kpis.get("total_execucoes", 0) or 0)
        total_frentes = int(dashboard.get("consolidado_periodo", {}).get("frentes", 0) or 0)
        total_ocorrencias = int(kpis.get("total_ocorrencias", 0) or 0)
        total_volume = float(dashboard.get("consolidado_periodo", {}).get("volume_total", 0.0) or 0.0)
        total_volume_fmt = str(dashboard.get("consolidado_periodo", {}).get("volume_total_fmt", "0") or "0")
        processamentos_alerta = int(kpis.get("processamentos_com_alerta", 0) or 0)
        processamentos_erro = int(kpis.get("processamentos_erro", 0) or 0)
        total_mapeados = int(kpis.get("total_mapeados", 0) or 0)
        total_nao_mapeados = int(kpis.get("total_nao_mapeados", 0) or 0)

        ranking_nucleos = list(
            dashboard.get("ranking_nucleos_volume", [])
            or dashboard.get("ranking_nucleos_processamentos", [])
            or []
        )
        ranking_equipes = list(
            dashboard.get("ranking_equipes_volume", [])
            or dashboard.get("ranking_equipes_processamentos", [])
            or []
        )
        ranking_servicos = list(dashboard.get("ranking_servicos", []) or [])[:top_n]
        ranking_ocorrencias = list(dashboard.get("ranking_ocorrencias", []) or [])[:top_n]
        ranking_categorias = list(dashboard.get("volume_por_categoria", []) or [])[:top_n]
        resumo_periodo = list(dashboard.get("resumo_por_periodo", []) or [])[:top_n]
        nao_mapeados_recorrentes = list(dashboard.get("ranking_nao_mapeados", []) or [])[:top_n]
        alertas_recorrentes = list(dashboard.get("alertas_recorrentes", []) or [])[:top_n]

        def _humanize_label(value: object, domain: str = "geral") -> str:
            return self._institutional_label(value, domain=domain)

        def _plural(valor: int, singular: str, plural: str) -> str:
            return singular if int(valor) == 1 else plural

        def _top_label(rows: List[dict], field: str) -> str:
            if not rows:
                return "-"
            return _humanize_label(rows[0].get(field, ""))

        def _fmt_pct(raw: object) -> str:
            return str(raw or "0")

        ranking_servicos_fmt: List[dict] = []
        for row in ranking_servicos:
            item = dict(row)
            item["servico"] = _humanize_label(item.get("servico", ""), domain="servico")
            item["categoria"] = _humanize_label(item.get("categoria", ""), domain="categoria")
            registros = int(item.get("registros", 0) or 0)
            volume_total_item = float(item.get("volume_total", 0.0) or 0.0)
            volume_total_fmt_item = self._display_number(volume_total_item)
            unidade_item = str(item.get("unidades", "") or "").strip()
            unidade_suffix = "" if not unidade_item or unidade_item == "-" else f" {unidade_item}"
            item["descricao_institucional"] = (
                f"{item['servico']}, com quantitativo consolidado de {volume_total_fmt_item}{unidade_suffix} no período"
            )
            if registros > 0:
                item["descricao_institucional"] += (
                    f" ({self._count_label(registros, 'registro', 'registros')})."
                )
            else:
                item["descricao_institucional"] += "."
            ranking_servicos_fmt.append(item)
        ranking_servicos = ranking_servicos_fmt

        ranking_categorias_fmt: List[dict] = []
        for row in ranking_categorias:
            item = dict(row)
            item["categoria"] = _humanize_label(item.get("categoria", ""), domain="categoria")
            item["descricao_institucional"] = (
                f"{item['categoria']}, com volume consolidado de {item.get('volume_total_fmt', '0')}."
            )
            ranking_categorias_fmt.append(item)
        ranking_categorias = ranking_categorias_fmt

        ranking_ocorrencias_fmt: List[dict] = []
        for row in ranking_ocorrencias:
            item = dict(row)
            item["tipo_ocorrencia"] = _humanize_label(
                item.get("tipo_ocorrencia", ""), domain="ocorrencia"
            )
            total = int(item.get("ocorrencias", 0) or 0)
            item["descricao_institucional"] = (
                f"{item['tipo_ocorrencia']}, com {self._count_label(total, 'ocorrência', 'ocorrências')}."
            )
            ranking_ocorrencias_fmt.append(item)
        ranking_ocorrencias = ranking_ocorrencias_fmt

        ranking_nao_mapeados_fmt: List[dict] = []
        for row in nao_mapeados_recorrentes:
            item = dict(row)
            item["termo"] = self._institutional_text_or_fallback(item.get("termo", ""))
            item["ultimo_nucleo"] = self._canonicalize_nucleo_for_aggregation(
                item.get("ultimo_nucleo", "")
            ) or self._institutional_text_or_fallback(item.get("ultimo_nucleo", ""))
            item["ultimo_municipio"] = self._canonicalize_municipio_for_aggregation(
                item.get("ultimo_municipio", "")
            ) or "Não informado"
            item["ultimo_equipe"] = self._canonicalize_equipe_for_aggregation(
                item.get("ultimo_equipe", "")
            ) or "Não informado"
            ranking_nao_mapeados_fmt.append(item)
        nao_mapeados_recorrentes = ranking_nao_mapeados_fmt

        entidades_nucleos = set()
        entidades_municipios = set()
        entidades_equipes = set()
        runs_multi_nucleos = 0
        runs_multi_logradouros = 0
        runs_multi_equipes = 0

        for row in filtered_history:
            entities = dict(row.get("_entities", {}) or {})
            if not entities:
                entities = self._collect_management_entities_for_row(row)

            nucleos_row = list(entities.get("nucleos", []) or [])
            municipios_row = list(entities.get("municipios", []) or [])
            equipes_row = list(entities.get("equipes", []) or [])
            logradouros_row = self._explode_history_values(row.get("logradouro", ""))

            for nuc in nucleos_row:
                if nuc:
                    entidades_nucleos.add(nuc)
            for mun in municipios_row:
                if mun:
                    entidades_municipios.add(mun)
            for eqp in equipes_row:
                if eqp:
                    entidades_equipes.add(eqp)

            if len(nucleos_row) > 1:
                runs_multi_nucleos += 1
            if len(logradouros_row) > 1:
                runs_multi_logradouros += 1
            if len(equipes_row) > 1:
                runs_multi_equipes += 1

        municipio_ausente = 0
        for row in filtered_history:
            municipios = self._explode_history_values(row.get("municipio", ""))
            if not municipios:
                municipio_ausente += 1

        comparativos = dict(dashboard.get("comparativos_basicos", {}) or {})
        periodo_label = str(dashboard.get("consolidado_periodo", {}).get("periodo_obra", "-") or "-")
        if comparativos.get("has_comparativo"):
            periodo_label = str(comparativos.get("label_atual", "-") or "-")

        top_nucleo = (
            str(analise_por_nucleo[0].get("nucleo", "") or "").strip()
            if analise_por_nucleo
            else _top_label(ranking_nucleos, "nucleo")
        ) or "-"
        top_equipe = _top_label(ranking_equipes, "equipe")
        top_servico = _top_label(ranking_servicos, "servico")
        top_ocorrencia = _top_label(ranking_ocorrencias, "tipo_ocorrencia")
        top_categoria = _top_label(ranking_categorias, "categoria")

        top_nucleo = self._institutional_text_or_fallback(top_nucleo)
        top_equipe = self._institutional_text_or_fallback(top_equipe)
        top_servico = self._institutional_text_or_fallback(top_servico)
        top_ocorrencia = self._institutional_text_or_fallback(top_ocorrencia)
        top_categoria = self._institutional_text_or_fallback(top_categoria)

        pontos_atencao: List[str] = []
        if processamentos_erro > 0:
            pontos_atencao.append(self._count_label(processamentos_erro, "mensagem com erro", "mensagens com erro"))
        if processamentos_alerta > 0:
            pontos_atencao.append(
                self._count_label(
                    processamentos_alerta,
                    "mensagem com alerta operacional",
                    "mensagens com alerta operacional",
                )
            )
        if total_nao_mapeados > 0:
            pontos_atencao.append(self._count_label(total_nao_mapeados, "item não mapeado", "itens não mapeados"))

        if pontos_atencao:
            linha_atencao = (
                "Pontos de atenção imediata: "
                f"{self._join_natural_phrases(pontos_atencao)}."
            )
        else:
            linha_atencao = "Pontos de atenção imediata: não foram identificadas inconformidades críticas no recorte."

        encaminhamentos_resumo: List[str] = []
        if total_nao_mapeados > 0:
            encaminhamentos_resumo.append("concluir o tratamento de itens não mapeados")
        if processamentos_erro > 0:
            encaminhamentos_resumo.append("encerrar o tratamento das mensagens com erro")
        if processamentos_alerta > 0:
            encaminhamentos_resumo.append("manter monitoramento ativo dos alertas recorrentes")
        if analise_por_nucleo:
            nucleo_prioritario = self._institutional_text_or_fallback(
                analise_por_nucleo[0].get("nucleo", "")
            )
            encaminhamentos_resumo.append(
                f"acompanhar de forma prioritária o núcleo {nucleo_prioritario}"
            )
        if not encaminhamentos_resumo:
            encaminhamentos_resumo.append("manter o ritmo operacional com revisão periódica dos indicadores")

        resumo_linhas = [
            (
                f"No período analisado, o quantitativo real consolidado foi de {total_volume_fmt}, "
                f"com {total_execucoes} {_plural(total_execucoes, 'execução registrada', 'execuções registradas')} "
                f"(apoio: {total_processamentos} {_plural(total_processamentos, 'mensagem processada', 'mensagens processadas')})."
            ),
            (
                f"A concentração operacional permaneceu em {top_nucleo}, com recorrência da equipe {top_equipe}, "
                f"predominância de {top_servico} e maior representatividade da categoria {top_categoria}."
            ),
            linha_atencao,
            "Encaminhamento gerencial recomendado: "
            f"{self._join_natural_phrases(encaminhamentos_resumo[:3])}.",
        ]

        inconsistencias_final: List[str] = []
        if processamentos_erro > 0:
            inconsistencias_final.append(f"Processamentos com erro no período: {processamentos_erro}.")
        if municipio_ausente > 0:
            inconsistencias_final.append(f"Registros sem município informado: {municipio_ausente}.")
        if total_nao_mapeados > 0:
            inconsistencias_final.append(f"Itens não mapeados no período: {total_nao_mapeados}.")
        if not inconsistencias_final:
            inconsistencias_final.append("Não foram identificadas pendências críticas no recorte aplicado.")

        inconsistencias_tecnicas: List[str] = []
        if runs_multi_nucleos > 0:
            inconsistencias_tecnicas.append(
                f"Mensagens com múltiplos núcleos detectados: {runs_multi_nucleos}."
            )
        if runs_multi_logradouros > 0:
            inconsistencias_tecnicas.append(
                f"Mensagens com múltiplos logradouros detectados: {runs_multi_logradouros}."
            )
        if runs_multi_equipes > 0:
            inconsistencias_tecnicas.append(
                f"Mensagens com múltiplas equipes detectadas: {runs_multi_equipes}."
            )
        for alerta in alertas_recorrentes[:5]:
            texto_alerta = str(alerta.get("alerta", "") or "").strip()
            qtd = int(alerta.get("ocorrencias", 0) or 0)
            if texto_alerta:
                inconsistencias_tecnicas.append(f"[Alerta interno] {texto_alerta} ({qtd}).")
        if not inconsistencias_tecnicas:
            inconsistencias_tecnicas.append("Sem alertas técnicos internos relevantes no recorte.")

        volume_total_fmt = str(
            dashboard.get("consolidado_periodo", {}).get("volume_total_fmt", "0") or "0"
        )
        panorama_topicos = [
            {
                "titulo": "Serviços com maior volume",
                "texto": (
                    str(ranking_servicos[0].get("descricao_institucional", "") or "").strip()
                    if ranking_servicos
                    else "Sem serviços relevantes no recorte."
                ),
            },
            {
                "titulo": "Categorias com maior representatividade",
                "texto": (
                    str(ranking_categorias[0].get("descricao_institucional", "") or "").strip()
                    if ranking_categorias
                    else "Sem categorias relevantes no recorte."
                ),
            },
            {
                "titulo": "Ocorrências mais recorrentes",
                "texto": (
                    str(ranking_ocorrencias[0].get("descricao_institucional", "") or "").strip()
                    if ranking_ocorrencias
                    else "Sem ocorrências recorrentes no recorte."
                ),
            },
        ]
        if total_processamentos > 0:
            nucleos_texto = self._count_label(
                len(entidades_nucleos), "núcleo", "núcleos"
            )
            equipes_texto = self._count_label(
                len(entidades_equipes), "equipe", "equipes"
            )
            panorama_texto = (
                f"A operação consolidou volume total de {volume_total_fmt} no período, com "
                f"{self._count_label(total_execucoes, 'execução', 'execuções')} distribuídas em {nucleos_texto} e {equipes_texto}. "
                "A leitura analítica está organizada por serviços, categorias e ocorrências prioritárias."
            )
        else:
            panorama_texto = "Não há dados suficientes no recorte para compor o panorama operacional."

        if total_processamentos == 0:
            conclusao = "Não há processamentos no recorte selecionado. Ajuste os filtros para gerar a leitura institucional."
        else:
            recomendacoes: List[str] = []
            if total_nao_mapeados > 0:
                recomendacoes.append("priorizar o fechamento dos itens não mapeados remanescentes")
            else:
                recomendacoes.append("manter o padrão atual de mapeamento e validação")
            if processamentos_erro > 0:
                recomendacoes.append("eliminar as mensagens com erro ainda pendentes")
            if processamentos_alerta > 0:
                recomendacoes.append("atuar sobre os alertas operacionais mais recorrentes")
            if analise_por_nucleo:
                recomendacoes.append(
                    f"acompanhar os núcleos com maior concentração operacional, com foco inicial em {top_nucleo}"
                )

            recomendacoes = recomendacoes[:3]
            recomendacao_texto = "; ".join(
                f"{idx + 1}) {texto}" for idx, texto in enumerate(recomendacoes)
            )
            conclusao = (
                f"No recorte analisado, o quantitativo consolidado foi de {total_volume_fmt}, "
                f"com maior presença em {top_nucleo}. "
                f"Como informação complementar, houve {self._count_label(total_processamentos, 'mensagem processada', 'mensagens processadas')}. "
                f"Encaminhamentos gerenciais recomendados: {recomendacao_texto}."
            )

        def _render_filter_value(value: object, fallback: str = "Todos") -> str:
            text = str(value or "").strip()
            return text if text else fallback

        filtros_aplicados = [
            {"label": "Período da obra (de)", "valor": _render_filter_value(raw_filters.get("obra_from", ""), "-")},
            {"label": "Período da obra (até)", "valor": _render_filter_value(raw_filters.get("obra_to", ""), "-")},
            {"label": "Processado em (de)", "valor": _render_filter_value(raw_filters.get("processed_from", ""), "-")},
            {"label": "Processado em (até)", "valor": _render_filter_value(raw_filters.get("processed_to", ""), "-")},
            {"label": "Núcleo", "valor": _render_filter_value(raw_filters.get("nucleo", ""))},
            {"label": "Município", "valor": _render_filter_value(raw_filters.get("municipio", ""))},
            {"label": "Equipe", "valor": _render_filter_value(raw_filters.get("equipe", ""))},
            {"label": "Status", "valor": _render_filter_value(status_filter, "Todos")},
            {"label": "Alertas", "valor": _render_filter_value(alertas_filter, "Todos")},
        ]

        final_payload = {
            "has_data": bool(total_processamentos),
            "header": {
                "titulo": "Relatório Institucional de Acompanhamento Operacional",
                "periodo": periodo_label,
                "emissao": datetime.now().strftime("%d/%m/%Y %H:%M"),
                "filtros_aplicados": filtros_aplicados,
            },
            "resumo_executivo": {
                "linhas": resumo_linhas,
            },
            "indicadores_principais": {
                "total_volume": total_volume,
                "total_volume_fmt": total_volume_fmt,
                "total_processamentos": total_processamentos,
                "total_execucoes": total_execucoes,
                "total_frentes": total_frentes,
                "total_ocorrencias": total_ocorrencias,
                "total_nucleos": len(entidades_nucleos) if entidades_nucleos else int(kpis.get("total_nucleos", 0) or 0),
                "total_municipios": len(entidades_municipios) if entidades_municipios else int(kpis.get("total_municipios", 0) or 0),
                "total_equipes": len(entidades_equipes) if entidades_equipes else int(kpis.get("total_equipes", 0) or 0),
                "total_mapeados": total_mapeados,
                "total_nao_mapeados": total_nao_mapeados,
                "percentual_mapeado_fmt": str(kpis.get("percentual_mapeado_fmt", "0") or "0"),
                "percentual_nao_mapeado_fmt": str(kpis.get("percentual_nao_mapeado_fmt", "0") or "0"),
                "processamentos_alerta": processamentos_alerta,
                "processamentos_erro": processamentos_erro,
            },
            "panorama_operacional": {
                "texto_analitico": panorama_texto,
                "topicos": panorama_topicos,
                "servicos_recorrentes": ranking_servicos,
                "categorias_recorrentes": ranking_categorias,
                "ocorrencias_recorrentes": ranking_ocorrencias,
                "resumo_periodo": resumo_periodo,
            },
            "analise_por_nucleo": analise_por_nucleo,
            "analise_por_nucleo_apendice": analise_por_nucleo_apendice,
            "alertas_pendencias": {
                "processamentos_alerta": processamentos_alerta,
                "processamentos_erro": processamentos_erro,
                "municipio_ausente": municipio_ausente,
                "nao_mapeados_recorrentes": nao_mapeados_recorrentes,
                "inconsistencias": inconsistencias_final,
            },
            "conclusao": conclusao,
            "top_n": top_n,
            "filters_applied": {
                "obra_from": str(raw_filters.get("obra_from", "") or ""),
                "obra_to": str(raw_filters.get("obra_to", "") or ""),
                "processed_from": str(raw_filters.get("processed_from", "") or ""),
                "processed_to": str(raw_filters.get("processed_to", "") or ""),
                "nucleo": str(raw_filters.get("nucleo", "") or ""),
                "municipio": str(raw_filters.get("municipio", "") or ""),
                "equipe": str(raw_filters.get("equipe", "") or ""),
                "status": status_filter,
                "alertas": alertas_filter,
            },
        }

        previa_tecnica = {
            "runs_analisadas": len(filtered_history),
            "analise_por_nucleo_tecnica": analise_tecnica_por_nucleo,
            "alertas_internos": inconsistencias_tecnicas,
            "totais": {
                "multiplos_nucleos": runs_multi_nucleos,
                "multiplos_logradouros": runs_multi_logradouros,
                "multiplas_equipes": runs_multi_equipes,
                "nao_mapeados": total_nao_mapeados,
                "processamentos_erro": processamentos_erro,
            },
        }

        return {
            **final_payload,
            "relatorio_final": dict(final_payload),
            "previa_tecnica": previa_tecnica,
        }

    def export_institutional_docx(self, report: Dict[str, object]) -> Tuple[bytes, str]:
        final_report = dict(report.get("relatorio_final", {}) or report or {})
        header = dict(final_report.get("header", {}) or {})
        indicadores = dict(final_report.get("indicadores_principais", {}) or {})
        panorama = dict(final_report.get("panorama_operacional", {}) or {})
        alertas = dict(final_report.get("alertas_pendencias", {}) or {})

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        title = str(
            header.get("titulo", "Relatório institucional")
            or "Relatório institucional"
        )
        doc.add_heading(title, 0)
        doc.add_paragraph(f"Período analisado: {header.get('periodo', '-')}")
        doc.add_paragraph(f"Data de emissão: {header.get('emissao', '-')}")

        filtros = list(header.get("filtros_aplicados", []) or [])
        if filtros:
            doc.add_heading("Filtros aplicados", level=1)
            for item in filtros:
                label = str(item.get("label", "") or "-")
                valor = str(item.get("valor", "") or "-")
                doc.add_paragraph(f"{label}: {valor}", style="List Bullet")

        doc.add_heading("Resumo executivo", level=1)
        for linha in list(final_report.get("resumo_executivo", {}).get("linhas", []) or []):
            doc.add_paragraph(str(linha), style="List Bullet")

        doc.add_heading("Indicadores principais", level=1)
        kpi_table = doc.add_table(rows=1, cols=2)
        kpi_table.rows[0].cells[0].text = "Indicador"
        kpi_table.rows[0].cells[1].text = "Valor"
        kpi_rows = [
            ("Quantitativo consolidado", indicadores.get("total_volume_fmt", "0")),
            ("Total de processamentos", indicadores.get("total_processamentos", 0)),
            ("Total de execuções", indicadores.get("total_execucoes", 0)),
            ("Total de frentes", indicadores.get("total_frentes", 0)),
            ("Total de ocorrências", indicadores.get("total_ocorrencias", 0)),
            ("Núcleos atendidos", indicadores.get("total_nucleos", 0)),
            ("Municípios atendidos", indicadores.get("total_municipios", 0)),
            ("Equipes envolvidas", indicadores.get("total_equipes", 0)),
            ("Percentual mapeado", f"{indicadores.get('percentual_mapeado_fmt', '0')}%"),
            (
                "Percentual não mapeado",
                f"{indicadores.get('percentual_nao_mapeado_fmt', '0')}%",
            ),
            ("Processamentos com alerta", indicadores.get("processamentos_alerta", 0)),
            ("Processamentos com erro", indicadores.get("processamentos_erro", 0)),
        ]
        for label, value in kpi_rows:
            row = kpi_table.add_row().cells
            row[0].text = str(label)
            row[1].text = str(value)

        doc.add_heading("Panorama operacional", level=1)
        texto_panorama = str(panorama.get("texto_analitico", "") or "").strip()
        if texto_panorama:
            doc.add_paragraph(texto_panorama)
        topicos_panorama = list(panorama.get("topicos", []) or [])
        for topico in topicos_panorama:
            titulo = self._institutional_text_or_fallback(topico.get("titulo", ""))
            texto = self._institutional_text_or_fallback(topico.get("texto", ""))
            doc.add_paragraph(f"{titulo}: {texto}", style="List Bullet")
        for servico in list(panorama.get("servicos_recorrentes", []) or [])[:5]:
            doc.add_paragraph(
                str(servico.get("descricao_institucional", "-") or "-"),
                style="List Bullet",
            )
        for categoria in list(panorama.get("categorias_recorrentes", []) or [])[:5]:
            doc.add_paragraph(
                str(categoria.get("descricao_institucional", "-") or "-"),
                style="List Bullet",
            )
        for ocorrencia in list(panorama.get("ocorrencias_recorrentes", []) or [])[:5]:
            doc.add_paragraph(
                str(ocorrencia.get("descricao_institucional", "-") or "-"),
                style="List Bullet",
            )

        doc.add_heading("Análise por núcleo", level=1)
        for nucleo in list(final_report.get("analise_por_nucleo", []) or []):
            doc.add_heading(
                str(nucleo.get("nucleo", "(sem núcleo)") or "(sem núcleo)"), level=2
            )
            doc.add_paragraph(
                f"Quantitativo consolidado no período: {nucleo.get('volume_total_fmt', '0')} "
                f"(apoio: {self._count_label(int(nucleo.get('processamentos', 0) or 0), 'registro operacional', 'registros operacionais')})"
            )
            logradouros = list(nucleo.get("logradouros", []) or [])
            doc.add_paragraph(
                f"Município de referência: {nucleo.get('municipio', '-')}"
            )
            doc.add_paragraph(
                f"{'Logradouro principal' if len(logradouros) <= 1 else 'Logradouros de referência'}: {nucleo.get('logradouro', '-')}"
            )

            equipes = list(nucleo.get("equipes", []) or [])
            if equipes:
                doc.add_paragraph("Equipes mobilizadas:")
                for equipe in equipes[:6]:
                    doc.add_paragraph(str(equipe), style="List Bullet")

            servicos = list(nucleo.get("principais_servicos", []) or [])
            if servicos:
                doc.add_paragraph("Serviços predominantes:")
                for servico in servicos[:5]:
                    doc.add_paragraph(
                        str(servico.get("descricao_institucional", "-") or "-"),
                        style="List Bullet",
                    )

            ocorrencias = list(nucleo.get("principais_ocorrencias", []) or [])
            if ocorrencias:
                doc.add_paragraph("Ocorrências relevantes:")
                for ocorrencia in ocorrencias[:5]:
                    doc.add_paragraph(
                        str(ocorrencia.get("descricao_institucional", "-") or "-"),
                        style="List Bullet",
                    )

            observacoes = list(nucleo.get("observacoes_relevantes", []) or [])
            if observacoes:
                doc.add_paragraph("Observações do período:")
                for obs in observacoes[:3]:
                    doc.add_paragraph(str(obs), style="List Bullet")

            leitura = str(nucleo.get("observacao_analitica", "") or "").strip()
            if leitura:
                doc.add_paragraph(f"Leitura executiva: {leitura}")

        analise_apendice = list(final_report.get("analise_por_nucleo_apendice", []) or [])
        if analise_apendice:
            doc.add_heading("Apêndice - Demais núcleos do recorte", level=1)
            for nucleo in analise_apendice:
                nome_nucleo = self._institutional_text_or_fallback(nucleo.get("nucleo", ""))
                volume_nucleo = str(nucleo.get("volume_total_fmt", "0") or "0")
                registros_nucleo = int(nucleo.get("processamentos", 0) or 0)
                alertas_nucleo = int(nucleo.get("processamentos_alerta", 0) or 0)
                doc.add_paragraph(
                    f"{nome_nucleo}: quantitativo {volume_nucleo}, "
                    f"{self._count_label(registros_nucleo, 'registro', 'registros')} "
                    f"e {self._count_label(alertas_nucleo, 'alerta', 'alertas')}.",
                    style="List Bullet",
                )

        doc.add_heading("Alertas e pendências", level=1)
        doc.add_paragraph(
            f"Mensagens com alerta operacional: {alertas.get('processamentos_alerta', 0)}",
            style="List Bullet",
        )
        doc.add_paragraph(
            f"Mensagens com erro: {alertas.get('processamentos_erro', 0)}",
            style="List Bullet",
        )
        doc.add_paragraph(
            f"Itens sem município informado: {alertas.get('municipio_ausente', 0)}",
            style="List Bullet",
        )

        for item in list(alertas.get("inconsistencias", []) or []):
            doc.add_paragraph(str(item), style="List Bullet")

        nao_mapeados = list(alertas.get("nao_mapeados_recorrentes", []) or [])
        if nao_mapeados:
            doc.add_paragraph("Termos não mapeados mais recorrentes:")
            for row in nao_mapeados[:5]:
                doc.add_paragraph(
                    f"{row.get('termo', '-')}: {self._count_label(int(row.get('ocorrencias', 0) or 0), 'ocorrência', 'ocorrências')}.",
                    style="List Bullet",
                )

        doc.add_heading("Conclusão", level=1)
        doc.add_paragraph(str(final_report.get("conclusao", "-") or "-"))

        buffer = BytesIO()
        doc.save(buffer)
        filename = f"relatorio_institucional_{datetime.now():%Y%m%d_%H%M%S}.docx"
        return buffer.getvalue(), filename

    def generate_from_draft(
        self,
        draft_id: str,
        overrides: Dict[str, object],
        contract_id: object = "",
        contract_label: object = "",
    ) -> dict:
        draft = self.load_draft(draft_id)
        raw_message = str(draft.get("raw_message", "") or "")
        parsed = self.apply_overrides(draft.get("parsed", {}), overrides)
        parsed = self.reconcile_with_nucleo_master(parsed)
        contract_value = str(contract_label or "").strip() or str(contract_id or "").strip()
        parsed = self.bind_contract_to_parsed(parsed, contract_value, force=True)
        if not str(contract_label or "").strip() and contract_value:
            contract_label = contract_value

        alert_items = self.build_alert_items(parsed)
        preview_alerts = [a["message"] for a in alert_items]
        output_dir = self._create_output_dir()
        entrada_path = output_dir / "entrada_original.txt"
        if raw_message.strip():
            entrada_path.write_text(raw_message.strip() + "\n", encoding="utf-8")

        save_parsed_outputs(parsed, output_dir)
        self.report_generator.generate_nucleus_reports(parsed, output_dir / "relatorios_nucleos")
        build_management_workbook(
            output_dir,
            self.dictionary_csv,
            nucleo_reference_file=self.nucleo_reference_file,
        )

        master_stats = update_master_from_output(
            output_dir,
            self.master_dir,
            self.dictionary_csv,
            nucleo_reference_file=self.nucleo_reference_file,
        )

        main = self.extract_main_fields(parsed)
        context = self.collect_context_overview(parsed)
        main_lists = self.build_main_lists(main, context)
        unmapped = self.collect_unmapped(parsed)

        history_nucleos = self._join_history_display_values(
            context.get("nucleos", []),
            canonicalizer=self._canonicalize_nucleo_for_aggregation,
        )
        history_nucleo = history_nucleos or main.get("nucleo", "")
        history_nucleo_values = self._split_nucleo_text_for_aggregation(history_nucleo)
        history_municipio = self._resolve_history_final_municipio(
            {
                "nucleo": history_nucleo,
                "municipio": main.get("municipio", ""),
                "municipio_oficial": parsed.get("municipio_oficial", ""),
            }
        )

        result = {
            "status": "sucesso",
            "output_dir": str(output_dir),
            "base_gerencial_path": str(output_dir / "base_gerencial.xlsx"),
            "master_dir": str(self.master_dir),
            "master_stats": master_stats,
            "alerts": preview_alerts,
            "alert_items": alert_items,
            "summary": self.summarize_preview(parsed),
            "main_fields": main,
            "main_lists": main_lists,
            "context_overview": context,
            "nao_mapeados": len(unmapped),
            "correcoes_manuais": list(parsed.get("correcoes_manuais", [])),
            "files": {
                "execucao_csv": str(output_dir / "execucao.csv"),
                "frentes_csv": str(output_dir / "frentes.csv"),
                "ocorrencias_csv": str(output_dir / "ocorrencias.csv"),
                "observacoes_csv": str(output_dir / "observacoes.csv"),
                "nao_mapeado_csv": str(output_dir / "servico_nao_mapeado.csv"),
                "base_gerencial": str(output_dir / "base_gerencial.xlsx"),
            },
        }

        self._append_history(
            {
                "processed_at": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
                "obra_data": main.get("data", ""),
                "nucleo": history_nucleo,
                "nucleo_detectado_texto": str(parsed.get("nucleo_detectado_texto", "") or ""),
                "nucleo_oficial": str(parsed.get("nucleo_oficial", "") or ""),
                "logradouro": main.get("logradouro", ""),
                "municipio": history_municipio,
                "municipio_detectado_texto": str(parsed.get("municipio_detectado_texto", "") or ""),
                "municipio_oficial": str(parsed.get("municipio_oficial", "") or ""),
                "nucleo_status_cadastro": str(parsed.get("nucleo_status_cadastro", "") or ""),
                "equipe": main.get("equipe", ""),
                "status": "sucesso",
                "contract_id": str(contract_id or "").strip(),
                "contract_label": str(contract_label or "").strip(),
                "output_dir": result["output_dir"],
                "base_gerencial_path": result["base_gerencial_path"],
                "master_dir": result["master_dir"],
                "nao_mapeados": str(result["nao_mapeados"]),
                "alertas": " | ".join(preview_alerts),
                "mensagem": raw_message[:350],
            }
        )
        return result

    def register_error_history(
        self,
        draft_id: str,
        overrides: Dict[str, object],
        error_message: str,
        contract_id: object = "",
        contract_label: object = "",
    ) -> None:
        try:
            draft = self.load_draft(draft_id)
            parsed = self.apply_overrides(draft.get("parsed", {}), overrides)
            parsed = self.reconcile_with_nucleo_master(parsed)
            contract_value = str(contract_label or "").strip() or str(contract_id or "").strip()
            parsed = self.bind_contract_to_parsed(parsed, contract_value, force=True)
            if not str(contract_label or "").strip() and contract_value:
                contract_label = contract_value
            raw_message = str(draft.get("raw_message", "") or "")
            main = self.extract_main_fields(parsed)
        except Exception:
            parsed = {}
            raw_message = ""
            main = {"data": "", "nucleo": "", "logradouro": "", "municipio": "", "equipe": ""}

        self._append_history(
            {
                "processed_at": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
                "obra_data": main.get("data", ""),
                "nucleo": main.get("nucleo", ""),
                "nucleo_detectado_texto": str(parsed.get("nucleo_detectado_texto", "") or ""),
                "nucleo_oficial": str(parsed.get("nucleo_oficial", "") or ""),
                "logradouro": main.get("logradouro", ""),
                "municipio": main.get("municipio", ""),
                "municipio_detectado_texto": str(parsed.get("municipio_detectado_texto", "") or ""),
                "municipio_oficial": str(parsed.get("municipio_oficial", "") or ""),
                "nucleo_status_cadastro": str(parsed.get("nucleo_status_cadastro", "") or ""),
                "equipe": main.get("equipe", ""),
                "status": f"erro: {error_message}",
                "contract_id": str(contract_id or "").strip(),
                "contract_label": str(contract_label or "").strip(),
                "output_dir": "",
                "base_gerencial_path": "",
                "master_dir": str(self.master_dir),
                "nao_mapeados": "",
                "alertas": error_message,
                "mensagem": raw_message[:350],
            }
        )

    def copy_output_to(self, source_output_dir: Path, target_dir: Path) -> None:
        target_dir.mkdir(parents=True, exist_ok=True)
        for file in source_output_dir.iterdir():
            src = Path(file)
            dst = target_dir / src.name
            if src.is_file():
                shutil.copyfile(src, dst)


























































