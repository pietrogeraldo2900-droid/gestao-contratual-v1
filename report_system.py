from __future__ import annotations

import csv
import json
import re
import unicodedata
from collections import defaultdict
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from input_layer import extrair_primeira_equipe


EMOJIS_TO_STRIP = ["ðŸ“Š", "ðŸ“…", "ðŸ“", "âœ…", "âš ï¸", "ðŸ“Œ", "â›ˆï¸"]


def strip_accents(text: str) -> str:
    return "".join(c for c in unicodedata.normalize("NFKD", text) if not unicodedata.combining(c))


def slugify(text: str) -> str:
    text = strip_accents(text).lower()
    text = re.sub(r"[^a-z0-9]+", "_", text).strip("_")
    return text or "item"


def repair_mojibake(text: str) -> str:
    if not isinstance(text, str):
        return text
    suspicious = ("Ã§", "Ã£", "Ã¡", "Ã©", "Ãª", "Â ", "Âº", "Âª", "â€", "â€“", "â€”", "â€¢", "âœ", "ðŸ", "ï¿½")
    if any(token in text for token in suspicious):
        try:
            fixed = text.encode("latin1", errors="ignore").decode("utf-8", errors="ignore")
            if fixed:
                return fixed
        except Exception:
            pass
    return text


def normalize_numeric(value):
    try:
        num = float(value)
    except (TypeError, ValueError):
        return value
    if num.is_integer():
        return int(num)
    return round(num, 2)


def normalize_row_numbers(row: dict) -> dict:
    numeric_fields = {
        "quantidade", "quantidade_total", "quantidade_servicos", "quantidade_insumos",
        "frentes", "itens", "ocorrencias", "registros", "frentes_total",
        "itens_total", "ocorrencias_total", "frentes_sem_producao"
    }
    out = {}
    for k, v in row.items():
        out[k] = normalize_numeric(v) if k in numeric_fields else v
    return out


@dataclass
class Front:
    id_frente: str
    data_referencia: str
    contrato: str
    programa: str
    nucleo: str
    equipe: str
    logradouro: str
    municipio: str
    status_frente: str
    observacao_frente: str
    arquivo_origem: str


@dataclass
class ExecutionItem:
    id_item: str
    id_frente: str
    data_referencia: str
    contrato: str
    programa: str
    nucleo: str
    equipe: str
    logradouro: str
    municipio: str
    item_original: str
    item_normalizado: str
    categoria_item: str
    tipo_registro: str
    quantidade: float
    unidade: str
    material: str
    especificacao: str
    complemento: str
    observacao_item: str
    arquivo_origem: str
    mensagem_origem: str


@dataclass
class Occurrence:
    id_ocorrencia: str
    id_frente: str
    data_referencia: str
    contrato: str
    programa: str
    nucleo: str
    equipe: str
    logradouro: str
    municipio: str
    tipo_ocorrencia: str
    descricao: str
    impacto_producao: str
    arquivo_origem: str


class ServiceDictionary:
    def __init__(self, csv_path: Path):
        self.entries: List[dict] = []
        with csv_path.open("r", encoding="utf-8-sig") as f:
            reader = csv.DictReader(f)
            for row in reader:
                clean_row = {}
                for k, v in row.items():
                    clean_row[k] = repair_mojibake(v) if isinstance(v, str) else v
                clean_row["match_key"] = slugify(clean_row.get("nome_original", ""))
                clean_row["alias_key"] = slugify(clean_row.get("alias", ""))
                self.entries.append(clean_row)

    def match(self, item_name: str) -> dict:
        item_name = repair_mojibake(item_name)
        key = slugify(item_name)
        # exact or alias exact
        for e in self.entries:
            if key == e["match_key"] or (e["alias_key"] and key == e["alias_key"]):
                return e
        # contains/startswith heuristics
        best: Optional[dict] = None
        best_score = 0
        for e in self.entries:
            keys = [e["match_key"], e["alias_key"]]
            for candidate in keys:
                if not candidate:
                    continue
                score = 0
                if candidate in key or key in candidate:
                    score = min(len(candidate), len(key))
                if score > best_score:
                    best = e
                    best_score = score
        if best:
            return best
        return {
            "nome_original": item_name,
            "nome_padronizado": item_name,
            "categoria": "servico_nao_mapeado",
            "unidade_padrao": "servico",
            "tipo_registro": "servico",
            "frase_tecnica": f"Execucao de {strip_accents(item_name).lower()}",
            "alias": "",
        }


class WhatsAppReportParser:
    def __init__(self, dictionary: ServiceDictionary):
        self.dictionary = dictionary

    def parse_text(self, text: str, source_name: str = "mensagem_whatsapp.txt") -> dict:
        lines = [self._clean_line(ln) for ln in text.splitlines() if self._clean_line(ln)]
        raw = "\n".join(lines)

        contrato = self._extract_contract(raw)
        programa = self._extract_program(raw)
        data_referencia = self._extract_date(raw)
        municipio_global = self._extract_municipio(raw)

        relevant_lines = []
        for line in lines:
            if re.search(r"TOTAL\s+(GERAL|EXECUTADO)\s+DO\s+DIA", strip_accents(line), re.I):
                break
            relevant_lines.append(line)

        fronts: List[Front] = []
        execution: List[ExecutionItem] = []
        occurrences: List[Occurrence] = []

        i = 0
        front_counter = 0
        item_counter = 0
        occ_counter = 0

        while i < len(relevant_lines):
            line = relevant_lines[i]
            if self._is_front_header(line):
                front_counter += 1
                front_id = f"F{front_counter:03d}"
                nucleo, equipe, logradouro, municipio = self._parse_front_header(line)
                if not municipio:
                    municipio = municipio_global
                block_lines: List[str] = []
                i += 1
                while i < len(relevant_lines) and not self._is_front_header(relevant_lines[i]):
                    block_lines.append(relevant_lines[i])
                    i += 1

                front_status = "com_producao"
                front_obs: List[str] = []
                has_item = False
                for bl in block_lines:
                    if self._is_no_production(bl):
                        front_status = "sem_producao"
                    elif self._is_observation(bl):
                        front_obs.append(self._normalize_obs(bl))
                    elif self._looks_like_item(bl):
                        has_item = True
                if front_status == "com_producao" and front_obs:
                    front_status = "producao_parcial" if has_item else "paralisada"

                front = Front(
                    id_frente=front_id,
                    data_referencia=data_referencia,
                    contrato=contrato,
                    programa=programa,
                    nucleo=nucleo,
                    equipe=equipe,
                    logradouro=logradouro,
                    municipio=municipio,
                    status_frente=front_status,
                    observacao_frente=" | ".join(front_obs),
                    arquivo_origem=source_name,
                )
                fronts.append(front)

                for bl in block_lines:
                    if self._is_no_production(bl):
                        occ_counter += 1
                        occurrences.append(Occurrence(
                            id_ocorrencia=f"O{occ_counter:04d}",
                            id_frente=front_id,
                            data_referencia=data_referencia,
                            contrato=contrato,
                            programa=programa,
                            nucleo=nucleo,
                            equipe=equipe,
                            logradouro=logradouro,
                            municipio=municipio,
                            tipo_ocorrencia="sem_producao",
                            descricao=self._normalize_obs(bl),
                            impacto_producao="sim",
                            arquivo_origem=source_name,
                        ))
                        continue

                    if self._is_observation(bl):
                        description = self._normalize_obs(bl)
                        occ_counter += 1
                        occurrences.append(Occurrence(
                            id_ocorrencia=f"O{occ_counter:04d}",
                            id_frente=front_id,
                            data_referencia=data_referencia,
                            contrato=contrato,
                            programa=programa,
                            nucleo=nucleo,
                            equipe=equipe,
                            logradouro=logradouro,
                            municipio=municipio,
                            tipo_ocorrencia=self._classify_occurrence(description),
                            descricao=description,
                            impacto_producao=self._impact_level(description),
                            arquivo_origem=source_name,
                        ))
                        continue

                    if self._looks_like_item(bl):
                        parsed = self._parse_item_line(bl)
                        if parsed is None:
                            continue
                        item_counter += 1
                        item_name, qty, unit, material, specification, complement, item_obs = parsed
                        service = self.dictionary.match(item_name)
                        if not unit or (unit == "un" and service.get("unidade_padrao") not in {"", "un"} and re.fullmatch(r"\d+[\.,]?\d*", complement.strip())):
                            unit = service.get("unidade_padrao", "servico")
                        execution.append(ExecutionItem(
                            id_item=f"I{item_counter:04d}",
                            id_frente=front_id,
                            data_referencia=data_referencia,
                            contrato=contrato,
                            programa=programa,
                            nucleo=nucleo,
                            equipe=equipe,
                            logradouro=logradouro,
                            municipio=municipio,
                            item_original=item_name,
                            item_normalizado=service["nome_padronizado"],
                            categoria_item=service["categoria"],
                            tipo_registro=service["tipo_registro"],
                            quantidade=qty,
                            unidade=unit,
                            material=material,
                            especificacao=specification,
                            complemento=complement,
                            observacao_item=item_obs,
                            arquivo_origem=source_name,
                            mensagem_origem=bl,
                        ))
                continue
            i += 1

        return {
            "data_referencia": data_referencia,
            "contrato": contrato,
            "programa": programa,
            "arquivo_origem": source_name,
            "frentes": [asdict(f) for f in fronts],
            "execucao": [asdict(e) for e in execution],
            "ocorrencias": [asdict(o) for o in occurrences],
        }

    def _extract_contract(self, text: str) -> str:
        m = re.search(r"RDO\s*[â€“-]\s*[^\n]*[â€“-]\s*([^\n*]+)", text, flags=re.IGNORECASE)
        if m:
            return m.group(1).strip()
        return "Oeste 1"

    def _extract_program(self, text: str) -> str:
        m = re.search(r"AL\s*\d+", text, flags=re.IGNORECASE)
        if m:
            return "Água Legal"
        return "Água Legal"

    def _extract_date(self, text: str) -> str:
        m = re.search(r"(\d{2}/\d{2}/\d{4})", text)
        return m.group(1) if m else ""

    def _extract_municipio(self, text: str) -> str:
        m = re.search(r"municipio\s*:\s*([^\n]+)", strip_accents(text), flags=re.IGNORECASE)
        if m:
            return m.group(1).strip()
        return ""

    def _split_nucleo_municipio(self, nucleo_bruto: str) -> Tuple[str, str]:
        nucleo = str(nucleo_bruto or "").strip()
        municipio = ""
        m = re.match(r"^(.*?)\s*[—–-]\s*([A-Za-zÀ-ÿ][A-Za-zÀ-ÿ\s]+)$", nucleo)
        if not m:
            return nucleo, municipio

        left = m.group(1).strip()
        right = m.group(2).strip()
        right_norm = strip_accents(right.lower())
        blocked_tokens = [
            "agua",
            "esgoto",
            "caixa",
            "rede",
            "recomposicao",
            "viela",
            "rua",
            "tv",
            "travessa",
            "av",
            "equipe",
            "resp",
            "operacional",
        ]
        if "/" in right or "caixa uma" in right_norm or any(token in right_norm for token in blocked_tokens):
            return nucleo, ""
        if not left:
            return nucleo, ""
        return left, right

    def _normalize_nucleo_label(self, nucleo: str) -> str:
        text = str(nucleo or "").strip()
        if not text:
            return ""

        # Remove qualificadores operacionais no sufixo do nucleo
        # (ex.: "Mississipi - Esgoto", "Mississipi - Caixa UMA").
        suffix_pattern = re.compile(
            r"^(.*?)\s*[—–-]\s*(agua(?:/esgoto)?|esgoto|caixa\s*uma|caixa)\s*$",
            re.I,
        )
        m = suffix_pattern.match(text)
        if m and m.group(1).strip():
            text = m.group(1).strip()

        return text

    def _looks_like_logradouro(self, text: str) -> bool:
        base = strip_accents(str(text or "").lower())
        if not base:
            return False
        return bool(
            re.search(
                r"\b(viela|rua|travessa|acesso|avenida|av\.?|tv\.?)\b",
                base,
            )
        )


    def _clean_line(self, line: str) -> str:
        line = line.replace("\ufeff", "")

        # Tenta recuperar texto com mojibake comum (ex.: NÃºcleo, âœ…).
        if any(token in line for token in ("Ã§", "Ã£", "Ã¡", "Ã©", "Ãª", "Â ", "Âº", "Âª", "â€", "â€“", "â€”", "â€¢", "âœ", "ðŸ", "ï¿½")):
            try:
                line = line.encode("latin1", errors="ignore").decode("utf-8", errors="ignore")
            except Exception:
                pass

        for emoji in EMOJIS_TO_STRIP:
            line = line.replace(emoji, "")
        line = line.replace("*", "")
        line = re.sub(r"\s+", " ", line).strip()
        return line

    def _is_front_header(self, line: str) -> bool:
        base = strip_accents(line.lower())
        base = re.sub(r"[^a-z0-9:\s]", "", base)
        return bool(re.search(r"n.{0,2}cleo\s*:\s*", base))

    def _parse_front_header(self, line: str) -> Tuple[str, str, str, str]:
        if ":" not in line:
            return line.strip(), "", "", ""

        _, right = line.split(":", 1)
        conteudo = right.strip()
        m = re.match(r"(.+?)(?:\s*\((.+)\))?$", conteudo)
        if not m:
            return conteudo, "", "", ""

        nucleo_bruto = m.group(1).strip()
        nucleo, municipio = self._split_nucleo_municipio(nucleo_bruto)
        nucleo = self._normalize_nucleo_label(nucleo)
        inside = (m.group(2) or "").strip()
        equipe = ""
        logradouro = ""
        if inside:
            inside = re.sub(r"^Equipe\s+", "", inside, flags=re.I)
            inside = re.sub(r"^Resp\.??\s+", "", inside, flags=re.I)
            parts = re.split(r"\s+[—–-]\s+", inside, maxsplit=1)
            if len(parts) == 2:
                equipe, logradouro = parts[0].strip(), parts[1].strip()
            else:
                equipe = inside
        equipe = extrair_primeira_equipe(equipe)

        # Regra oficial: nao usar logradouro (viela/rua/travessa) como nucleo.
        if self._looks_like_logradouro(nucleo):
            if not logradouro:
                logradouro = nucleo
            nucleo = ""

        return nucleo, equipe, logradouro, municipio

    def _is_no_production(self, line: str) -> bool:
        s = strip_accents(line.lower())
        return "sem producao" in s

    def _is_observation(self, line: str) -> bool:
        s = strip_accents(line.lower())
        return s.startswith("obs:") or " obs:" in s or s.startswith("observacao") or s.startswith("obs ")

    def _looks_like_item(self, line: str) -> bool:
        s = strip_accents(line.lower())
        if not s:
            return False
        if self._is_observation(line) or self._is_no_production(line):
            return False
        if ":" in line:
            return True
        if re.match(r"^\d+[\.,]?\d*\s+valas?\s+recompostas", s):
            return True
        if any(k in s for k in ["concretagem de vala", "corte de pavimento", "retirada de big bag", "cortes d'agua", "cortes dâ€™agua"]):
            return True
        return False

    def _parse_item_line(self, line: str) -> Optional[Tuple[str, float, str, str, str, str, str]]:
        clean = line.strip().rstrip(".")
        item_obs = ""
        material = ""
        specification = ""
        complement = ""

        # Special case: 10 valas recompostas
        norm = strip_accents(clean.lower())
        m = re.match(r"^(\d+[\.,]?\d*)\s+valas?\s+recompostas(?:\s*\((.*?)\))?$", norm, re.I)
        if m:
            qty = float(m.group(1).replace(",", "."))
            extra = m.group(2) or ""
            item_name = "RecomposiÃ§Ã£o asfÃ¡ltica" if "asfalt" in extra else "RecomposiÃ§Ã£o de valas"
            return item_name, qty, "vala", "", "", clean, item_obs

        if ":" in clean:
            left, right = clean.split(":", 1)
            item_name = left.strip()
            details = right.strip()
        else:
            item_name = clean
            details = ""

        # Regra operacional: quando vier "Servico complementar: ...",
        # considera o texto apos ":" como servico principal.
        if strip_accents(item_name.lower()) in {"servico complementar", "servico comp", "atividade complementar"} and details:
            item_name = details
            m_comp = re.match(
                r"^(\d+[\.,]?\d*)\s*(mÂ²|m2|m|unidades|unidade|un|valas|vala|servi[cÃ§]os?|ramais?)\b\s*(.*)$",
                item_name,
                re.I,
            )
            if m_comp and m_comp.group(3).strip():
                details = f"{m_comp.group(1)} {m_comp.group(2)}"
                item_name = m_comp.group(3).strip()
            else:
                details = ""

        # Parenthetical metadata in item name
        paren = re.search(r"\((.*?)\)", item_name)
        if paren:
            content = paren.group(1)
            mat, spec = self._split_material_spec(content)
            material = mat or material
            specification = spec or specification
            if not mat and content and "mureta" in strip_accents(content.lower()):
                item_obs = "com mureta"
            item_name = re.sub(r"\(.*?\)", "", item_name).strip()

        # Material/spec in item name itself
        mat2 = re.search(r"\b(PEAD|PVC|concreto)\b\s*(Ã˜?\s*\d+)?", item_name, re.I)
        if mat2:
            material = mat2.group(1).upper()
            specification = (mat2.group(2) or "").replace(" ", "")
            item_name = re.sub(r"\b(PEAD|PVC|concreto)\b\s*(Ã˜?\s*\d+)?", "", item_name, flags=re.I).strip()

        qty = None
        unit = ""

        m = re.search(r"(\d+[\.,]?\d*)\s*(mÂ²|m2|m|unidades|unidade|un|valas|vala|servi[cÃ§]os?|ramais?)\b", details, re.I)
        if m:
            qty = float(m.group(1).replace(",", "."))
            unit = self._normalize_unit(m.group(2))
        else:
            m_simple = re.search(r"^(\d+[\.,]?\d*)$", details)
            if m_simple:
                qty = float(m_simple.group(1).replace(",", "."))
                unit = "un"

        if qty is None:
            qty = 1.0
            unit = "servico"

        # Material/spec in details
        if not material:
            m_mat = re.search(r"\b(PEAD|PVC|concreto)\b\s*(Ã˜?\s*\d+)?", details, re.I)
            if m_mat:
                material = m_mat.group(1).upper()
                specification = (m_mat.group(2) or "").replace(" ", "")
        if not specification:
            m_spec = re.search(r"Ã˜\s*\d+", details)
            if m_spec:
                specification = m_spec.group(0).replace(" ", "")

        if re.search(r"mureta", details, re.I):
            item_obs = "com mureta"

        item_name = item_name.strip()
        if re.match(r"^\d+\s+valas?\s+recompostas$", strip_accents(item_name.lower())):
            item_name = "RecomposiÃ§Ã£o de valas"
            unit = "vala"
        if strip_accents(item_name.lower()) == "recomposicao asfaltica":
            unit = "vala" if unit == "un" else unit

        complement = details
        complement = re.sub(r"\b(PEAD|PVC|concreto)\b\s*(Ã˜?\s*\d+)?", "", complement, flags=re.I)
        complement = re.sub(r"\b\d+[\.,]?\d*\s*(mÂ²|m2|m|unidades|unidade|un|valas|vala|servi[cÃ§]os?|ramais?)\b", "", complement, flags=re.I)
        complement = re.sub(r"^\s*\d+[\.,]?\d*\s*$", "", complement)
        complement = re.sub(r"\(\s*com mureta\s*\)", "", complement, flags=re.I)
        complement = re.sub(r"\bcom mureta\b", "", complement, flags=re.I)
        complement = re.sub(r"\s+", " ", complement).strip(" -;,.")

        return item_name, qty, unit, material, specification, complement, item_obs

    def _split_material_spec(self, content: str) -> Tuple[str, str]:
        m = re.search(r"(PEAD|PVC|concreto)\s*(Ã˜?\s*\d+)?", content, re.I)
        if not m:
            return "", content.strip()
        material = m.group(1).upper()
        specification = (m.group(2) or "").replace(" ", "")
        return material, specification

    def _normalize_unit(self, raw: str) -> str:
        raw = strip_accents(raw.lower())
        mapping = {
            "unidade": "un",
            "unidades": "un",
            "un": "un",
            "ramal": "un",
            "ramais": "un",
            "m": "m",
            "m2": "m2",
            "mÂ²": "m2",
            "vala": "vala",
            "valas": "vala",
            "servico": "servico",
            "servicos": "servico",
            "servico executado": "servico",
        }
        return mapping.get(raw, raw)

    def _normalize_obs(self, line: str) -> str:
        line = re.sub(r"^(Obs|ObservaÃ§Ã£o geral|Observacao geral)\s*:\s*", "", line, flags=re.I)
        line = line.strip(" -")
        return line.strip()

    def _classify_occurrence(self, desc: str) -> str:
        d = strip_accents(desc.lower())
        if any(k in d for k in ["chuva", "tempestade", "temporal"]):
            return "clima"
        if "vistoria" in d:
            return "vistoria"
        if any(k in d for k in ["vazamento", "reparo"]):
            return "reparo"
        if any(k in d for k in ["contas altas", "direcionada", "atividade paralela"]):
            return "atividade_paralela"
        if any(k in d for k in ["espera", "compressor", "mangueira", "material"]):
            return "restricao_operacional"
        return "ocorrencia_operacional"

    def _impact_level(self, desc: str) -> str:
        d = strip_accents(desc.lower())
        if any(k in d for k in ["nao foi possivel", "impossibilitando", "interrompida", "interrompido", "paralisadas", "sem producao"]):
            return "sim"
        if any(k in d for k in ["impactada", "impactando", "parcial"]):
            return "parcial"
        return "parcial"


class ReportGenerator:
    def __init__(self, dictionary: ServiceDictionary):
        self.dictionary = dictionary

    def generate_nucleus_reports(self, parsed: dict, output_dir: Path) -> None:
        output_dir.mkdir(parents=True, exist_ok=True)
        by_nucleus: Dict[str, dict] = defaultdict(lambda: {"items": [], "fronts": [], "occurrences": []})
        fronts_by_id = {f["id_frente"]: f for f in parsed["frentes"]}

        for item in parsed["execucao"]:
            front = fronts_by_id[item["id_frente"]]
            bucket = by_nucleus[front["nucleo"]]
            bucket["items"].append(item)
            if front not in bucket["fronts"]:
                bucket["fronts"].append(front)

        for occ in parsed["ocorrencias"]:
            front = fronts_by_id[occ["id_frente"]]
            bucket = by_nucleus[front["nucleo"]]
            bucket["occurrences"].append(occ)
            if front not in bucket["fronts"]:
                bucket["fronts"].append(front)

        for nucleus, data in by_nucleus.items():
            md = self._build_markdown_report(parsed, nucleus, data)
            safe = slugify(nucleus)
            (output_dir / f"{safe}.md").write_text(md, encoding="utf-8")
            self._write_docx(parsed, nucleus, data, output_dir / f"{safe}.docx")
            self._write_pdf(md, output_dir / f"{safe}.pdf")

    def _aggregate_items(self, items: List[dict]) -> List[dict]:
        agg: Dict[Tuple[str, str, str, str, str, str], dict] = {}
        for item in items:
            key = (
                item["item_normalizado"],
                item["unidade"],
                item.get("material", ""),
                item.get("especificacao", ""),
                item.get("tipo_registro", ""),
                item.get("observacao_item", ""),
            )
            if key not in agg:
                agg[key] = item.copy()
            else:
                agg[key]["quantidade"] += item["quantidade"]
        return list(agg.values())

    def _build_markdown_report(self, parsed: dict, nucleus: str, data: dict) -> str:
        date_ref = parsed["data_referencia"]
        contrato = parsed["contrato"]
        programa = parsed["programa"]
        activity_lines = [self._technical_sentence(v) for v in self._aggregate_items(data["items"])]
        analysis = self._build_analysis(data["fronts"], data["occurrences"], data["items"])
        obs = [o["descricao"] for o in data["occurrences"]]
        lines = [
            "# RELATÓRIO TÉCNICO DE EVOLUÇÃO DE OBRA",
            "",
            f"**Contrato:** {contrato}",
            f"**Programa:** {programa}",
            f"**Data de referência:** {date_ref}",
            f"**Núcleo:** {nucleus}",
            "",
            "## Objetivo",
            f"Apresentar o registro das atividades executadas no núcleo {nucleus}, com base nas informações operacionais consolidadas do dia {date_ref}, no âmbito do Contrato {contrato} – Programa {programa}.",
            "",
            "## Introdução",
            f"As informações a seguir refletem a evolução das frentes de serviço executadas no núcleo {nucleus}, considerando os quantitativos reportados em campo, bem como as ocorrências operacionais que impactaram o andamento das atividades.",
            "",
            "## Atividades executadas",
        ]
        if activity_lines:
            lines.extend([f"- {x}" for x in activity_lines])
        else:
            lines.append("- Não houve quantitativos executados no período.")
        lines.extend(["", "## Análise crítica", analysis])
        if obs:
            lines.extend(["", "## Observações operacionais"])
            lines.extend([f"- {o}" for o in obs])
        lines.extend([
            "",
            "## Conclusão",
            f"Conclui-se que as atividades registradas no período contribuíram para o avanço operacional do núcleo {nucleus}, observadas as condicionantes de campo relatadas pelas equipes e os impactos operacionais identificados no consolidado diário.",
        ])
        return "\n".join(lines)

    def _clean_report_complement(self, item: dict) -> str:
        complement = (item.get("complemento") or "").strip()
        if not complement:
            return ""
        text = complement
        if item.get("material"):
            text = re.sub(rf"\b{re.escape(item['material'])}\b", "", text, flags=re.I)
        if item.get("especificacao"):
            text = re.sub(re.escape(item["especificacao"]), "", text, flags=re.I)
        text = re.sub(r"\b\d+[\.,]?\d*\s*(m²|m2|m|unidades|unidade|un|valas|vala|servi[cç]os?|ramais?)\b", "", text, flags=re.I)
        obs = (item.get("observacao_item") or "").strip()
        if obs:
            text = re.sub(re.escape(obs), "", text, flags=re.I)
            text = re.sub(rf"\({re.escape(obs)}\)", "", text, flags=re.I)
        text = re.sub(r"\s+", " ", text).strip(" -;,.")
        normalized = strip_accents(text.lower())
        if normalized in {"", "com mureta", "servico executado", "servico", "executado", "realizada", "realizado"}:
            return ""
        return text

    def _technical_sentence(self, item: dict) -> str:
        service = self.dictionary.match(item.get("item_normalizado") or item["item_original"])
        phrase = repair_mojibake(service.get("frase_tecnica") or f"Execução de {item['item_normalizado'].lower()}")
        qty_value = float(item["quantidade"])
        qty = int(qty_value) if qty_value.is_integer() else qty_value
        unit_text = self._unit_text(item["unidade"], qty_value)
        suffix_parts = []
        if item.get("material"):
            mat = item["material"]
            if item.get("especificacao"):
                mat += f" {item['especificacao']}"
            suffix_parts.append(f"em {mat}")
        elif item.get("especificacao"):
            suffix_parts.append(item["especificacao"])
        clean_complement = self._clean_report_complement(item)
        if clean_complement:
            suffix_parts.append(clean_complement)
        if item.get("observacao_item"):
            suffix_parts.append(item["observacao_item"])
        dedup = []
        seen = set()
        for part in suffix_parts:
            key = strip_accents(part.lower()).strip()
            if key and key not in seen:
                seen.add(key)
                dedup.append(part)
        suffix = ""
        if dedup:
            first, *rest = dedup
            suffix = f" {first}"
            if rest:
                suffix += f" ({'; '.join(rest)})"
        return repair_mojibake(f"{phrase}: {qty} {unit_text}{suffix}.")

    def _unit_text(self, unit: str, qty: float) -> str:
        singular = float(qty) == 1.0
        mapping = {
            "un": "unidade" if singular else "unidades",
            "m": "metro" if singular else "metros",
            "vala": "vala" if singular else "valas",
            "servico": "serviço" if singular else "serviços",
            "m2": "m²",
        }
        return mapping.get(unit, unit)

    def _build_analysis(self, fronts: List[dict], occurrences: List[dict], items: List[dict]) -> str:
        total_fronts = len(fronts)
        productive_fronts = sum(1 for f in fronts if f.get("status_frente") == "com_producao")
        partial_fronts = sum(1 for f in fronts if f.get("status_frente") == "producao_parcial")
        stopped_fronts = sum(1 for f in fronts if f.get("status_frente") in {"sem_producao", "paralisada"})
        occ_text = " ".join(o["descricao"].lower() for o in occurrences)
        occ_norm = strip_accents(occ_text)

        if total_fronts and productive_fronts == 0 and partial_fronts == 0:
            return "Não houve produção no período, em razão de intercorrências operacionais devidamente registradas, sem caracterização de ociosidade injustificada das equipes mobilizadas."

        if any(k in occ_norm for k in ["chuva", "tempestade", "temporal"]):
            return f"O núcleo apresentou avanço operacional parcial no período, com produção registrada em {productive_fronts + partial_fronts} frente(s). Contudo, as condições climáticas adversas impactaram o ritmo previsto dos serviços, especialmente nas frentes com atividades de escavação, abertura de vala e execução a céu aberto."

        if any(k in occ_norm for k in ["vistoria", "vazamento", "reparo", "contas altas"]):
            return "Houve produção no período, porém parte da capacidade operacional foi direcionada para atividades complementares de vistoria, verificação e atendimento corretivo em campo, reduzindo o potencial de avanço físico das frentes originalmente programadas."

        if items and all(i.get("tipo_registro") in {"apoio_operacional", "insumo"} or i.get("unidade") == "servico" for i in items):
            return "As atividades registradas no período corresponderam predominantemente a etapas preparatórias, apoio operacional e serviços acessórios, importantes para a continuidade e a viabilização das etapas executivas subsequentes."

        if productive_fronts + partial_fronts > 0:
            if stopped_fronts > 0:
                return f"O núcleo apresentou produção em {productive_fronts + partial_fronts} frente(s), embora {stopped_fronts} frente(s) tenham registrado restrições ou paralisações pontuais. De modo geral, os quantitativos consolidados demonstram avanço físico compatível com as condições operacionais observadas no período."
            return "A produção consolidada do núcleo foi compatível com as atividades programadas para o período, indicando avanço físico da obra e regularidade operacional das frentes executivas mobilizadas."

        return "Não foram identificados quantitativos válidos para análise no período."

    def _set_cell_shading(self, cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    def _write_docx(self, parsed: dict, nucleus: str, data: dict, path: Path) -> None:
        date_ref = parsed["data_referencia"]
        contrato = parsed["contrato"]
        programa = parsed["programa"]
        activity_items = self._aggregate_items(data["items"])
        analysis = self._build_analysis(data["fronts"], data["occurrences"], data["items"])

        doc = Document()
        sec = doc.sections[0]
        sec.top_margin = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(2.1)
        sec.right_margin = Cm(2.1)

        styles = doc.styles
        styles["Normal"].font.name = "Arial"
        styles["Normal"].font.size = Pt(10)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("RELATÓRIO TÉCNICO DE EVOLUÇÃO DE OBRA")
        r.bold = True
        r.font.size = Pt(15)
        r.font.color.rgb = RGBColor(0, 82, 155)

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = sub.add_run("Sistema de Gestão de Relatórios de Campo")
        rr.italic = True
        rr.font.size = Pt(9)
        rr.font.color.rgb = RGBColor(90, 90, 90)

        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        table.autofit = False
        widths = [Cm(7.0), Cm(8.5)]
        info = [
            ("Contrato", contrato),
            ("Programa", programa),
            ("Data de referência", date_ref),
            ("Núcleo", nucleus),
        ]
        idx = 0
        for row in table.rows:
            for cell in row.cells:
                cell.width = widths[idx % 2]
                label, value = info[idx]
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                r1 = p.add_run(f"{label}: ")
                r1.bold = True
                p.add_run(repair_mojibake(value))
                self._set_cell_shading(cell, "EAF3FB")
                idx += 1

        doc.add_paragraph("")
        for title, text in [
            (
                "Objetivo",
                f"Apresentar o registro das atividades executadas no núcleo {nucleus}, com base nas informações operacionais consolidadas do dia {date_ref}, no âmbito do Contrato {contrato} – Programa {programa}.",
            ),
            (
                "Introdução",
                f"As informações a seguir refletem a evolução das frentes de serviço executadas no núcleo {nucleus}, considerando os quantitativos reportados em campo, bem como as ocorrências operacionais que impactaram o andamento das atividades.",
            ),
        ]:
            hp = doc.add_paragraph()
            rr = hp.add_run(title)
            rr.bold = True
            rr.font.size = Pt(11)
            rr.font.color.rgb = RGBColor(0, 82, 155)
            body = doc.add_paragraph(repair_mojibake(text))
            body.paragraph_format.space_after = Pt(8)

        hp = doc.add_paragraph()
        rr = hp.add_run("Atividades executadas")
        rr.bold = True
        rr.font.size = Pt(11)
        rr.font.color.rgb = RGBColor(0, 82, 155)

        if activity_items:
            for sentence in [self._technical_sentence(v) for v in activity_items]:
                bullet = doc.add_paragraph(style="List Bullet")
                bullet.paragraph_format.space_after = Pt(2)
                bullet.add_run(repair_mojibake(sentence))
        else:
            doc.add_paragraph("Não houve quantitativos executados no período.")

        hp = doc.add_paragraph()
        rr = hp.add_run("Análise crítica")
        rr.bold = True
        rr.font.size = Pt(11)
        rr.font.color.rgb = RGBColor(0, 82, 155)
        doc.add_paragraph(repair_mojibake(analysis))

        if data["occurrences"]:
            hp = doc.add_paragraph()
            rr = hp.add_run("Observações operacionais")
            rr.bold = True
            rr.font.size = Pt(11)
            rr.font.color.rgb = RGBColor(0, 82, 155)
            for occ in data["occurrences"]:
                bullet = doc.add_paragraph(style="List Bullet")
                bullet.add_run(repair_mojibake(occ["descricao"]))

        hp = doc.add_paragraph()
        rr = hp.add_run("Conclusão")
        rr.bold = True
        rr.font.size = Pt(11)
        rr.font.color.rgb = RGBColor(0, 82, 155)
        doc.add_paragraph(
            f"Conclui-se que as atividades registradas no período contribuíram para o avanço operacional do núcleo {nucleus}, observadas as condicionantes de campo relatadas pelas equipes e os impactos operacionais identificados no consolidado diário."
        )

        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "Relatório gerado automaticamente pelo Sistema de Gestão de Relatórios de Obra"
        doc.save(str(path))

    def _write_pdf(self, markdown_text: str, path: Path) -> None:
        markdown_text = repair_mojibake(markdown_text)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("CustomTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=16, textColor="#00529B", spaceAfter=10)
        h2 = ParagraphStyle("CustomH2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=11.5, textColor="#00529B", spaceBefore=10, spaceAfter=4)
        body = ParagraphStyle("CustomBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.2, leading=12)
        story = []
        for line in markdown_text.splitlines():
            if line.startswith("# "):
                story.append(Paragraph(line[2:], title_style))
            elif line.startswith("## "):
                story.append(Paragraph(line[3:], h2))
            elif line.startswith("- "):
                story.append(Paragraph(f"• {line[2:]}", body))
            elif line.strip():
                story.append(Paragraph(line.replace("**", ""), body))
            else:
                story.append(Spacer(1, 0.18 * cm))
        doc = SimpleDocTemplate(str(path), pagesize=A4, leftMargin=2 * cm, rightMargin=2 * cm, topMargin=1.7 * cm, bottomMargin=1.7 * cm)
        doc.build(story)


def save_parsed_outputs(parsed: dict, output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    with (output_dir / "relatorio_consolidado.json").open("w", encoding="utf-8") as f:
        json.dump(parsed, f, ensure_ascii=False, indent=2)

    def write_csv(name: str, rows: List[dict]):
        csv_path = output_dir / name
        if not rows:
            # Evita contaminacao por arquivo antigo quando reprocessa a mesma pasta de saida.
            csv_path.write_text("", encoding="utf-8-sig")
            return
        cleaned_rows = []
        for row in rows:
            clean_row = dict(row)
            clean_row["equipe"] = extrair_primeira_equipe(clean_row.get("equipe", ""))
            cleaned_rows.append(normalize_row_numbers(clean_row))
        with csv_path.open("w", encoding="utf-8-sig", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=list(cleaned_rows[0].keys()))
            writer.writeheader()
            writer.writerows(cleaned_rows)

    write_csv("frentes.csv", parsed["frentes"])
    write_csv("execucao.csv", parsed["execucao"])
    write_csv("ocorrencias.csv", parsed["ocorrencias"])
    write_csv("observacoes.csv", parsed.get("observacoes", []))
    write_csv("servico_nao_mapeado.csv", parsed.get("servicos_nao_mapeados", []))

    def _exec_scope_key(item: dict) -> str:
        id_frente = str(item.get("id_frente", "") or "").strip()
        if id_frente:
            return f"id:{id_frente}"
        return "ctx:{0}|{1}|{2}|{3}".format(
            str(item.get("data_referencia", "") or "").strip(),
            str(item.get("nucleo", "") or "").strip(),
            extrair_primeira_equipe(item.get("equipe", "")),
            str(item.get("logradouro", "") or "").strip(),
        )

    def _norm_text(value: object) -> str:
        return strip_accents(str(value or "").lower()).strip()

    def _service_text(item: dict) -> str:
        raw = item.get("servico_bruto") or item.get("item_original") or item.get("item_normalizado")
        return _norm_text(raw)

    def _category(item: dict) -> str:
        raw = item.get("categoria_item") or item.get("categoria")
        return _norm_text(raw).replace(" ", "_")

    def _is_caixa_uma_total(item: dict) -> bool:
        if _category(item) != "caixa_uma":
            return False
        s = _service_text(item)
        if not s:
            return False
        return ("caixa" in s and "uma" in s) and ("embutid" not in s and "mureta" not in s and "reboco" not in s)

    def _is_caixa_uma_breakdown(item: dict) -> bool:
        if _category(item) != "caixa_uma":
            return False
        s = _service_text(item)
        return any(token in s for token in ("embutid", "mureta", "reboco"))

    caixa_uma_total_scopes = {_exec_scope_key(e) for e in parsed.get("execucao", []) if _is_caixa_uma_total(e)}

    def _is_countable_exec(item: dict) -> bool:
        if _is_caixa_uma_breakdown(item) and _exec_scope_key(item) in caixa_uma_total_scopes:
            return False
        return True

    execucao_contabil = [e for e in parsed.get("execucao", []) if _is_countable_exec(e)]

    resumo = []
    por_nucleo = defaultdict(lambda: {
        "frentes": 0,
        "itens": 0,
        "ocorrencias": 0,
        "quantidade_total": 0.0,
        "quantidade_servicos": 0.0,
        "quantidade_insumos": 0.0,
    })

    def _nucleo_label(row: dict) -> str:
        nuc = str(row.get("nucleo", "") or "").strip()
        if not nuc:
            return ""
        return nuc

    for f in parsed["frentes"]:
        por_nucleo[_nucleo_label(f)]["frentes"] += 1
    for e in execucao_contabil:
        nuc_key = _nucleo_label(e)
        por_nucleo[nuc_key]["itens"] += 1
        qty = float(e.get("quantidade", 0) or 0)
        por_nucleo[nuc_key]["quantidade_total"] += qty
        if e.get("tipo_registro") == "insumo":
            por_nucleo[nuc_key]["quantidade_insumos"] += qty
        else:
            por_nucleo[nuc_key]["quantidade_servicos"] += qty
    for o in parsed["ocorrencias"]:
        por_nucleo[_nucleo_label(o)]["ocorrencias"] += 1
    for nucleo, vals in por_nucleo.items():
        resumo.append({"nucleo": nucleo, **vals})
    write_csv("resumo_nucleos.csv", resumo)

    indicadores = []
    por_categoria = defaultdict(lambda: {"registros": 0, "quantidade_total": 0.0})
    for e in execucao_contabil:
        chave = (e.get("nucleo", ""), e.get("categoria_item", ""), e.get("unidade", ""), e.get("tipo_registro", ""))
        por_categoria[chave]["registros"] += 1
        por_categoria[chave]["quantidade_total"] += float(e.get("quantidade", 0) or 0)
    for (nucleo, categoria, unidade, tipo), vals in por_categoria.items():
        indicadores.append({
            "data_referencia": parsed.get("data_referencia", ""),
            "contrato": parsed.get("contrato", ""),
            "programa": parsed.get("programa", ""),
            "nucleo": nucleo,
            "categoria_item": categoria,
            "unidade": unidade,
            "tipo_registro": tipo,
            "registros": vals["registros"],
            "quantidade_total": vals["quantidade_total"],
        })
    write_csv("indicadores_dashboard.csv", indicadores)

    indicadores_ocorrencias = []
    por_ocorrencia = defaultdict(int)
    for o in parsed["ocorrencias"]:
        chave = (o.get("nucleo", ""), o.get("tipo_ocorrencia", ""), o.get("impacto_producao", ""))
        por_ocorrencia[chave] += 1
    for (nucleo, tipo, impacto), total in por_ocorrencia.items():
        indicadores_ocorrencias.append({
            "data_referencia": parsed.get("data_referencia", ""),
            "contrato": parsed.get("contrato", ""),
            "programa": parsed.get("programa", ""),
            "nucleo": nucleo,
            "tipo_ocorrencia": tipo,
            "impacto_producao": impacto,
            "ocorrencias": total,
        })
    write_csv("indicadores_ocorrencias.csv", indicadores_ocorrencias)

    painel_geral = [{
        "data_referencia": parsed.get("data_referencia", ""),
        "contrato": parsed.get("contrato", ""),
        "programa": parsed.get("programa", ""),
        "frentes_total": len(parsed.get("frentes", [])),
        "itens_total": len(execucao_contabil),
        "ocorrencias_total": len(parsed.get("ocorrencias", [])),
        "quantidade_total": sum(float(e.get("quantidade", 0) or 0) for e in execucao_contabil),
        "frentes_sem_producao": sum(1 for f in parsed.get("frentes", []) if f.get("status_frente") in {"sem_producao", "paralisada"}),
    }]
    write_csv("painel_geral.csv", painel_geral)





























